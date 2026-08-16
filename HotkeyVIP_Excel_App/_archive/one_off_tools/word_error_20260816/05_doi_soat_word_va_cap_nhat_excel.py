# -*- coding: utf-8 -*-
"""Đối soát Word phục hồi và cập nhật trạng thái trong Excel gốc."""

import re
import shutil
from datetime import datetime
from pathlib import Path

import openpyxl
import xlwings as xw
from docx import Document

APP_ROOT = Path(__file__).resolve().parents[3]
MAPPING = APP_ROOT / "outputs" / "worker_profile_word_errors_20260816" / "word_error_worker_mapping.xlsx"
SOURCE = Path(r"D:\CodexProjects\Hotkeyvip\04_excel\hotkeyvip_test.xlsm")
SHEET_NAME = "VIET_BAI"
MIN_WORDS = 80


def word_count(path):
    try:
        doc = Document(path)
        parts = [paragraph.text for paragraph in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return len(re.findall(r"\S+", "\n".join(parts)))
    except Exception:
        return 0


def read_mapping():
    book = openpyxl.load_workbook(MAPPING, read_only=False, data_only=True)
    sheet = book["WORD_ERROR theo Worker"]
    header_row = None
    headers = {}
    for row in sheet.iter_rows(min_row=1, max_row=20, values_only=True):
        values = [str(value or "").strip() for value in row]
        if "Dòng Excel" in values and "Đường dẫn Word" in values:
            header_row = row
            headers = {value: index for index, value in enumerate(values) if value}
            break
    if header_row is None:
        raise RuntimeError("Không tìm thấy tiêu đề file mapping")

    jobs = []
    start = list(sheet.iter_rows(min_row=1, max_row=20, values_only=True)).index(header_row) + 2
    for values in sheet.iter_rows(min_row=start, values_only=True):
        excel_row = values[headers["Dòng Excel"]]
        word_path = str(values[headers["Đường dẫn Word"]] or "").strip()
        keyword = str(values[headers["Từ khóa"]] or "").strip()
        if excel_row and word_path:
            jobs.append({
                "excel_row": int(excel_row),
                "word_path": word_path,
                "keyword": keyword,
            })
    book.close()
    return jobs


def main():
    jobs = read_mapping()
    valid = []
    invalid = []
    for job in jobs:
        count = word_count(job["word_path"])
        job["actual_words"] = count
        if count >= MIN_WORDS:
            valid.append(job)
        else:
            invalid.append(job)

    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = SOURCE.with_name(f"{SOURCE.stem}_backup_truoc_cap_nhat_word_{stamp}{SOURCE.suffix}")
    shutil.copy2(SOURCE, backup)

    app = xw.App(visible=False, add_book=False)
    app.display_alerts = False
    app.screen_updating = False
    book = None
    try:
        book = app.books.open(str(SOURCE), update_links=False, read_only=False)
        sheet = book.sheets[SHEET_NAME]
        last_col = sheet.used_range.last_cell.column
        header_values = sheet.range((1, 1), (1, last_col)).value
        headers = {
            str(value or "").strip(): index + 1
            for index, value in enumerate(header_values)
            if str(value or "").strip()
        }
        required = ["Từ khóa", "Trạng thái viết", "Lỗi viết", "Đường dẫn Word"]
        missing = [name for name in required if name not in headers]
        if missing:
            raise RuntimeError("Excel gốc thiếu cột: " + ", ".join(missing))

        updated = 0
        for job in valid:
            row = job["excel_row"]
            source_keyword = str(sheet.cells(row, headers["Từ khóa"]).value or "").strip()
            if source_keyword != job["keyword"]:
                raise RuntimeError(
                    f"Dòng {row} đã đổi nhận dạng: mapping={job['keyword']!r}, Excel={source_keyword!r}"
                )
            sheet.cells(row, headers["Đường dẫn Word"]).value = job["word_path"]
            sheet.cells(row, headers["Trạng thái viết"]).value = "OK"
            sheet.cells(row, headers["Lỗi viết"]).value = ""
            updated += 1

        for job in invalid:
            row = job["excel_row"]
            source_keyword = str(sheet.cells(row, headers["Từ khóa"]).value or "").strip()
            if source_keyword != job["keyword"]:
                raise RuntimeError(f"Dòng {row} đã đổi nhận dạng, không cập nhật")
            sheet.cells(row, headers["Trạng thái viết"]).value = "WORD_ERROR"

        book.save()
        print(f"Đã cập nhật OK: {updated}")
        print(f"Giữ WORD_ERROR: {len(invalid)}")
        for job in invalid:
            print(
                f"WORD_ERROR dòng {job['excel_row']}: {job['keyword']} "
                f"({job['actual_words']} từ)"
            )
        print(f"Backup: {backup}")
    finally:
        if book is not None:
            book.close()
        app.quit()


if __name__ == "__main__":
    main()
