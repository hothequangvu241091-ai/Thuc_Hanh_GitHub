# -*- coding: utf-8 -*-
"""
FILE TỔNG: VIẾT BÀI CHATGPT -> LƯU WORD -> XIN BRIEF ẢNH -> TẠO ẢNH GEMINI

Dùng lại file Excel cũ, không cần đổi cấu trúc cột.
Có thể chạy song song với 3 file lẻ cũ. Code tổng sẽ kiểm tra thiếu gì thì làm phần đó.

CỘT EXCEL ĐANG DÙNG:
B  = Web / thư mục website
C  = Tên bài / keyword / tên file
D  = Prompt viết bài
F  = Link GPT gốc
G  = Path Word
H  = URL chat bài viết
I  = Status bài viết
J  = Lỗi bài viết
K  = Số từ Word (chỉ lưu số nguyên)
O  = BRIEF_1
P  = BRIEF_2
Q  = Status brief
T  = Status ảnh 1
U  = Path ảnh 1
V  = Status ảnh 2
W  = Path ảnh 2
X  = Done tổng
Y  = URL Gemini ảnh 1
Z  = URL Gemini ảnh 2
AA = Mốc thủ công: nhập đúng "OK OK" để lần sau chạy từ dòng kế tiếp

AB = Retry Count / số lần bảo hiểm
AC = Retry Step / bước đang retry
AD = Retry Error / mã lỗi ngắn
AE = Retry Time / thời điểm ghi lỗi
"""

# =====================================================
# V1.5: HOÀN THIỆN BẢO HIỂM 3 LỚP
# - Giữ nguyên luồng viết bài / Word / Brief / Gemini.
# - Bổ sung checkpoint và phân loại lỗi để dễ theo dõi.
# - Truyền Edge driver mới về process_row và main sau khi Lớp 2 reset Edge.
# - Giữ tham chiếu driver đang hoạt động kể cả khi lỗi phát sinh giữa chừng.
# =====================================================
VERSION = "V2.0_vietbai_3cap_baohiem_anh_khong_ghi_de"

ERROR_CODES = {
    "GPT": "GPT_ERROR",
    "WORD": "WORD_ERROR",
    "BRIEF": "BRIEF_ERROR",
    "IMAGE": "IMAGE_ERROR",
}


import os
import ctypes
import queue
import re
import threading
import time
import winsound
import shutil
import tempfile
import base64
import psutil
from ctypes import wintypes
from html.parser import HTMLParser
import xlwings as xw
from PIL import Image
from docx import Document

from selenium import webdriver
from selenium.webdriver.edge.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import StaleElementReferenceException

# Các HWND Edge thuộc lần chạy hiện tại, dùng cho nút Hiện/Ẩn trên bảng tiến độ.
_HIDDEN_EDGE_HANDLES = set()

# Khóa cấp phát tên ảnh: tránh hai luồng cùng chọn một đường dẫn khi tên bài trùng.
_IMAGE_PATH_LOCK = threading.RLock()
_RESERVED_IMAGE_BASES = set()

# =====================================================
# CẤU HÌNH EXCEL / FILE
# =====================================================
WORD_MACRO = "FullProcess_AllSteps"
SHEET_NAME = "VIET_BAI"
START_ROW = 2
END_ROW = 10000
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ROOT_OUTPUT = os.path.join(
    PROJECT_ROOT,
    "07_ket_qua",
    "bai_viet",
)
TEMP_DOWNLOAD_DIR = r"D:\autodangky\hotkey\GeminiDownloads"

COL_WEB = "B"
COL_NAME = "C"
COL_PROMPT = "D"
COL_GPT_URL = "F"
COL_WORD_PATH = "G"
COL_CHAT_URL = "H"
COL_ARTICLE_STATUS = "I"
COL_ARTICLE_ERROR = "J"
COL_WORD_COUNT = "K"

COL_BRIEF_1 = "O"
COL_BRIEF_2 = "P"
COL_BRIEF_STATUS = "Q"

COL_STATUS_IMG1 = "T"
COL_PATH_IMG1 = "U"
COL_STATUS_IMG2 = "V"
COL_PATH_IMG2 = "W"
COL_DONE = "X"
COL_GEMINI_URL_IMG1 = "Y"
COL_GEMINI_URL_IMG2 = "Z"
COL_MANUAL_MARK = "AA"
MANUAL_MARK_TEXT = "OK OK"

# Cột ghi chú lỗi phụ, không ảnh hưởng logic code cũ.
# Không ghi lỗi vào G/O/P/U/W để code cũ vẫn tự chạy tiếp được.
COL_RETRY_COUNT = "AB"
COL_RETRY_STEP = "AC"
COL_RETRY_ERROR = "AD"
COL_RETRY_TIME = "AE"

# Chỉ ánh xạ Excel theo tên tiêu đề. Phần xử lý bài viết giữ nguyên code cũ.
COLUMN_HEADER_BINDINGS = {
    "COL_WEB": "Tên Miền",
    "COL_NAME": "Từ khóa",
    "COL_PROMPT": "Prompt viết bài",
    "COL_GPT_URL": "URL GPT gốc",
    "COL_WORD_PATH": "Đường dẫn Word",
    "COL_CHAT_URL": "URL ChatGPT",
    "COL_ARTICLE_STATUS": "Trạng thái viết",
    "COL_ARTICLE_ERROR": "Lỗi viết",
    "COL_WORD_COUNT": "Số từ Word",
    "COL_BRIEF_1": "Brief ảnh 1",
    "COL_BRIEF_2": "Brief ảnh 2",
    "COL_BRIEF_STATUS": "Trạng thái brief",
    "COL_STATUS_IMG1": "Trạng thái ảnh 1",
    "COL_PATH_IMG1": "Đường dẫn ảnh 1",
    "COL_STATUS_IMG2": "Trạng thái ảnh 2",
    "COL_PATH_IMG2": "Đường dẫn ảnh 2",
    "COL_DONE": "Trạng thái hoàn tất",
    "COL_GEMINI_URL_IMG1": "URL Gemini ảnh 1",
    "COL_GEMINI_URL_IMG2": "URL Gemini ảnh 2",
    "COL_MANUAL_MARK": "Mốc bắt đầu",
    "COL_RETRY_COUNT": "Số lần thử lại",
    "COL_RETRY_STEP": "Bước thử lại",
    "COL_RETRY_ERROR": "Lỗi thử lại",
    "COL_RETRY_TIME": "Thời gian thử lại",
}

_RESOLVED_WORKBOOK = None
_LAST_STABLE_ARTICLE = None
_ACTIVE_DRIVER = None
_ACTIVE_WAIT = None


def remember_active_driver(driver, wait):
    """Lưu phiên Edge mới nhất để không mất session khi Lớp 2 reset Edge."""
    global _ACTIVE_DRIVER, _ACTIVE_WAIT
    _ACTIVE_DRIVER = driver
    _ACTIVE_WAIT = wait
    return driver, wait


def get_active_driver(driver=None, wait=None):
    """Lấy phiên Edge còn hoạt động mới nhất, có fallback về đối số hiện tại."""
    return _ACTIVE_DRIVER or driver, _ACTIVE_WAIT or wait

# Thời gian bảo hiểm theo yêu cầu
ARTICLE_WAIT_SECONDS = 600
ARTICLE_RELOAD_WAIT_SECONDS = 20
BRIEF_WAIT_SECONDS = 40
BRIEF_RELOAD_WAIT_SECONDS = 10
GEMINI_IMAGE_WAIT_SECONDS = 60
GEMINI_RELOAD_WAIT_SECONDS = 20

STATUS_OK = "OK"
STATUS_RUNNING = "RUNNING"
STATUS_ERROR = "ERROR"
STATUS_OK_BRIEF = "OK BRIEF ẢNH"
STATUS_RUNNING_BRIEF = "ĐANG TẠO BRIEF ẢNH"
STATUS_ERROR_BRIEF = "LỖI BRIEF ẢNH"
STATUS_SENT_IMG1 = "Đã gửi Gemini ảnh 1"
STATUS_SAVED_IMG1 = "Đã lưu ảnh 1 Gemini"
STATUS_SENT_IMG2 = "Đã gửi Gemini ảnh 2"
STATUS_SAVED_IMG2 = "Đã lưu ảnh 2 Gemini"

MIN_WORDS = 700

# Nếu đủ số bài Word mới liên tiếp có số từ dưới ngưỡng này,
# chương trình sẽ nghỉ trước khi bắt đầu bài tiếp theo.
SHORT_ARTICLE_WORD_LIMIT = 1300
SHORT_ARTICLE_PAUSE_MINUTES = 15
SHORT_ARTICLE_STREAK_LIMIT = 2

# Worker Recycling: Tắt/mở lại Edge sau bao nhiêu bài hoàn thành để xả RAM
RECYCLE_EVERY_N_ROWS = 30

MAX_IMAGE_SIZE = 800
GEMINI_URL = "https://gemini.google.com/app"

# Bản thử một dòng: luôn tạo lại đúng tên Word cũ, không sinh thêm file đánh số.
TEST_OVERWRITE_WORD = False

# Lần đầu để False nhằm quan sát selector menu tải ảnh. Sau khi test ổn có thể
# đổi thành True để toàn bộ Edge Selenium chạy ẩn.
RUN_HEADLESS = False

SECOND_PROMPT = """Các file ĐẦU VÀO đã được cung cấp đầy đủ trong PROMT của tôi
=> Hãy kiểm tra lại thật kỹ để thực hiện đúng"""

ASK_BRIEF_PROMPT = """
Nhiệm vụ:

Đọc toàn bộ bài viết trong cuộc trò chuyện hiện tại.

Không tạo ảnh.

Hãy tạo 2 mô tả cảnh dùng để tạo ảnh minh họa cho bài viết.

Cách lấy nội dung:

Mô tả cảnh 1:
- Dựa trên toàn bộ phần từ H1 đến ngay trước H2 đầu tiên.
- Đây là ảnh mở đầu của bài viết.

Mô tả cảnh 2:
- Dựa trên toàn bộ phần H2 đầu tiên đến ngay trước H2 tiếp theo.
- Đây là ảnh minh họa cho kiến thức trong H2 đầu tiên.

Yêu cầu:

Không mô tả lại tiêu đề.

Hãy đọc toàn bộ nội dung của từng section rồi xác định ý chính mà người đọc cần hình dung.

Từ ý chính đó, xây dựng một concept ảnh phù hợp.

Nếu hai section có nội dung gần giống nhau thì phải tạo hai concept khác nhau để thể hiện hai góc nhìn khác nhau của cùng chủ đề.

Không được chỉ thay đổi:

- Góc chụp
- Khoảng cách camera
- Bố cục

mà phải thay đổi cách truyền tải nội dung.

Mỗi mô tả cảnh cần có:

- Chủ thể hoặc đối tượng trung tâm
- Bối cảnh phù hợp
- Hành động hoặc trạng thái
- Không khí hoặc cảm xúc
- Các chi tiết nổi bật giúp người xem hiểu đúng nội dung section

Viết tự nhiên bằng tiếng Việt.

Không viết theo dạng checklist.

Không thêm giải thích.

Chỉ trả về:

===BRIEF_1===
...

===BRIEF_2===
...
""".strip()

# =====================================================
# EXCEL UTILS
# =====================================================
def get_sheet():
    global _RESOLVED_WORKBOOK
    app = xw.apps.active
    if not app:
        raise Exception("Vui lòng mở file Excel dữ liệu lên trước khi chạy code!")
    wb = app.books.active
    sh = wb.sheets[SHEET_NAME]
    workbook_key = (wb.fullname, sh.name)
    if _RESOLVED_WORKBOOK != workbook_key:
        resolve_columns_by_header(sh)
        _RESOLVED_WORKBOOK = workbook_key
    return wb, sh


def column_letter(number):
    result = ""
    while number:
        number, remainder = divmod(number - 1, 26)
        result = chr(65 + remainder) + result
    return result


def resolve_columns_by_header(sh):
    headers = {}
    last_col = sh.used_range.last_cell.column
    for col in range(1, last_col + 1):
        value = sh.range((1, col)).value
        if value is not None:
            headers[str(value).strip().casefold()] = column_letter(col)

    missing = []
    for global_name, header_name in COLUMN_HEADER_BINDINGS.items():
        letter = headers.get(header_name.casefold())
        if letter:
            globals()[global_name] = letter
        else:
            missing.append(header_name)

    if missing:
        raise Exception(
            "Sheet VIET_BAI thiếu tiêu đề: " + ", ".join(missing)
        )


def is_blank(value):
    return value is None or str(value).strip() == ""


def cell(row, col):
    _, sh = get_sheet()
    return sh.range(f"{col}{row}")


def safe_filename(name):
    name = str(name or "").strip()
    name = re.sub(r'[\\/:*?"<>|]', "-", name)
    name = re.sub(r"\s+", " ", name)
    return name[:160]


def make_output_path(web, file_name):
    folder = os.path.join(ROOT_OUTPUT, safe_filename(web))
    os.makedirs(folder, exist_ok=True)

    base_name = safe_filename(file_name)

    # Ưu tiên tên gốc
    path = os.path.join(folder, base_name + ".docx")
    if TEST_OVERWRITE_WORD:
        return path
    if not os.path.exists(path):
        return path

    # Nếu đã tồn tại thì thêm số phía sau
    i = 1
    while True:
        path = os.path.join(folder, f"{base_name} {i}.docx")

        if not os.path.exists(path):
            return path

        i += 1


def file_exists(path):
    return bool(path) and os.path.exists(str(path).strip())


def read_task(row):
    _, sh = get_sheet()
    return {
        "row": row,
        "web": str(sh.range(f"{COL_WEB}{row}").value or "").strip(),
        "name": str(sh.range(f"{COL_NAME}{row}").value or "").strip(),
        "prompt": str(sh.range(f"{COL_PROMPT}{row}").value or ""),
        "gpt_url": str(sh.range(f"{COL_GPT_URL}{row}").value or "").strip(),
        "word_path": str(sh.range(f"{COL_WORD_PATH}{row}").value or "").strip(),
        "chat_url": str(sh.range(f"{COL_CHAT_URL}{row}").value or "").strip(),
        "article_status": str(sh.range(f"{COL_ARTICLE_STATUS}{row}").value or "").strip(),
        "brief1": str(sh.range(f"{COL_BRIEF_1}{row}").value or "").strip(),
        "brief2": str(sh.range(f"{COL_BRIEF_2}{row}").value or "").strip(),
        "path_img1": str(sh.range(f"{COL_PATH_IMG1}{row}").value or "").strip(),
        "path_img2": str(sh.range(f"{COL_PATH_IMG2}{row}").value or "").strip(),
        "done": str(sh.range(f"{COL_DONE}{row}").value or "").strip().upper(),
    }


def get_last_row():
    _, sh = get_sheet()
    last_row = sh.range(f"{COL_NAME}{sh.cells.last_cell.row}").end("up").row
    return min(max(last_row, START_ROW), END_ROW)


def get_start_row_by_manual_mark():
    """
    Cột AA là mốc thủ công do người dùng tự nhập.
    Nếu có ô AA ghi đúng "OK OK", lấy dòng "OK OK" cuối cùng từ dưới lên
    rồi bắt đầu xử lý từ dòng kế tiếp. Code không tự ghi vào cột AA.
    Nếu không có mốc, giữ nguyên cách chạy cũ từ START_ROW.
    """
    _, sh = get_sheet()
    values = sh.range(
        f"{COL_MANUAL_MARK}{START_ROW}:{COL_MANUAL_MARK}{END_ROW}"
    ).options(ndim=1).value
    for index in range(len(values) - 1, -1, -1):
        row = START_ROW + index
        value = str(values[index] or "").strip()
        if value == MANUAL_MARK_TEXT:
            return row + 1, row
    return START_ROW, None


def write_value(row, col, value, save=True):
    wb, sh = get_sheet()
    target = sh.range(f"{col}{row}")
    target.value = value
    if col in {
        COL_ARTICLE_STATUS,
        COL_ARTICLE_ERROR,
        COL_BRIEF_STATUS,
        COL_DONE,
        COL_RETRY_COUNT,
        COL_RETRY_STEP,
        COL_RETRY_ERROR,
        COL_RETRY_TIME,
    }:
        target.api.WrapText = False
    if save:
        wb.save()


def write_retry_note(row, retry_count, step, error_code, detail=""):
    """
    Ghi chú lỗi phụ ở AB/AC/AD/AE.
    Không ghi vào các cột dữ liệu thật G/O/P/U/W để code cũ vẫn resume bình thường.
    """
    try:
        wb, sh = get_sheet()
        sh.range(f"{COL_RETRY_COUNT}{row}").value = retry_count
        sh.range(f"{COL_RETRY_STEP}{row}").value = step
        sh.range(f"{COL_RETRY_ERROR}{row}").value = f"{error_code}: {str(detail)[:180]}" if detail else error_code
        sh.range(f"{COL_RETRY_TIME}{row}").value = time.strftime("%Y-%m-%d %H:%M:%S")
        for col in (
            COL_RETRY_COUNT,
            COL_RETRY_STEP,
            COL_RETRY_ERROR,
            COL_RETRY_TIME,
        ):
            sh.range(f"{col}{row}").api.WrapText = False
        wb.save()
    except Exception as e:
        print(f"⚠️ Không ghi được retry note dòng {row}: {e}")


def write_image_final_error(row):
    """Ghi mã lỗi ảnh cuối cùng dạng ngắn; không chụp hoặc lưu ảnh màn hình lỗi."""
    try:
        _wb, sh = get_sheet()
        saved_error = str(
            sh.range(f"{COL_RETRY_ERROR}{row}").value or ""
        ).strip()
        if saved_error.startswith("GEMINI_"):
            return
    except Exception:
        pass
    write_retry_note(row, 9, "IMAGE", "IMAGE_FINAL_ERROR")


def reload_current_url(driver, wait_seconds, label=""):
    """
    Không bấm phím F5 thật. Load lại URL hiện tại bằng Selenium để ổn định hơn.
    """
    try:
        current_url = driver.current_url
        print(f"-> Load lại URL hiện tại {label}: {current_url}")
        driver.get(current_url)
        time.sleep(wait_seconds)
        return True
    except Exception as e:
        print(f"⚠️ Không load lại được URL hiện tại {label}: {e}")
        return False


class ArticleHTMLTextExtractor(HTMLParser):
    """Tách chữ để kiểm tra từ HTML gốc dùng cho Word, không sửa HTML."""

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.parts = []
        self.ignored_depth = 0

    def handle_starttag(self, tag, attrs):
        if tag in {"script", "style", "svg"}:
            self.ignored_depth += 1
        elif not self.ignored_depth and tag in {
            "p", "div", "br", "li", "tr", "td", "th",
            "h1", "h2", "h3", "h4", "h5", "h6",
        }:
            self.parts.append(" ")

    def handle_endtag(self, tag):
        if tag in {"script", "style", "svg"} and self.ignored_depth:
            self.ignored_depth -= 1
        elif not self.ignored_depth and tag in {
            "p", "div", "li", "tr", "td", "th",
            "h1", "h2", "h3", "h4", "h5", "h6",
        }:
            self.parts.append(" ")

    def handle_data(self, data):
        if not self.ignored_depth and data:
            self.parts.append(data)

    def text(self):
        return re.sub(r"\s+", " ", " ".join(self.parts)).strip()


def get_word_source_text(article_snapshot):
    """
    Lấy text từ chính HTML sẽ được copy_and_save_perfect() đưa vào Word.
    Chỉ tạo bản text trong bộ nhớ để kiểm tra; HTML gốc được giữ nguyên.
    """
    html_content = str((article_snapshot or {}).get("html") or "").strip()
    if not html_content:
        return ""
    parser = ArticleHTMLTextExtractor()
    parser.feed(html_content)
    parser.close()
    return parser.text()


def get_gpt_content_after_wait(driver, timeout_seconds, label=""):
    """
    Chờ ChatGPT xong trong thời gian giới hạn rồi lấy markdown cuối.
    Nếu timeout/lag thì trả None để lớp bảo hiểm xử lý tiếp.
    """
    try:
        wait_for_gpt_done(driver, max_timeout_seconds=timeout_seconds)
    except Exception as e:
        print(f"⚠️ {label} chờ GPT quá lâu hoặc lỗi: {e}")
    try:
        global _LAST_STABLE_ARTICLE
        _LAST_STABLE_ARTICLE = capture_stable_assistant_article(
            driver,
            stable_seconds=3.0,
            max_timeout_seconds=max(15, min(45, timeout_seconds)),
        )
        content = get_word_source_text(_LAST_STABLE_ARTICLE)
        if content and str(content).strip():
            return content
    except Exception as e:
        print(f"⚠️ {label} không lấy được nội dung GPT: {e}")
    return None


def write_article_running(row):
    wb, sh = get_sheet()
    sh.range(f"{COL_ARTICLE_STATUS}{row}").value = STATUS_RUNNING
    sh.range(f"{COL_ARTICLE_ERROR}{row}").value = ""
    sh.range(f"{COL_ARTICLE_STATUS}{row}").api.WrapText = False
    sh.range(f"{COL_ARTICLE_ERROR}{row}").api.WrapText = False


def write_article_error(row, text):
    wb, sh = get_sheet()
    sh.range(f"{COL_ARTICLE_STATUS}{row}").value = STATUS_ERROR
    sh.range(f"{COL_ARTICLE_ERROR}{row}").value = str(text)[:500]
    sh.range(f"{COL_ARTICLE_STATUS}{row}").api.WrapText = False
    sh.range(f"{COL_ARTICLE_ERROR}{row}").api.WrapText = False
    wb.save()


def write_article_success(row, word_path, chat_url, word_count):
    wb, sh = get_sheet()
    sh.range(f"{COL_WORD_PATH}{row}").value = word_path
    sh.range(f"{COL_CHAT_URL}{row}").value = chat_url
    sh.range(f"{COL_ARTICLE_STATUS}{row}").value = STATUS_OK
    sh.range(f"{COL_ARTICLE_ERROR}{row}").value = ""
    sh.range(f"{COL_WORD_COUNT}{row}").value = int(word_count)
    sh.range(f"{COL_ARTICLE_STATUS}{row}").api.WrapText = False
    sh.range(f"{COL_ARTICLE_ERROR}{row}").api.WrapText = False
    wb.save()


def mark_done_if_complete(row):
    task = read_task(row)
    if is_word_ok(task["word_path"]) and task["brief1"] and task["brief2"] and file_exists(task["path_img1"]) and file_exists(task["path_img2"]):
        write_value(row, COL_DONE, STATUS_OK)
        return True
    return False

# =====================================================
# KIỂM TRA WORD
# =====================================================
def count_words(text_content):
    if not text_content:
        return 0
    return len(re.findall(r"\S+", text_content.strip()))


def read_word_text(path):
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell_obj in row.cells:
                if cell_obj.text:
                    parts.append(cell_obj.text)
    return "\n".join(parts)


def is_word_ok(path, min_words=80):
    try:
        if not file_exists(path):
            return False
        if os.path.getsize(path) < 5 * 1024:
            return False
        text = read_word_text(path)
        return count_words(text) >= min_words
    except Exception:
        return False


def check_keyword_exists(text_content, keyword):
    return str(keyword).lower() in str(text_content).lower()

# =====================================================
# DRIVER
# =====================================================
def ensure_temp_download_dir():
    os.makedirs(TEMP_DOWNLOAD_DIR, exist_ok=True)


def add_background_running_options(options):
    """Giữ Chromium tiếp tục xử lý trang khi cửa sổ bị minimize/che khuất."""
    options.add_argument("--disable-background-timer-throttling")
    options.add_argument("--disable-backgrounding-occluded-windows")
    options.add_argument("--disable-renderer-backgrounding")
    options.add_argument("--disable-features=CalculateNativeWinOcclusion")


def keep_page_lifecycle_active(driver, quiet=False):
    """
    Yêu cầu Chromium giữ trang ở lifecycle active và giả lập trang vẫn có
    focus. Không gọi Page.bringToFront nên không tự mở/khôi phục cửa sổ Edge.
    """
    errors = []
    try:
        driver.execute_cdp_cmd(
            "Page.setWebLifecycleState",
            {"state": "active"},
        )
    except Exception as e:
        errors.append(f"lifecycle: {e}")

    try:
        driver.execute_cdp_cmd(
            "Emulation.setFocusEmulationEnabled",
            {"enabled": True},
        )
    except Exception as e:
        errors.append(f"focus: {e}")

    if errors and not quiet:
        print(
            "⚠️ Không giữ được đầy đủ trạng thái active/focus, "
            f"tiếp tục bình thường: {'; '.join(errors)}"
        )
    return not errors


def hide_edge_window_by_driver(driver, timeout=10):
    """Đưa Edge ra ngoài màn hình; không dùng trạng thái Windows SW_HIDE."""
    global _HIDDEN_EDGE_HANDLES
    service_process = getattr(
        getattr(driver, "service", None),
        "process",
        None,
    )
    if service_process is None:
        raise Exception("Không lấy được PID của EdgeDriver để chạy SW_HIDE.")

    root_pid = int(service_process.pid)
    process_ids = {root_pid}
    user32 = ctypes.windll.user32
    hidden_handles = set()
    deadline = time.time() + timeout
    callback_type = ctypes.WINFUNCTYPE(
        wintypes.BOOL,
        wintypes.HWND,
        wintypes.LPARAM,
    )

    while time.time() < deadline and not hidden_handles:
        candidates = []
        try:
            root_process = psutil.Process(root_pid)
            process_ids.update(
                child.pid for child in root_process.children(recursive=True)
            )
        except (psutil.Error, OSError):
            pass

        @callback_type
        def enum_callback(hwnd, _lparam):
            window_pid = wintypes.DWORD()
            user32.GetWindowThreadProcessId(
                hwnd,
                ctypes.byref(window_pid),
            )
            if int(window_pid.value) not in process_ids:
                return True

            class_buffer = ctypes.create_unicode_buffer(256)
            user32.GetClassNameW(hwnd, class_buffer, len(class_buffer))
            if class_buffer.value.startswith("Chrome_WidgetWin"):
                title_length = user32.GetWindowTextLengthW(hwnd)
                title_buffer = ctypes.create_unicode_buffer(
                    max(1, title_length + 1)
                )
                user32.GetWindowTextW(
                    hwnd,
                    title_buffer,
                    len(title_buffer),
                )
                rect = wintypes.RECT()
                user32.GetWindowRect(hwnd, ctypes.byref(rect))
                area = max(0, rect.right - rect.left) * max(
                    0,
                    rect.bottom - rect.top,
                )
                candidates.append(
                    {
                        "hwnd": int(hwnd),
                        "title": title_buffer.value.strip(),
                        "visible": bool(user32.IsWindowVisible(hwnd)),
                        "area": area,
                    }
                )
            return True

        user32.EnumWindows(enum_callback, 0)

        # Cửa sổ trình duyệt chính có tiêu đề trang. Các Chrome_WidgetWin
        # không tiêu đề là cửa sổ phụ nội bộ, không được hiện/ẩn cùng nút.
        main_candidates = [
            item for item in candidates
            if item["title"] and item["visible"] and item["area"] > 0
        ]
        if not main_candidates:
            main_candidates = [
                item for item in candidates
                if item["visible"] and item["area"] > 0
            ]

        if main_candidates:
            main_window = max(
                main_candidates,
                key=lambda item: item["area"],
            )
            virtual_right = (
                user32.GetSystemMetrics(76)
                + user32.GetSystemMetrics(78)
            )
            user32.ShowWindow(main_window["hwnd"], 9)  # SW_RESTORE
            user32.SetWindowPos(
                main_window["hwnd"],
                0,
                virtual_right + 100,
                0,
                0,
                0,
                0x0001 | 0x0004 | 0x0010,
            )
            hidden_handles.add(main_window["hwnd"])
        else:
            time.sleep(0.25)

    if not hidden_handles:
        raise Exception("Không tìm thấy cửa sổ Edge Selenium để đưa ra ngoài màn hình.")

    _HIDDEN_EDGE_HANDLES = set(hidden_handles)
    print(
        "-> Đã đưa Edge Selenium ra ngoài phạm vi màn hình."
    )


def set_current_edge_visible(visible):
    """Hiện hoặc ẩn lại các cửa sổ Edge Selenium của lần chạy hiện tại."""
    user32 = ctypes.windll.user32
    valid_handles = [
        hwnd for hwnd in _HIDDEN_EDGE_HANDLES
        if user32.IsWindow(hwnd)
    ]
    if not valid_handles:
        print("⚠️ Chưa tìm thấy cửa sổ Edge của lần chạy hiện tại.")
        return False

    for hwnd in valid_handles:
        user32.ShowWindow(hwnd, 9)  # SW_RESTORE
        if visible:
            user32.SetWindowPos(
                hwnd, 0, 60, 80, 0, 0, 0x0001 | 0x0004
            )
        else:
            virtual_right = (
                user32.GetSystemMetrics(76)
                + user32.GetSystemMetrics(78)
            )
            user32.SetWindowPos(
                hwnd,
                0,
                virtual_right + 100,
                0,
                0,
                0,
                0x0001 | 0x0004 | 0x0010,
            )

    if visible:
        try:
            user32.SetForegroundWindow(valid_handles[0])
        except Exception:
            pass
        print("-> Đã hiện lại Edge Selenium.")
    else:
        print("-> Đã đưa Edge Selenium ra ngoài phạm vi màn hình.")
    return True


def close_existing_edge_for_profile(profile_dir, timeout=8):
    """
    Đóng Edge Selenium còn sót đang dùng đúng profile của bản CHAY_HIDE.
    Không tác động Edge cá nhân hoặc worker dùng profile khác.
    """
    normalized_profile = os.path.normcase(
        os.path.abspath(profile_dir)
    )
    browser_roots = []

    for process in psutil.process_iter(["pid", "name", "cmdline"]):
        try:
            if (process.info.get("name") or "").lower() != "msedge.exe":
                continue

            command_parts = process.info.get("cmdline") or []
            command_line = " ".join(command_parts)
            normalized_command = os.path.normcase(command_line)
            uses_profile = (
                f"--user-data-dir={normalized_profile}" in normalized_command
                or f'--user-data-dir="{normalized_profile}"'
                in normalized_command
            )
            is_browser_root = not any(
                str(part).startswith("--type=")
                for part in command_parts
            )

            if uses_profile and is_browser_root:
                browser_roots.append(process)
        except (psutil.Error, OSError):
            continue

    if not browser_roots:
        return 0

    targets = []
    for root_process in browser_roots:
        try:
            targets.extend(root_process.children(recursive=True))
        except psutil.Error:
            pass
        targets.append(root_process)

    unique_targets = {
        process.pid: process for process in targets
    }
    for process in unique_targets.values():
        try:
            process.terminate()
        except psutil.Error:
            pass

    _, alive = psutil.wait_procs(
        list(unique_targets.values()),
        timeout=timeout,
    )
    for process in alive:
        try:
            process.kill()
        except psutil.Error:
            pass

    if alive:
        psutil.wait_procs(alive, timeout=3)

    print(
        "-> Đã đóng Edge Selenium cũ dùng profile CHAY_HIDE: "
        f"{len(unique_targets)} tiến trình."
    )
    return len(unique_targets)


def create_chatgpt_driver():
    options = Options()
    options.add_argument(r"--user-data-dir=D:\autodangky\hotkey\SeleniumData")
    options.add_argument("--disable-blink-features=AutomationControlled")
    options.add_argument("--remote-debugging-port=9222")
    options.add_argument("--disable-notifications")
    options.add_argument("--start-maximized")
    add_background_running_options(options)
    driver = webdriver.Edge(options=options)
    wait = WebDriverWait(driver, 45)
    return driver, wait


def create_gemini_driver():
    ensure_temp_download_dir()
    options = Options()
    options.add_argument(r"--user-data-dir=D:\autodangky\hotkey\SeleniumData")
    options.add_argument("--disable-blink-features=AutomationControlled")
    options.add_argument("--disable-notifications")
    options.add_argument("--start-maximized")
    add_background_running_options(options)
    prefs = {
        "download.default_directory": TEMP_DOWNLOAD_DIR,
        "download.prompt_for_download": False,
        "download.directory_upgrade": True,
        "safebrowsing.enabled": True,
        "profile.default_content_setting_values.automatic_downloads": 1,
    }
    options.add_experimental_option("prefs", prefs)
    driver = webdriver.Edge(options=options)
    wait = WebDriverWait(driver, 45)
    return driver, wait


def create_shared_driver():
    """
    Mở duy nhất 1 cửa sổ Edge cho toàn bộ quá trình.
    Driver này dùng chung cho cả ChatGPT và Gemini để tránh đóng/mở trình duyệt
    ở mỗi dòng hoặc khi chuyển từ bước viết bài sang bước tạo ảnh.
    """
    ensure_temp_download_dir()
    selenium_profile = r"D:\autodangky\hotkey\SeleniumData"
    close_existing_edge_for_profile(selenium_profile)

    options = Options()
    options.add_argument(f"--user-data-dir={selenium_profile}")
    options.add_argument("--disable-blink-features=AutomationControlled")
    options.add_argument("--disable-notifications")
    options.add_argument("--start-maximized")
    add_background_running_options(options)
    if RUN_HEADLESS:
        options.add_argument("--headless=new")
        options.add_argument("--window-size=1920,1080")
    prefs = {
        "download.default_directory": TEMP_DOWNLOAD_DIR,
        "download.prompt_for_download": False,
        "download.directory_upgrade": True,
        "safebrowsing.enabled": True,
        "profile.default_content_setting_values.automatic_downloads": 1,
    }
    options.add_experimental_option("prefs", prefs)
    driver = webdriver.Edge(options=options)
    if not RUN_HEADLESS:
        hide_edge_window_by_driver(driver)
    wait = WebDriverWait(driver, 45)
    return remember_active_driver(driver, wait)

# =====================================================
# CHATGPT: GỬI PROMPT / LẤY NỘI DUNG / LƯU WORD
# =====================================================
def get_chatgpt_input_box(driver, wait, timeout=25):
    """
    Lấy đúng ô nhập prompt thật của ChatGPT ở dưới cùng.
    Tránh dán/gửi nhầm vào Article Detail / Edit / Canvas vì các vùng đó cũng có contenteditable/ProseMirror.
    """
    start = time.time()
    last_error = None

    while time.time() - start < timeout:
        try:
            # Chỉ dùng id prompt-textarea. Không dùng XPath contenteditable rộng.
            candidates = driver.find_elements(By.XPATH, "//*[@id='prompt-textarea']")
            valid_boxes = []

            win_h = driver.execute_script("return window.innerHeight || 0;")

            for el in candidates:
                try:
                    if not el.is_displayed():
                        continue

                    rect = driver.execute_script("""
                        const r = arguments[0].getBoundingClientRect();
                        return {
                            top: r.top,
                            bottom: r.bottom,
                            left: r.left,
                            right: r.right,
                            width: r.width,
                            height: r.height
                        };
                    """, el)

                    # Bỏ qua editor ở phần nội dung bài viết nếu nó nằm quá cao.
                    if rect["bottom"] < win_h * 0.50:
                        continue

                    # Bỏ qua phần tử quá nhỏ / không phải ô nhập.
                    if rect["width"] < 200 or rect["height"] < 15:
                        continue

                    valid_boxes.append((rect["top"], el))
                except Exception as e:
                    last_error = e
                    continue

            if valid_boxes:
                # Nếu có nhiều phần tử cùng id, lấy phần tử thấp nhất trên màn hình.
                valid_boxes.sort(key=lambda x: x[0], reverse=True)
                return valid_boxes[0][1]

        except Exception as e:
            last_error = e

        time.sleep(0.5)

    raise Exception(f"Không tìm thấy ô nhập prompt thật ở dưới cùng. Lỗi gần nhất: {last_error}")


def focus_chatgpt_input_safely(driver, wait, timeout=25):
    """
    Lấy đúng ô nhập ChatGPT thật ở dưới cùng. Việc gửi dùng JavaScript trực
    tiếp vào chatbox nên không bắt buộc document.activeElement phải là hộp.
    get_chatgpt_input_box vẫn kiểm tra selector, hiển thị, kích thước và vị trí
    để loại trừ vùng Article Detail/Edit.
    """
    chatbox = get_chatgpt_input_box(driver, wait, timeout=timeout)

    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", chatbox)
    time.sleep(0.5)

    # Focus chỉ để giao diện cập nhật thuận lợi; thất bại focus không phải lỗi
    # vì nội dung sẽ được gán trực tiếp bằng JavaScript vào đúng chatbox.
    driver.execute_script("arguments[0].focus();", chatbox)
    time.sleep(0.3)

    return chatbox


def get_chatgpt_send_button(driver, wait, timeout=35):
    """
    Tìm nút gửi prompt ChatGPT có bảo hiểm riêng:
    - đợi UI render
    - xử lý khi GPT đang generate
    - thử selector mới/cũ
    - refresh URL hiện tại 1 lần nếu không thấy nút gửi
    """

    def find_button_once():
        return driver.execute_script("""
            const input = document.activeElement && document.activeElement.id === 'prompt-textarea'
                ? document.activeElement
                : document.querySelector('#prompt-textarea');

            if (!input) return null;

            // Nếu ChatGPT còn đang trả lời thì chưa tìm nút Send
            const stopBtn = document.querySelector(
                'button[aria-label="Stop generating"], button[data-testid="stop-button"]'
            );
            if (stopBtn) return "GPT_STILL_GENERATING";

            const form = input.closest('form');

            const selectors = [
                'button[data-testid="send-button"]',
                'button[aria-label="Send prompt"]',
                'button[aria-label="Send message"]',
                'button[type="submit"]',
                'button svg[data-testid="send-button"]'
            ];

            if (form) {
                for (const sel of selectors) {
                    let b = form.querySelector(sel);
                    if (b) {
                        if (b.tagName.toLowerCase() === 'svg') b = b.closest('button');
                        const r = b.getBoundingClientRect();
                        const st = getComputedStyle(b);
                        if (r.width > 0 && r.height > 0 && st.display !== 'none' && st.visibility !== 'hidden') {
                            return b;
                        }
                    }
                }
            }

            // Fallback: lấy button thấp nhất gần composer
            const buttons = Array.from(document.querySelectorAll('button')).filter(b => {
                const r = b.getBoundingClientRect();
                const st = getComputedStyle(b);
                const html = (b.outerHTML || '').toLowerCase();
                const label = (
                    b.getAttribute('aria-label') ||
                    b.getAttribute('data-testid') ||
                    b.innerText ||
                    ''
                ).toLowerCase();

                const looksSend =
                    label.includes('send') ||
                    label.includes('gửi') ||
                    html.includes('send-button') ||
                    html.includes('arrow-up') ||
                    html.includes('submit');

                return looksSend &&
                    r.width > 0 &&
                    r.height > 0 &&
                    st.display !== 'none' &&
                    st.visibility !== 'hidden';
            });

            if (!buttons.length) return null;

            buttons.sort((a, b) => b.getBoundingClientRect().top - a.getBoundingClientRect().top);
            return buttons[0];
        """)

    def wait_find(label, seconds):
        end_time = time.time() + seconds
        last_state = None

        while time.time() < end_time:
            try:
                btn = find_button_once()

                if btn == "GPT_STILL_GENERATING":
                    last_state = "GPT vẫn đang generate"
                    time.sleep(2)
                    continue

                if btn:
                    return btn

                last_state = "Chưa thấy nút Send"
            except Exception as e:
                last_state = str(e)

            time.sleep(0.7)

        raise Exception(f"{label}: Không tìm thấy nút gửi prompt của ChatGPT. Trạng thái cuối: {last_state}")

    try:
        return wait_find("Lần 1", timeout)

    except Exception as first_error:
        print(f"⚠️ Không thấy nút gửi lần 1: {first_error}")
        print("-> Bảo hiểm: load lại URL hiện tại rồi tìm lại nút gửi.")

        try:
            current_url = driver.current_url
            driver.get(current_url)
            time.sleep(10)

            # focus lại ô nhập sau khi reload
            focus_chatgpt_input_safely(driver, wait, timeout=25)

            return wait_find("Sau reload", 25)

        except Exception as second_error:
            raise Exception(
                "Không tìm thấy nút gửi prompt của ChatGPT sau khi đã reload URL. "
                f"Lỗi đầu: {first_error} | Lỗi sau reload: {second_error}"
            )


def send_prompt_by_js(driver, wait, text_to_send):
    """
    Gửi prompt bằng JS nhưng vẫn bắt buộc chọn đúng ô nhập ChatGPT.
    Đã bỏ selector contenteditable/ProseMirror rộng để tránh dán nhầm vào Article Detail/Edit.
    """
    chatbox = focus_chatgpt_input_safely(driver, wait, timeout=25)

    tag_name = (chatbox.tag_name or "").lower()
    if tag_name == "textarea":
        driver.execute_script("""
            arguments[0].value = arguments[1];
            arguments[0].dispatchEvent(new Event('input', { bubbles: true }));
            arguments[0].dispatchEvent(new Event('change', { bubbles: true }));
        """, chatbox, text_to_send)
    else:
        # Với div/p contenteditable id=prompt-textarea.
        driver.execute_script("""
            arguments[0].focus();
            arguments[0].innerText = arguments[1];
            arguments[0].dispatchEvent(new InputEvent('input', {
                bubbles: true,
                inputType: 'insertText',
                data: arguments[1]
            }));
        """, chatbox, text_to_send)

    time.sleep(1)

    send_button = get_chatgpt_send_button(driver, wait, timeout=20)
    driver.execute_script("arguments[0].removeAttribute('disabled');", send_button)
    time.sleep(0.2)
    driver.execute_script("arguments[0].click();", send_button)
    print("-> Đã gửi prompt bằng JS vào đúng ô nhập ChatGPT.")
    time.sleep(3)


def send_prompt_by_real_paste(driver, wait, text_to_send):
    """
    ĐÃ ĐỔI SANG JS THUẦN (không còn dùng Clipboard hệ điều hành, không còn
    giả lập Ctrl+V/Enter thật). Giữ nguyên tên hàm để không phải sửa nơi gọi.

    Lý do đổi: pyperclip dùng chung 1 Clipboard cho toàn máy. Khi chạy nhiều
    worker song song, worker này pyperclip.copy() có thể đè nội dung ngay
    lúc worker khác chuẩn bị Ctrl+V, gây dán nhầm bài. Cách cũ cũng bắt buộc
    cửa sổ Edge phải đang có focus thật tại đúng thời điểm bấm phím.
    """
    chatbox = focus_chatgpt_input_safely(driver, wait, timeout=25)

    tag_name = (chatbox.tag_name or "").lower()
    if tag_name == "textarea":
        driver.execute_script("""
            arguments[0].value = arguments[1];
            arguments[0].dispatchEvent(new Event('input', { bubbles: true }));
            arguments[0].dispatchEvent(new Event('change', { bubbles: true }));
        """, chatbox, text_to_send)
    else:
        driver.execute_script("""
            arguments[0].focus();
            arguments[0].innerText = arguments[1];
            arguments[0].dispatchEvent(new InputEvent('input', {
                bubbles: true,
                inputType: 'insertText',
                data: arguments[1]
            }));
        """, chatbox, text_to_send)
    print("-> Đã gõ prompt bằng JS vào đúng ô nhập ChatGPT dưới cùng.")

    time.sleep(1)

    send_button = get_chatgpt_send_button(driver, wait, timeout=20)
    driver.execute_script("arguments[0].removeAttribute('disabled');", send_button)
    time.sleep(0.2)
    driver.execute_script("arguments[0].click();", send_button)
    print("-> Đã bấm gửi prompt.")
    time.sleep(2)


def normalize_prompt_for_compare(text):
    """Chuẩn hóa khoảng trắng để so prompt trong code với tin nhắn trên DOM."""
    return re.sub(r"\s+", " ", str(text or "")).strip().casefold()


def prompt_already_sent(driver, text_to_send):
    """Kiểm tra các tin nhắn user gần nhất để chống gửi trùng prompt."""
    expected = normalize_prompt_for_compare(text_to_send)
    if not expected:
        return False
    try:
        messages = driver.execute_script("""
            const direct = Array.from(
                document.querySelectorAll('[data-message-author-role="user"]')
            );
            const fromTurns = Array.from(
                document.querySelectorAll('div[data-testid^="conversation-turn-"]')
            ).map(turn => turn.querySelector('[data-message-author-role="user"]'))
             .filter(Boolean);
            const unique = Array.from(new Set([...direct, ...fromTurns]));
            return unique.slice(-8).map(el => (el.innerText || el.textContent || '').trim());
        """) or []
    except Exception:
        return False

    for message in messages:
        actual = normalize_prompt_for_compare(message)
        if not actual:
            continue
        if actual == expected:
            return True
        # ChatGPT có thể chèn nhãn/nút phụ vào cùng khối DOM. Chỉ cần toàn bộ
        # prompt nằm trong tin nhắn user là đã gửi; không bắt buộc bằng tuyệt đối.
        if len(expected) >= 40 and expected in actual:
            return True
        # DOM đôi lúc thêm/bớt ít ký tự định dạng. Với prompt dài, so cả
        # đầu và cuối để tránh nhận nhầm một tin nhắn khác.
        if len(expected) >= 240 and len(actual) >= 240:
            if expected[:120] == actual[:120] and expected[-120:] == actual[-120:]:
                return True
    return False


def wait_prompt_send_signal(driver, text_to_send, timeout=10):
    """Quan sát dấu hiệu gửi thành công nhưng không dùng nó để tự gửi lại."""
    deadline = time.time() + timeout
    while time.time() < deadline:
        if prompt_already_sent(driver, text_to_send):
            return "prompt đã xuất hiện trong lịch sử chat"
        try:
            state = driver.execute_script("""
                const box = document.querySelector('#prompt-textarea');
                const boxText = box
                    ? ((box.value || box.innerText || box.textContent || '').trim())
                    : '';
                const stop = !!document.querySelector(
                    'button[aria-label="Stop generating"], '
                    + 'button[data-testid="stop-button"], .result-streaming'
                );
                return { boxEmpty: !!box && !boxText, generating: stop };
            """) or {}
            if state.get("generating"):
                return "ChatGPT đã bắt đầu trả lời"
            if state.get("boxEmpty"):
                return "ô nhập đã trống sau khi bấm gửi"
        except Exception:
            pass
        time.sleep(0.5)
    return ""


def send_once_unless_present(driver, wait, text_to_send, send_func, label):
    """Không bao giờ gửi lại chỉ vì bước quan sát DOM không nhận ra prompt."""
    if prompt_already_sent(driver, text_to_send):
        print(f"-> {label}: Prompt đã có trong cuộc trò chuyện. Không gửi trùng.")
        return
    send_func(driver, wait, text_to_send)
    signal = wait_prompt_send_signal(driver, text_to_send, timeout=10)
    if signal:
        print(f"-> {label}: Xác nhận gửi thành công ({signal}).")
    else:
        # send_func chỉ trả về sau khi đã tìm đúng nút Send và click xong.
        # Trạng thái DOM không rõ không được phép kích hoạt L1/L2, vì như vậy
        # có thể gửi cùng prompt 2-3 lần như lỗi của V1.7.
        print(
            f"⚠️ {label}: Đã click gửi nhưng DOM chưa cho dấu hiệu rõ. "
            "Coi là đã gửi và KHÔNG gửi lại để tránh trùng prompt."
        )


def send_prompt_with_3_layers(driver, wait, text_to_send, send_func, task=None):
    """
    HỆ THỐNG BẢO HIỂM 3 LỚP KHI GỬI PROMPT SANG CHATGPT.
    LỚP 0: Gửi bình thường.
    LỚP 1: Reload URL hiện tại rồi gửi lại.
    LỚP 2: Reset Edge, khôi phục đúng URL, chờ ô nhập tối đa 60 giây rồi gửi lại.
    LỚP 3: Ghi FINAL_ERROR và bỏ dòng.
    """
    errors = []

    try:
        send_once_unless_present(driver, wait, text_to_send, send_func, "LỚP 0")
        return driver, wait
    except Exception as e:
        errors.append(f"L0:{e}")
        print(f"⚠️ LỚP 0 lỗi: {e}")

    try:
        reload_current_url(driver, wait_seconds=8, label="Bảo hiểm Lớp 1")
        send_once_unless_present(driver, wait, text_to_send, send_func, "LỚP 1")
        print("-> LỚP 1 cứu thành công.")
        return driver, wait
    except Exception as e:
        errors.append(f"L1:{e}")
        print(f"⚠️ LỚP 1 lỗi: {e}")

    try:
        global _LAST_STABLE_ARTICLE
        _LAST_STABLE_ARTICLE = None

        try:
            current_url_before_reset = str(driver.current_url or "").strip()
        except Exception:
            current_url_before_reset = ""

        try:
            driver.quit()
        except Exception:
            pass

        time.sleep(3)
        driver, wait = create_shared_driver()

        task = task or {}
        chat_url = str(task.get("chat_url") or "").strip()
        gpt_url = str(task.get("gpt_url") or "").strip()
        candidates = [current_url_before_reset, chat_url, gpt_url]
        target_url = next(
            (url for url in candidates if url.lower().startswith(("http://", "https://"))),
            "",
        )
        if not target_url:
            raise Exception("LỚP 2 không có URL H/F hoặc URL hiện tại để khôi phục.")

        print(f"-> LỚP 2 mở lại URL công việc: {target_url}")
        driver.get(target_url)

        ready_deadline = time.time() + 60
        ready_box = None
        last_ready_error = None
        while time.time() < ready_deadline:
            try:
                ready_box = get_chatgpt_input_box(driver, wait, timeout=3)
                if ready_box is not None:
                    break
            except Exception as ready_error:
                last_ready_error = ready_error
            time.sleep(0.5)

        if ready_box is None:
            raise Exception(
                "LỚP 2: Sau 60 giây vẫn không tìm thấy ô nhập ChatGPT. "
                f"Lỗi gần nhất: {last_ready_error}"
            )

        send_once_unless_present(driver, wait, text_to_send, send_func, "LỚP 2")
        print("-> LỚP 2 cứu thành công.")
        return driver, wait

    except Exception as e:
        errors.append(f"L2:{e}")
        row = (task or {}).get("row", 0)
        write_retry_note(
            row,
            9,
            "FINAL_ERROR",
            "CHATGPT_3_LAYER_FAILED",
            " | ".join(errors)
        )
        raise Exception(
            "Bảo hiểm 3 lớp thất bại: " + " | ".join(errors)
        )


def wait_for_gpt_done(driver, max_timeout_seconds=300):
    start_time = time.time()
    while True:
        if time.time() - start_time > max_timeout_seconds:
            raise Exception(f"ChatGPT phản hồi quá lâu hoặc bị treo (> {max_timeout_seconds}s).")
        # Giữ renderer ChatGPT tiếp tục cập nhật DOM khi Edge bị minimize.
        # Không mở lại cửa sổ và không chiếm chuột.
        keep_page_lifecycle_active(driver, quiet=True)
        is_writing = driver.execute_script("""
            let stopBtn = document.querySelector('button[aria-label="Stop generating"], button[data-testid="stop-button"]');
            let streaming = document.querySelector('.result-streaming');
            return (stopBtn !== null || streaming !== null);
        """)
        if not is_writing:
            break
        time.sleep(2)
    time.sleep(2)


def capture_stable_assistant_article(driver, stable_seconds=3.0, max_timeout_seconds=45):
    """
    Chụp text + HTML của assistant cuối sau khi cả hai đã đứng yên.
    Dùng textContent thay cho innerText để không phụ thuộc cửa sổ ChatGPT
    đang được render/hiển thị hay đang minimize.
    """
    script = """
        const turns = Array.from(
            document.querySelectorAll('div[data-testid^="conversation-turn-"]')
        );
        const candidates = [];
        for (const turn of turns) {
            const assistant = turn.querySelector('[data-message-author-role="assistant"]');
            if (!assistant) continue;
            const markdown = assistant.querySelector('.markdown') || turn.querySelector('.markdown');
            if (markdown) candidates.push(markdown);
        }
        if (!candidates.length) {
            document.querySelectorAll(
                '[data-message-author-role="assistant"] .markdown, div.agent-turn .markdown'
            ).forEach(node => candidates.push(node));
        }
        if (!candidates.length) return null;

        const source = candidates[candidates.length - 1];
        const clone = source.cloneNode(true);
        clone.querySelectorAll(
            'button, svg, [data-testid*="copy"], [aria-label*="Copy"], [aria-label*="Sao chép"]'
        ).forEach(el => el.remove());
        const originalText = (clone.textContent || source.textContent || '').trim();
        return {text: originalText, html: clone.innerHTML || ''};
    """

    deadline = time.time() + max_timeout_seconds
    last_text = None
    last_html = None
    stable_since = None
    next_cdp_wakeup = 0.0

    while time.time() < deadline:
        # Nhắc lại active/focus định kỳ để React tiếp tục đưa nội dung vào DOM
        # ngay cả khi cửa sổ Edge đang minimize.
        now = time.time()
        if now >= next_cdp_wakeup:
            keep_page_lifecycle_active(driver, quiet=True)
            next_cdp_wakeup = now + 2.0

        article = driver.execute_script(script)
        text_value = str((article or {}).get("text") or "").strip()
        html_value = str((article or {}).get("html") or "").strip()

        if not text_value or not html_value:
            last_text = None
            last_html = None
            stable_since = None
            time.sleep(0.5)
            continue

        if text_value == last_text and html_value == last_html:
            if stable_since is None:
                stable_since = time.time()
            elif time.time() - stable_since >= stable_seconds:
                print(
                    f"-> Nội dung assistant đã ổn định {stable_seconds:.0f}s "
                    f"({count_words(text_value)} từ)."
                )
                return {"text": text_value, "html": html_value}
        else:
            last_text = text_value
            last_html = html_value
            stable_since = None

        time.sleep(0.5)

    raise Exception(
        f"Nội dung assistant chưa ổn định sau {max_timeout_seconds}s."
    )


def extract_content_by_js(driver):
    return driver.execute_script("""
        let blocks = document.querySelectorAll('div[data-testid*="-turn-assistant"] .markdown, div.agent-turn .markdown, .markdown');
        if (blocks.length > 0) return blocks[blocks.length - 1].textContent;
        return null;
    """)


def copy_and_save_perfect(driver, file_path):
    import win32com.client as win32

    def copy_save_once(article_snapshot):
        word_app = None
        doc = None
        html_path = None

        try:
            article = article_snapshot
            if not article or not article.get("html"):
                raise Exception("Không có snapshot HTML assistant đã ổn định.")
            if count_words(article.get("text")) < 80:
                raise Exception("HTML lấy được có quá ít nội dung.")

            html_document = f"""<!doctype html>
<html><head><meta charset="utf-8">
<style>
body {{ font-family: Arial, sans-serif; font-size: 11pt; line-height: 1.45; }}
h1 {{ font-size: 20pt; }} h2 {{ font-size: 16pt; }} h3 {{ font-size: 13pt; }}
p {{ margin: 0 0 6pt 0; }} li {{ margin-bottom: 3pt; }}
table {{ border-collapse: collapse; width: 100%; }}
th, td {{ border: 1px solid #999; padding: 5px; }}
</style></head><body>{article['html']}</body></html>"""

            output_folder = os.path.dirname(os.path.abspath(file_path))
            os.makedirs(output_folder, exist_ok=True)
            with tempfile.NamedTemporaryFile(
                mode="w", suffix=".html", prefix="chatgpt_",
                dir=output_folder, encoding="utf-8", delete=False
            ) as temp_html:
                temp_html.write(html_document)
                html_path = temp_html.name

            word_app = win32.gencache.EnsureDispatch('Word.Application')
            word_app.Visible = False
            word_app.DisplayAlerts = 0

            doc = word_app.Documents.Open(
                html_path, ConfirmConversions=False, ReadOnly=False,
                AddToRecentFiles=False, Encoding=65001
            )

            for para in doc.Paragraphs:
                para.SpaceAfter = 4
                para.SpaceBefore = 0

            word_app.Run(WORD_MACRO)

            if TEST_OVERWRITE_WORD and os.path.exists(file_path):
                os.remove(file_path)
            doc.SaveAs2(file_path, FileFormat=16)

            doc.Close()
            doc = None

        finally:
            if doc is not None:
                try:
                    doc.Close(False)
                except Exception:
                    pass

            if word_app is not None:
                try:
                    word_app.Quit()
                except Exception:
                    pass

            if html_path and os.path.exists(html_path):
                try:
                    os.remove(html_path)
                except Exception:
                    pass

    try:
        global _LAST_STABLE_ARTICLE
        if not _LAST_STABLE_ARTICLE:
            _LAST_STABLE_ARTICLE = capture_stable_assistant_article(driver)
        copy_save_once(_LAST_STABLE_ARTICLE)

    except Exception as first_error:
        print(f"⚠️ Copy/lưu Word lần 1 lỗi: {first_error}")
        print("-> Bảo hiểm: load lại URL hiện tại rồi thử copy/lưu Word thêm 1 lần.")

        try:
            current_url = driver.current_url
            driver.get(current_url)
            time.sleep(ARTICLE_RELOAD_WAIT_SECONDS)

            # Chờ ChatGPT ổn định lại rồi thử copy lần 2
            get_gpt_content_after_wait(driver, 10, "Copy Word sau reload")

            copy_save_once(_LAST_STABLE_ARTICLE)

        except Exception as second_error:
            raise Exception(
                "Copy/lưu Word lỗi sau khi đã reload URL và thử lại 1 lần. "
                f"Lỗi lần 1: {first_error} | Lỗi lần 2: {second_error}"
            )


def parse_briefs_logic(text):
    text = str(text or "").strip()
    m = re.search(r"===BRIEF_1===\s*(.*?)\s*===BRIEF_2===\s*(.*)", text, flags=re.I | re.S)
    if not m:
        raise Exception("Không tìm thấy đúng định dạng ===BRIEF_1=== và ===BRIEF_2===")
    brief1 = re.sub(r"^```.*?\n|\n```$", "", m.group(1).strip(), flags=re.S).strip()
    brief2 = re.sub(r"^```.*?\n|\n```$", "", m.group(2).strip(), flags=re.S).strip()
    if not brief1 or not brief2:
        raise Exception("Brief 1 hoặc Brief 2 bị rỗng.")
    return brief1, brief2

def save_word_from_current_chat(driver, task, label=""):
    row = task["row"]

    content_text = get_gpt_content_after_wait(driver, ARTICLE_WAIT_SECONDS, label)

    if not content_text:
        write_retry_note(row, 1, "ARTICLE_RESUME", "NO_CONTENT", "Không lấy được nội dung từ URL H.")
        reload_current_url(driver, ARTICLE_RELOAD_WAIT_SECONDS, "Resume URL H")
        content_text = get_gpt_content_after_wait(driver, 10, "Resume sau reload")

    if not content_text:
        return None, None, "Không lấy được nội dung từ URL hiện tại."

    word_count = count_words(content_text)
    print(f"-> Số từ lấy từ chat hiện tại: {word_count}")

    if word_count < MIN_WORDS:
        print(f"-> Nội dung dưới {MIN_WORDS} từ. Gửi prompt bổ sung trong chính URL hiện tại.")
        send_prompt_by_real_paste(driver, WebDriverWait(driver, 45), SECOND_PROMPT)
        content_text = get_gpt_content_after_wait(driver, ARTICLE_WAIT_SECONDS, "Kéo dài bài trong URL H")

        if not content_text:
            return None, None, "Không lấy được nội dung sau prompt bổ sung trong URL H."

        word_count = count_words(content_text)
        print(f"-> Số từ sau bổ sung: {word_count}")

        if word_count < MIN_WORDS:
            return None, None, f"Bài vẫn dưới {MIN_WORDS} từ: {word_count}"

    if not check_keyword_exists(content_text, task["name"]):
        return None, None, f"Nội dung không chứa từ khóa/tên file: {task['name']}"

    word_path = make_output_path(task["web"], task["name"])
    print("-> Đang lưu Word từ chat hiện tại...")
    copy_and_save_perfect(driver, word_path)

    if not is_word_ok(word_path):
        return None, None, "Word đã lưu nhưng kiểm tra lại không đạt."

    return word_path, word_count, None

def write_article_if_needed(driver, wait, task):
    row = task["row"]

    if is_word_ok(task["word_path"]) and not TEST_OVERWRITE_WORD:
        print(f"Dòng {row}: Word đã có và đọc được. Bỏ qua viết bài.")
        if task["article_status"].upper() != STATUS_OK:
            write_value(row, COL_ARTICLE_STATUS, STATUS_OK)
        return read_task(row), driver, wait

    if not task["web"] or not task["name"] or not task["prompt"]:
        raise Exception("Thiếu B/C/D nên không thể viết bài.")

    print(f"\n[DÒNG {row}] Bắt đầu xử lý tạo Word...")

    write_article_running(row)

    # =====================================================
    # ƯU TIÊN 1: Nếu cột H có URL thì mở H trước
    # =====================================================
    if task["chat_url"]:
        print(f"-> Cột H có URL. Ưu tiên mở lại chat cũ: {task['chat_url']}")

        try:
            driver.get(task["chat_url"])
            time.sleep(6)

            word_path, word_count, err = save_word_from_current_chat(driver, task, "Resume từ URL H")

            if word_path:
                write_article_success(row, word_path, driver.current_url, word_count)
                print(f"-> Dòng {row}: Đã lưu Word OK từ URL H.")
                return read_task(row), driver, wait

            print(f"⚠️ Không lưu được Word từ URL H: {err}")
            write_retry_note(row, 1, "ARTICLE_RESUME", "RESUME_H_FAILED", err)

        except Exception as e:
            print(f"⚠️ URL H lỗi, sẽ xét tạo mới từ F: {e}")
            write_retry_note(row, 2, "ARTICLE_RESUME", "URL_H_ERROR", e)

    # =====================================================
    # ƯU TIÊN 2: Chỉ tạo mới từ F khi H trống hoặc H lỗi
    # =====================================================
    if not task["gpt_url"]:
        raise Exception("Không có URL H dùng được và thiếu F nên không thể tạo mới.")

    print("-> Không dùng được URL H. Bắt đầu tạo mới từ URL F...")

    driver.get(task["gpt_url"])
    time.sleep(3)

    driver, wait = send_prompt_with_3_layers(driver, wait, task["prompt"], send_prompt_by_js, task=task)
    write_value(row, COL_CHAT_URL, driver.current_url)

    content_text = get_gpt_content_after_wait(driver, ARTICLE_WAIT_SECONDS, "Viết bài lần 1")

    if not content_text:
        write_retry_note(row, 1, "ARTICLE", "ARTICLE_EMPTY_RELOAD", "Không có nội dung sau lần chờ đầu, load lại URL hiện tại.")
        reload_current_url(driver, ARTICLE_RELOAD_WAIT_SECONDS, "ChatGPT bài viết")
        content_text = get_gpt_content_after_wait(driver, 5, "Viết bài sau load URL")

    if not content_text:
        write_retry_note(row, 2, "ARTICLE", "ARTICLE_RESEND_PROMPT", "Load lại URL vẫn không có nội dung, gửi lại prompt viết bài 1 lần.")
        driver, wait = send_prompt_with_3_layers(driver, wait, task["prompt"], send_prompt_by_js, task=task)
        write_value(row, COL_CHAT_URL, driver.current_url)
        content_text = get_gpt_content_after_wait(driver, ARTICLE_WAIT_SECONDS, "Viết bài gửi lại prompt")

    if not content_text:
        raise Exception("Không lấy được nội dung bài viết sau load URL + gửi lại prompt 1 lần.")

    word_count = count_words(content_text)
    print(f"-> Số từ ước tính bài viết: {word_count}")

    if word_count < MIN_WORDS:
        print(f"-> Bài dưới {MIN_WORDS} từ. Gửi prompt kéo dài bài.")
        send_prompt_by_real_paste(driver, wait, SECOND_PROMPT)
        content_text = get_gpt_content_after_wait(driver, ARTICLE_WAIT_SECONDS, "Viết bài prompt 2")

        if not content_text:
            raise Exception("Không lấy được nội dung sau prompt kéo dài.")

        word_count = count_words(content_text)
        print(f"-> Số từ sau kéo dài: {word_count}")

        if word_count < MIN_WORDS:
            raise Exception(f"Bài viết vẫn dưới {MIN_WORDS} từ sau prompt kéo dài: {word_count} từ.")

    if not check_keyword_exists(content_text, task["name"]):
        raise Exception(f"Bài viết không chứa từ khóa/tên file: {task['name']}")

    word_path = make_output_path(task["web"], task["name"])
    print("-> Đang lưu Word giữ định dạng...")
    copy_and_save_perfect(driver, word_path)

    if not is_word_ok(word_path):
        raise Exception("Word đã lưu nhưng kiểm tra lại thấy rỗng/lỗi/không đủ nội dung.")

    write_article_success(row, word_path, driver.current_url, word_count)
    print(f"-> Dòng {row}: Đã lưu Word OK từ URL mới.")
    return read_task(row), driver, wait

def try_get_briefs_from_current_answer(driver):
    raw_answer = extract_content_by_js(driver)
    if not raw_answer:
        raise Exception("Không lấy được phản hồi Brief.")

    brief1, brief2 = parse_briefs_logic(raw_answer)

    def invalid(text):
        text = str(text or "").strip()

        if not text:
            return True

        # chỉ toàn dấu chấm: ., .., ..., .....
        if text.replace(".", "").strip() == "":
            return True

        # quá ngắn thì coi như lỗi
        if len(text) < 20:
            return True

        return False

    if invalid(brief1):
        raise Exception("Brief 1 không hợp lệ.")

    if invalid(brief2):
        raise Exception("Brief 2 không hợp lệ.")

    return brief1, brief2


def brief_if_needed(driver, wait, task):
    row = task["row"]
    if task["brief1"] and task["brief2"]:
        print(f"Dòng {row}: Brief đã có. Bỏ qua xin Brief.")
        if str(cell(row, COL_BRIEF_STATUS).value or "").strip().upper() != STATUS_OK_BRIEF:
            write_value(row, COL_BRIEF_STATUS, STATUS_OK_BRIEF)
        return read_task(row), driver, wait

    if not is_word_ok(task["word_path"]):
        raise Exception("Chưa có Word hợp lệ nên không xin Brief.")

    print(f"\n[DÒNG {row}] Bắt đầu xin Brief ảnh...")
    write_value(row, COL_BRIEF_STATUS, STATUS_RUNNING_BRIEF, save=False)

    # Ưu tiên dùng tab ChatGPT hiện tại nếu vừa viết xong. Nếu đang resume thì mở lại URL H.
    current_url = driver.current_url or ""
    target_url = task["chat_url"]
    if target_url and target_url not in current_url:
        print("-> Đang mở lại URL chat cũ ở cột H để xin Brief.")
        driver.get(target_url)
        time.sleep(5)

    # LẦN 1: gửi prompt xin Brief bằng Bảo Hiểm 3 Lớp
    driver, wait = send_prompt_with_3_layers(driver, wait, ASK_BRIEF_PROMPT, send_prompt_by_real_paste, task=task)
    get_gpt_content_after_wait(driver, BRIEF_WAIT_SECONDS, "Brief lần 1")

    brief1 = brief2 = None
    first_error = None
    try:
        brief1, brief2 = try_get_briefs_from_current_answer(driver)
    except Exception as e:
        first_error = e
        print(f"⚠️ Brief lần 1 chưa đúng: {e}")

    # BẢO HIỂM 1: sai/rỗng thì load lại URL hiện tại, chờ 10s, parse lại
    if not brief1 or not brief2:
        write_retry_note(row, 1, "BRIEF", "BRIEF_RELOAD", first_error or "Brief rỗng/sai cấu trúc, load lại URL hiện tại.")
        reload_current_url(driver, BRIEF_RELOAD_WAIT_SECONDS, "ChatGPT Brief")
        try:
            brief1, brief2 = try_get_briefs_from_current_answer(driver)
        except Exception as e:
            print(f"⚠️ Brief sau load URL vẫn chưa đúng: {e}")
            first_error = e

    # BẢO HIỂM 2: vẫn sai thì gửi lại prompt Brief đúng 1 lần
    if not brief1 or not brief2:
        write_retry_note(row, 2, "BRIEF", "BRIEF_RESEND_PROMPT", first_error or "Load URL vẫn chưa có Brief đúng, gửi lại prompt Brief 1 lần.")
        send_prompt_by_real_paste(driver, wait, ASK_BRIEF_PROMPT)
        get_gpt_content_after_wait(driver, BRIEF_WAIT_SECONDS, "Brief gửi lại prompt")
        try:
            brief1, brief2 = try_get_briefs_from_current_answer(driver)
        except Exception as e:
            raise Exception(f"Không lấy được Brief đúng sau load URL + gửi lại prompt 1 lần: {e}")

    wb, sh = get_sheet()
    sh.range(f"{COL_BRIEF_1}{row}").value = brief1
    sh.range(f"{COL_BRIEF_2}{row}").value = brief2
    sh.range(f"{COL_BRIEF_STATUS}{row}").value = STATUS_OK_BRIEF
    sh.range(f"{COL_BRIEF_1}{row}").api.WrapText = False
    sh.range(f"{COL_BRIEF_2}{row}").api.WrapText = False
    wb.save()
    print(f"-> Dòng {row}: Đã lưu Brief vào O/P.")
    return read_task(row), driver, wait


# =====================================================
# GEMINI: TẠO ẢNH / DOWNLOAD
# =====================================================
def build_gemini_prompt(brief):
    brief = str(brief or "").strip()
    return f"""
Hãy tạo ra chính xác 1 hình minh họa chất lượng cao.

=========================
BRIEF ẢNH CHI TIẾT
=========================

{brief}

=========================
CÁCH TRÌNH BÀY HÌNH ẢNH
=========================

Thể hiện Brief bằng ngôn ngữ editorial tối giản dành cho website cao cấp.

Brief là nguồn nội dung chính. Cách trình bày hình ảnh chỉ điều chỉnh cách tổ chức, nhấn mạnh và thể hiện Brief; không được làm sai thông điệp hoặc bản chất sự việc trong Brief.

=========================
YÊU CẦU BẮT BUỘC VỀ ẢNH
=========================

- Tạo đúng 1 ảnh duy nhất.
- Ảnh ngang, tỷ lệ 4:3.
- Giữ đúng nội dung chính của Brief.
- Mọi chi tiết phải đúng logic thực tế.

Chỉ xuất ra ảnh, không giải thích gì thêm.
""".strip()


def get_output_folder(word_path):
    folder = os.path.dirname(str(word_path).strip())
    if not folder:
        raise Exception("Không lấy được thư mục từ cột G.")
    os.makedirs(folder, exist_ok=True)
    return folder


def get_unique_image_path(folder, base_name):
    """
    Cấp phát đường dẫn ảnh không trùng tên.

    Ví dụ:
    - Tên bài 1.png
    - Tên bài 1 2.png
    - Tên bài 1 3.png

    Kiểm tra đồng thời PNG/JPG/JPEG/WEBP và khóa trong RAM để các worker
    không thể cùng chọn một tên trước khi file ảnh được ghi xuống ổ đĩa.
    """
    folder = os.path.abspath(folder)
    base_name = safe_filename(base_name)
    supported_exts = (".png", ".jpg", ".jpeg", ".webp")

    with _IMAGE_PATH_LOCK:
        def name_is_available(candidate_base):
            reserved_key = os.path.normcase(
                os.path.join(folder, candidate_base)
            )
            if reserved_key in _RESERVED_IMAGE_BASES:
                return False
            return not any(
                os.path.exists(os.path.join(folder, candidate_base + ext))
                for ext in supported_exts
            )

        candidate_base = base_name
        if not name_is_available(candidate_base):
            i = 2
            while True:
                candidate_base = f"{base_name} {i}"
                if name_is_available(candidate_base):
                    break
                i += 1

        reserved_key = os.path.normcase(
            os.path.join(folder, candidate_base)
        )
        _RESERVED_IMAGE_BASES.add(reserved_key)
        return os.path.join(folder, candidate_base + ".png")

def list_download_files():
    ensure_temp_download_dir()
    files = set()
    for name in os.listdir(TEMP_DOWNLOAD_DIR):
        path = os.path.join(TEMP_DOWNLOAD_DIR, name)
        if not os.path.isfile(path):
            continue
        if name.endswith(".crdownload") or name.endswith(".tmp"):
            continue
        files.add(path)
    return files


def wait_new_download(before_files, timeout=180):
    print("-> Đang chờ file tải về...")
    start = time.time()
    while time.time() - start < timeout:
        now_files = list_download_files()
        candidates = list(now_files - before_files)
        if candidates:
            newest = max(candidates, key=os.path.getmtime)
            size1 = os.path.getsize(newest)
            time.sleep(1)
            size2 = os.path.getsize(newest)
            if size1 == size2 and size2 > 10 * 1024:
                return newest
        time.sleep(1)
    raise Exception("Không thấy file ảnh mới trong thư mục tải xuống.")


def resize_image_max_800(image_path, max_size=MAX_IMAGE_SIZE):
    try:
        with Image.open(image_path) as img:
            img.load()
            w, h = img.size
            longest = max(w, h)
            if longest <= max_size:
                return image_path
            ratio = max_size / float(longest)
            new_w = int(w * ratio)
            new_h = int(h * ratio)
            img = img.resize((new_w, new_h), Image.LANCZOS)
            ext = os.path.splitext(image_path)[1].lower()
            if ext in [".jpg", ".jpeg"]:
                if img.mode in ("RGBA", "P"):
                    img = img.convert("RGB")
                img.save(image_path, quality=92, optimize=True)
            else:
                img.save(image_path, optimize=True)
            print(f"-> Đã resize ảnh về tối đa {max_size}px: {new_w}x{new_h}")
            return image_path
    except Exception as e:
        print(f"⚠️ Không resize được ảnh, giữ file gốc: {e}")
        return image_path


def move_downloaded_image(downloaded_path, final_path):
    ext = os.path.splitext(downloaded_path)[1].lower()
    if ext not in [".png", ".jpg", ".jpeg", ".webp"]:
        ext = ".png"
    final_path = os.path.splitext(final_path)[0] + ext
    if os.path.exists(final_path):
        os.remove(final_path)
    shutil.move(downloaded_path, final_path)
    resize_image_max_800(final_path)
    return final_path


def save_image_directly_from_element(driver, img_element, final_path):
    """Đọc blob ảnh ngay trong trang Gemini; không phụ thuộc nút Download."""
    result = driver.execute_async_script("""
        const img = arguments[0];
        const done = arguments[arguments.length - 1];
        const src = img.currentSrc || img.src || '';
        if (!src) { done({ok:false, error:'Ảnh không có src'}); return; }
        if (!img.complete || !img.naturalWidth || !img.naturalHeight) {
            done({
                ok:false,
                error:'Thẻ ảnh chưa load/decode xong',
                complete:Boolean(img.complete),
                naturalWidth:img.naturalWidth || 0,
                naturalHeight:img.naturalHeight || 0
            });
            return;
        }

        // Đọc pixel đã decode bằng canvas. Cách này không cần fetch lại blob
        // URL vốn đã được kiểm chứng là có thể bị Gemini thu hồi.
        try {
            const canvas = document.createElement('canvas');
            canvas.width = img.naturalWidth;
            canvas.height = img.naturalHeight;
            const context = canvas.getContext('2d');
            context.drawImage(img, 0, 0);
            const data = canvas.toDataURL('image/png');
            if (data && data.length > 20000) {
                done({ok:true, data:data, type:'image/png', method:'canvas'});
                return;
            }
        } catch (canvasError) {
            // Nếu canvas bị giới hạn bảo mật thì thử fetch blob bên dưới.
        }

        fetch(src, {credentials:'include'})
            .then(response => {
                if (!response.ok) throw new Error('HTTP ' + response.status);
                return response.blob();
            })
            .then(blob => {
                const reader = new FileReader();
                reader.onloadend = () => done({
                    ok:true,
                    data:reader.result,
                    type:blob.type || ''
                });
                reader.onerror = () => done({ok:false, error:'FileReader lỗi'});
                reader.readAsDataURL(blob);
            })
            .catch(error => done({ok:false, error:String(error)}));
    """, img_element)

    if not result or not result.get("ok"):
        raise Exception("Không đọc trực tiếp được blob Gemini: " + str(result))

    data_url = result.get("data") or ""
    if "," not in data_url:
        raise Exception("Gemini trả về data URL không hợp lệ.")
    header, encoded = data_url.split(",", 1)
    mime = (result.get("type") or header).lower()
    if "jpeg" in mime or "jpg" in mime:
        ext = ".jpg"
    elif "webp" in mime:
        ext = ".webp"
    else:
        ext = ".png"

    final_path = os.path.splitext(final_path)[0] + ext
    image_bytes = base64.b64decode(encoded)
    if len(image_bytes) < 10 * 1024:
        raise Exception(f"Blob ảnh quá nhỏ: {len(image_bytes)} bytes.")

    with open(final_path, "wb") as image_file:
        image_file.write(image_bytes)

    # Mở thử bằng PIL để chắc chắn bytes vừa lưu là ảnh thật.
    try:
        with Image.open(final_path) as image:
            image.verify()
    except Exception:
        try:
            os.remove(final_path)
        except Exception:
            pass
        raise

    resize_image_max_800(final_path)
    print("-> Đã lưu ảnh trực tiếp từ blob Gemini, không cần mở viewer.")
    return final_path


def send_prompt_to_gemini(driver, wait, text_to_send):
    """
    ĐÃ ĐỔI SANG JS THUẦN: không còn pyperclip (Clipboard hệ điều hành dùng
    chung cho cả máy) và không còn send_keys Ctrl+A/Backspace/Ctrl+V/Enter
    thật (cần cửa sổ đang focus). Toàn bộ thao tác gõ + gửi đều qua
    execute_script nên chạy được cả khi cửa sổ không active và cả headless.
    """
    chatbox_xpath = (
        "//div[@contenteditable='true' and @role='textbox']"
        "|//rich-textarea//div[@contenteditable='true']"
        "|//div[contains(@class,'textarea') and @contenteditable='true']"
    )

    for attempt in range(1, 4):
        enter_started = False
        try:
            # Gemini có thể vừa hiện ô nhập đã render lại ngay. Chờ ngắn rồi
            # tìm lại để luôn thao tác trên phần tử mới nhất.
            wait.until(EC.element_to_be_clickable((By.XPATH, chatbox_xpath)))
            time.sleep(0.8)
            chatbox = wait.until(
                EC.element_to_be_clickable((By.XPATH, chatbox_xpath))
            )
            driver.execute_script(
                "arguments[0].scrollIntoView({block:'center'});", chatbox
            )

            # Gõ nội dung bằng JS thay vì Ctrl+A/Backspace/Ctrl+V thật.
            driver.execute_script("""
                const el = arguments[0];
                const text = arguments[1];
                el.focus();
                el.innerText = text;
                el.dispatchEvent(new InputEvent('input', {
                    bubbles: true,
                    inputType: 'insertText',
                    data: text
                }));
            """, chatbox, text_to_send)
            time.sleep(0.5)

            # Từ thời điểm gọi gửi, không tự gửi lại nếu trạng thái không rõ,
            # tránh Gemini nhận hai prompt giống nhau.
            enter_started = True

            # Ưu tiên bấm nút gửi thật bằng JS; chỉ dùng phím Enter giả lập
            # qua JS làm phương án dự phòng nếu không thấy nút gửi.
            send_clicked = driver.execute_script("""
                const el = arguments[0];
                const form = el.closest('form') || el.closest('rich-textarea') || document;
                const btn = form.querySelector(
                    'button[aria-label*="Send" i], button[aria-label*="Gửi" i], ' +
                    'button[aria-label*="gửi" i], button[type="submit"]'
                );
                if (btn && !btn.disabled) { btn.click(); return true; }
                return false;
            """, chatbox)

            if not send_clicked:
                driver.execute_script("""
                    const el = arguments[0];
                    const opts = {
                        key: 'Enter', code: 'Enter', keyCode: 13, which: 13,
                        bubbles: true, cancelable: true
                    };
                    el.dispatchEvent(new KeyboardEvent('keydown', opts));
                    el.dispatchEvent(new KeyboardEvent('keyup', opts));
                """, chatbox)

            print("-> Đã gửi prompt sang Gemini bằng JS.")
            time.sleep(3)
            return
        except StaleElementReferenceException:
            if enter_started:
                raise Exception(
                    "Ô nhập Gemini đổi đúng lúc gửi; dừng để tránh gửi trùng prompt."
                )
            print(
                f"⚠️ Ô nhập Gemini vừa được render lại, "
                f"đang tìm lại ({attempt}/3)..."
            )
            time.sleep(0.5)

    raise Exception("Ô nhập Gemini liên tục thay đổi sau 3 lần thử.")


def open_gemini_and_wait_ready(driver, timeout=30):
    """Mở Gemini và đi tiếp ngay khi ô nhập thật sự sẵn sàng."""
    started = time.time()
    driver.get(GEMINI_URL)
    keep_page_lifecycle_active(driver)
    chatbox_xpath = (
        "//div[@contenteditable='true' and @role='textbox']"
        "|//rich-textarea//div[@contenteditable='true']"
        "|//div[contains(@class,'textarea') and @contenteditable='true']"
    )
    WebDriverWait(driver, timeout).until(
        EC.element_to_be_clickable((By.XPATH, chatbox_xpath))
    )
    print(f"-> Gemini đã sẵn sàng sau {time.time() - started:.1f}s.")


def get_current_large_image_srcs(driver):
    srcs = set()
    imgs = driver.find_elements(By.XPATH, "//img")
    for img in imgs:
        try:
            src = img.get_attribute("src") or ""
            w = driver.execute_script("return arguments[0].naturalWidth || 0;", img)
            h = driver.execute_script("return arguments[0].naturalHeight || 0;", img)
            if src and w >= 250 and h >= 250:
                srcs.add(src)
        except Exception:
            pass
    return srcs


def get_current_gemini_blob_srcs(driver):
    """Lấy URL blob ảnh Gemini hiện có, không phụ thuộc ảnh đã render hay chưa."""
    return set(driver.execute_script("""
        return Array.from(document.querySelectorAll('img'))
            .map(img => img.currentSrc || img.src || '')
            .filter(src => src.startsWith(
                'blob:https://gemini.google.com/'
            ));
    """))


def wait_new_gemini_completed_image(driver, old_srcs, timeout=240):
    """
    Chờ Gemini hoàn tất đúng phản hồi ảnh mới.

    Nút "Tải hình ảnh có kích thước đầy đủ xuống" chỉ được dùng làm tín hiệu
    hoàn tất. Code không click nút này; sau đó lấy thẻ img blob trong cùng
    phản hồi để lưu bằng canvas.
    """
    print(
        "-> Đang chờ phản hồi ảnh hoàn tất "
        "(xuất hiện nút tải hình ảnh kích thước đầy đủ)..."
    )
    deadline = time.time() + timeout
    old_srcs = set(old_srcs or [])
    next_cdp_wakeup = 0.0

    while time.time() < deadline:
        # Edge có thể giảm hoạt động của renderer khi bị minimize. Nhắc lại
        # lifecycle/focus định kỳ, không khôi phục cửa sổ và không chiếm chuột.
        now = time.time()
        if now >= next_cdp_wakeup:
            keep_page_lifecycle_active(driver, quiet=True)
            next_cdp_wakeup = now + 2.0

        result = driver.execute_script("""
            const oldSrcs = new Set(arguments[0]);
            const buttons = Array.from(document.querySelectorAll(
                'button[aria-label]'
            )).filter(button => {
                const label = (button.getAttribute('aria-label') || '')
                    .toLocaleLowerCase();
                return (
                    label.includes('tải hình ảnh có kích thước đầy đủ') ||
                    label.includes('download full-size image') ||
                    label.includes('download full size image')
                );
            });

            // Duyệt từ phản hồi mới nhất về trước.
            for (const button of buttons.reverse()) {
                const response =
                    button.closest('message-content') ||
                    button.closest('model-response') ||
                    button.closest('response-element') ||
                    button.closest('generated-image');
                if (!response) continue;

                const images = Array.from(response.querySelectorAll('img'));
                const fresh = images.filter(img => {
                    const src = img.currentSrc || img.src || '';
                    return src.startsWith(
                        'blob:https://gemini.google.com/'
                    ) && !oldSrcs.has(src);
                });
                if (fresh.length) return fresh[fresh.length - 1];
            }
            return null;
        """, list(old_srcs))

        if result is not None:
            src = result.get_attribute("src") or ""
            print(
                "-> Gemini đã hoàn tất phản hồi ảnh; "
                f"đã tìm thấy blob mới: {src[:100]}"
            )
            return result
        time.sleep(0.5)

    raise Exception(
        "Phản hồi mới chưa xuất hiện nút tải hình ảnh kích thước đầy đủ "
        "hoặc không có blob ảnh mới."
    )


def force_gemini_image_load(driver, img_element, timeout=45):
    """
    Bỏ lazy-load và chủ động yêu cầu Edge load/decode ảnh. Cơ chế này đã
    được kiểm chứng khi cửa sổ Edge đang minimize.
    """
    keep_page_lifecycle_active(driver)
    driver.set_script_timeout(timeout + 5)
    result = driver.execute_async_script("""
        const img = arguments[0];
        const timeoutMs = arguments[1] * 1000;
        const done = arguments[arguments.length - 1];
        let finished = false;

        function finish(value) {
            if (finished) return;
            finished = true;
            done({
                ...value,
                complete: Boolean(img.complete),
                naturalWidth: img.naturalWidth || 0,
                naturalHeight: img.naturalHeight || 0,
                src: img.currentSrc || img.src || ''
            });
        }

        // Không chờ trình duyệt tự xử lý loading="lazy" khi minimize.
        img.loading = 'eager';
        img.removeAttribute('loading');
        try { img.fetchPriority = 'high'; } catch (_) {}
        img.style.contentVisibility = 'visible';
        img.style.visibility = 'visible';
        img.style.display = 'block';
        img.scrollIntoView({block:'center', inline:'center'});

        if (img.complete && img.naturalWidth > 0 && img.naturalHeight > 0) {
            finish({ok:true, method:'already_complete'});
            return;
        }

        const timer = setTimeout(() => {
            finish({ok:false, method:'decode_timeout'});
        }, timeoutMs);

        img.addEventListener('load', () => {
            clearTimeout(timer);
            finish({ok:true, method:'load_event'});
        }, {once:true});

        img.addEventListener('error', () => {
            clearTimeout(timer);
            finish({ok:false, method:'error_event'});
        }, {once:true});

        if (typeof img.decode === 'function') {
            img.decode().then(() => {
                clearTimeout(timer);
                finish({ok:true, method:'img_decode'});
            }).catch(() => {
                // Vẫn chờ sự kiện load cho tới timeout.
            });
        }
    """, img_element, timeout)

    width = int(result.get("naturalWidth") or 0) if result else 0
    height = int(result.get("naturalHeight") or 0) if result else 0
    if not result or not result.get("ok") or width <= 0 or height <= 0:
        raise Exception(f"Không ép tải/decode được ảnh Gemini: {result}")

    print(
        f"-> Đã ép tải ảnh Gemini thành công bằng "
        f"{result.get('method')}: {width}x{height}"
    )
    return result


def wait_new_large_image(driver, old_srcs, timeout=240):
    print("-> Đang chờ ảnh mới Gemini xuất hiện...")
    start = time.time()
    while time.time() - start < timeout:
        imgs = driver.find_elements(By.XPATH, "//img")
        for img in reversed(imgs):
            try:
                src = img.get_attribute("src") or ""
                w = driver.execute_script("return arguments[0].naturalWidth || 0;", img)
                h = driver.execute_script("return arguments[0].naturalHeight || 0;", img)
                if src and w >= 250 and h >= 250 and src not in old_srcs:
                    print(f"-> Đã tìm thấy ảnh mới: {w}x{h}")
                    return img
            except Exception:
                pass
        time.sleep(2)
    raise Exception("Không tìm thấy ảnh mới sau khi gửi prompt.")


def debug_images_on_page(driver):
    imgs = driver.find_elements(By.XPATH, "//img")
    print(f"\nTổng số img trên trang: {len(imgs)}")
    for i, img in enumerate(imgs):
        try:
            src = img.get_attribute("src") or ""
            w = driver.execute_script("return arguments[0].naturalWidth || 0;", img)
            h = driver.execute_script("return arguments[0].naturalHeight || 0;", img)
            rect = driver.execute_script("""
                const r = arguments[0].getBoundingClientRect();
                return {
                    x: Math.round(r.x),
                    y: Math.round(r.y),
                    width: Math.round(r.width),
                    height: Math.round(r.height)
                };
            """, img)
            if w >= 200 and h >= 200:
                print(f"[IMG {i}] size={w}x{h} rect={rect} src={src[:100]}")
        except Exception as e:
            print(f"[IMG {i}] lỗi đọc: {e}")


def open_viewer_and_download(driver, img_element, timeout=30):
    print("-> Click ảnh mới để mở viewer...")
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", img_element)
    time.sleep(1)

    driver.execute_script("""
        const img = arguments[0];
        img.scrollIntoView({block:'center'});
        img.click();
    """, img_element)

    download_xpath = (
        "(//button[contains(@aria-label,'Tải xuống') "
        "or contains(@aria-label,'Download') "
        "or .//mat-icon[contains(text(),'download')] "
        "or .//mat-icon[contains(@data-mat-icon-name,'download')]])[last()]"
    )

    for attempt in range(1, 4):
        print(f"-> Tìm nút tải xuống lần {attempt}...")
        try:
            btn = WebDriverWait(driver, timeout).until(
                EC.element_to_be_clickable((By.XPATH, download_xpath))
            )
            driver.execute_script("arguments[0].click();", btn)
            print(f"-> Đã click nút tải xuống lần {attempt}.")
            return True
        except Exception as e:
            print(f"⚠️ Không tìm thấy/click được nút tải lần {attempt}: {e}")
            try:
                driver.execute_script("""
                    const img = arguments[0];
                    img.scrollIntoView({block:'center'});
                    img.click();
                """, img_element)
                time.sleep(2)
            except Exception:
                pass

    raise Exception("Không tìm được nút tải xuống sau 3 lần thử.")


def click_more_button_near_image(driver, img_element):
    print("-> Đang tìm nút ... gần ảnh mới...")
    button = driver.execute_script("""
        const img = arguments[0];
        const ir = img.getBoundingClientRect();
        const cx = ir.left + ir.width / 2;
        const cy = ir.top + ir.height / 2;
        function visible(el){
            const r = el.getBoundingClientRect();
            const st = getComputedStyle(el);
            return r.width > 0 && r.height > 0 && st.visibility !== 'hidden' && st.display !== 'none';
        }
        const candidates = Array.from(document.querySelectorAll('button, gem-icon-button'))
            .filter(visible)
            .map(el => {
                const txt = (el.innerText || el.getAttribute('aria-label') || el.getAttribute('arialabel') || '').toLowerCase();
                const html = (el.outerHTML || '').toLowerCase();
                const r = el.getBoundingClientRect();
                const bx = r.left + r.width / 2;
                const by = r.top + r.height / 2;
                const dist = Math.abs(bx - cx) + Math.abs(by - cy);
                const isMore = txt.includes('thêm') || txt.includes('more') || txt.includes('tuỳ chọn') || txt.includes('tùy chọn') || html.includes('more_vert');
                return {el, txt, x:r.left, y:r.top, dist, isMore};
            })
            .filter(o => o.isMore)
            .sort((a,b) => a.dist - b.dist);
        if (!candidates.length) return null;
        return candidates[0].el;
    """, img_element)
    if button is None:
        raise Exception("Không tìm thấy nút dấu ba chấm gần ảnh mới.")
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", button)
    time.sleep(0.3)
    driver.execute_script("arguments[0].click();", button)
    print("-> Đã mở menu dấu ba chấm gần ảnh mới.")
    time.sleep(0.8)


def click_download_image_menu_item(driver, timeout=10):
    menu_xpath = (
        "//*[@role='menuitem' and (contains(normalize-space(.),'Tải hình ảnh xuống') "
        "or contains(translate(normalize-space(.),'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'download image'))]"
        "|//button[contains(normalize-space(.),'Tải hình ảnh xuống') "
        "or contains(translate(normalize-space(.),'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'download image')]"
    )
    item = WebDriverWait(driver, timeout).until(
        EC.element_to_be_clickable((By.XPATH, menu_xpath))
    )
    driver.execute_script("arguments[0].click();", item)
    print("-> Đã click 'Tải hình ảnh xuống' trong menu Gemini.")


def generate_and_download_image(driver, wait, prompt, final_save_path, row=None, step="IMAGE"):
    # Luồng chính:
    # 1) lưu danh sách blob ảnh đang có
    # 2) gửi prompt
    # 3) chờ nút tải ảnh kích thước đầy đủ trong đúng phản hồi mới
    # 4) nếu lần đầu lỗi, ghi mã lỗi ngắn rồi mở chat Gemini mới và thử lại
    # 5) nếu lần bảo hiểm cũng lỗi, ghi mã lỗi ngắn và dừng bước ảnh
    # 6) lấy thẻ img blob trong cùng phản hồi (không click nút)
    # 7) bỏ lazy-load, ép Edge load/decode cả khi minimize
    # 8) lưu pixel trực tiếp bằng canvas
    #
    # Bảo hiểm này chỉ chạy khi giai đoạn gửi/chờ ảnh gặp lỗi.
    img = None

    for image_attempt in (1, 2):
        print(
            "\n===== DEBUG TRƯỚC KHI GỬI PROMPT"
            + (" (LẦN BẢO HIỂM)" if image_attempt == 2 else "")
            + " ====="
        )
        debug_images_on_page(driver)

        try:
            old_srcs = get_current_gemini_blob_srcs(driver)
            send_prompt_to_gemini(driver, wait, prompt)
            img = wait_new_gemini_completed_image(
                driver,
                old_srcs,
                timeout=GEMINI_IMAGE_WAIT_SECONDS,
            )
            break
        except Exception:
            if image_attempt == 1:
                if row is not None:
                    write_retry_note(
                        row,
                        1,
                        step,
                        "GEMINI_IMAGE_ATTEMPT_1_FAILED",
                    )
                print(
                    "⚠️ Lần tạo ảnh đầu tiên thất bại. "
                    "Sẽ thử lại đúng 1 lần bằng chat Gemini mới."
                )
                try:
                    open_gemini_and_wait_ready(driver)
                except Exception:
                    if row is not None:
                        write_retry_note(
                            row,
                            2,
                            step,
                            "GEMINI_RETRY_OPEN_FAILED",
                        )
                    raise
                continue

            if row is not None:
                write_retry_note(
                    row,
                    2,
                    step,
                    "GEMINI_IMAGE_ATTEMPT_2_FAILED",
                )
            raise

    print("\n===== PHẢN HỒI ẢNH GEMINI ĐÃ HOÀN TẤT =====")
    debug_images_on_page(driver)

    try:
        force_gemini_image_load(driver, img, timeout=45)
        return save_image_directly_from_element(
            driver,
            img,
            final_save_path,
        )
    except Exception as direct_error:
        if row is not None:
            write_retry_note(
                row,
                2,
                step,
                "GEMINI_CANVAS_FAILED",
                direct_error,
            )
        raise Exception(
            "Gemini đã báo hoàn tất ảnh nhưng không decode/lưu canvas được: "
            f"{direct_error}"
        ) from direct_error

def images_if_needed(driver, wait, task):
    row = task["row"]
    if not is_word_ok(task["word_path"]):
        raise Exception("Chưa có Word hợp lệ nên không tạo ảnh.")
    if not task["brief1"] or not task["brief2"]:
        raise Exception("Chưa có đủ Brief O/P nên không tạo ảnh.")

    output_folder = get_output_folder(task["word_path"])

    if not file_exists(task["path_img1"]):
        print(f"\n[DÒNG {row}] Bắt đầu tạo ảnh 1 Gemini...")
        write_value(row, COL_STATUS_IMG1, "Đang sinh ảnh 1...", save=False)
        open_gemini_and_wait_ready(driver)
        prompt1 = build_gemini_prompt(task["brief1"])
        write_value(row, COL_STATUS_IMG1, STATUS_SENT_IMG1, save=False)
        save_path1 = get_unique_image_path(output_folder, f"{task['name']} 1")
        final_path1 = generate_and_download_image(driver, wait, prompt1, save_path1, row=row, step="IMG1")
        wb, sh = get_sheet()
        sh.range(f"{COL_PATH_IMG1}{row}").value = final_path1
        sh.range(f"{COL_STATUS_IMG1}{row}").value = STATUS_SAVED_IMG1
        sh.range(f"{COL_GEMINI_URL_IMG1}{row}").value = driver.current_url
        wb.save()
        task = read_task(row)
    else:
        print(f"Dòng {row}: Ảnh 1 đã có. Bỏ qua.")
        write_value(row, COL_STATUS_IMG1, STATUS_SAVED_IMG1, save=False)

    if not file_exists(task["path_img2"]):
        print(f"\n[DÒNG {row}] Bắt đầu tạo ảnh 2 Gemini...")
        write_value(row, COL_STATUS_IMG2, "Đang sinh ảnh 2...", save=False)
        open_gemini_and_wait_ready(driver)
        prompt2 = build_gemini_prompt(task["brief2"])
        write_value(row, COL_STATUS_IMG2, STATUS_SENT_IMG2, save=False)
        save_path2 = get_unique_image_path(output_folder, f"{task['name']} 2")
        final_path2 = generate_and_download_image(driver, wait, prompt2, save_path2, row=row, step="IMG2")
        wb, sh = get_sheet()
        sh.range(f"{COL_PATH_IMG2}{row}").value = final_path2
        sh.range(f"{COL_STATUS_IMG2}{row}").value = STATUS_SAVED_IMG2
        sh.range(f"{COL_GEMINI_URL_IMG2}{row}").value = driver.current_url
        wb.save()
    else:
        print(f"Dòng {row}: Ảnh 2 đã có. Bỏ qua.")
        write_value(row, COL_STATUS_IMG2, STATUS_SAVED_IMG2, save=False)

    if mark_done_if_complete(row):
        print(f"=====> HOÀN THÀNH HOÀN TOÀN DÒNG {row} <=====")
    return read_task(row)

# =====================================================
# MAIN TỔNG
# =====================================================
def should_skip_empty_row(task):
    # Dòng trống hoặc thiếu tối thiểu C thì bỏ qua.
    return not task["name"]




def get_resume_checkpoint(task):
    """
    Xác định bước còn thiếu để chạy tiếp.
    Không tạo lại phần đã hoàn thành.
    """
    if task.get("done") == STATUS_OK and not TEST_OVERWRITE_WORD:
        return "DONE"
    if not is_word_ok(task.get("word_path")):
        return "WORD"
    if not task.get("brief1") or not task.get("brief2"):
        return "BRIEF"
    if not file_exists(task.get("path_img1")):
        return "IMG1"
    if not file_exists(task.get("path_img2")):
        return "IMG2"
    return "DONE"


def process_row(row, driver, wait, progress=None):
    """
    Dùng chung một Edge driver đã được mở từ main().
    Không tự quit driver sau từng dòng nên trình duyệt không còn đóng rồi mở lại.
    """
    task = read_task(row)

    if should_skip_empty_row(task):
        return get_active_driver(driver, wait)

    print("\n" + "=" * 70)
    print(f"KIỂM TRA DÒNG {row}: {task['name']} | CHECKPOINT={get_resume_checkpoint(task)}")

    # 1) Nếu X đã OK thì bỏ qua tuyệt đối
    if task["done"] == STATUS_OK and not TEST_OVERWRITE_WORD:
        print(f"Dòng {row}: Cột X đã OK. Bỏ qua toàn dòng.")
        return get_active_driver(driver, wait)

    # 2) Nếu đã có ảnh 2 hoặc status ảnh 2 đã lưu
    status_img2 = str(cell(row, COL_STATUS_IMG2).value or "").strip()
    if (task["path_img2"] or status_img2 == STATUS_SAVED_IMG2) and not TEST_OVERWRITE_WORD:
        print(f"Dòng {row}: Đã có ảnh 2. Ghi X = OK và bỏ qua toàn dòng.")
        write_value(row, COL_DONE, STATUS_OK)
        return get_active_driver(driver, wait)

    # 3) Nếu đã có ảnh 1 thì chỉ tạo tiếp ảnh 2
    if task["path_img1"] and not TEST_OVERWRITE_WORD:
        print(f"Dòng {row}: Đã có ảnh 1. Bỏ qua Word/Brief/Ảnh 1, tạo tiếp ảnh 2.")
        if not task["brief2"]:
            write_value(row, COL_DONE, "Lỗi: Có ảnh 1 nhưng thiếu Brief 2")
            print(f"❌ Dòng {row}: Có ảnh 1 nhưng thiếu Brief 2.")
            return get_active_driver(driver, wait)
        try:
            if progress:
                progress.update(row, "Đang tạo tiếp ảnh 2 trên Gemini")
            images_if_needed(driver, wait, task)
        except Exception as e:
            write_image_final_error(row)
            write_value(row, COL_DONE, f"Lỗi ảnh: {str(e)[:120]}")
            print(f"❌ Dòng {row}: Lỗi bước ảnh: {e}")
        return get_active_driver(driver, wait)

    # 4) Nếu đã có đủ Brief thì chỉ tạo ảnh
    if task["brief1"] and task["brief2"] and not TEST_OVERWRITE_WORD:
        print(f"Dòng {row}: Đã có Brief. Bỏ qua Word, tạo ảnh.")
        try:
            if progress:
                progress.update(row, "Đang tạo ảnh 1/2 trên Gemini")
            images_if_needed(driver, wait, task)
        except Exception as e:
            write_image_final_error(row)
            write_value(row, COL_DONE, f"Lỗi ảnh: {str(e)[:120]}")
            print(f"❌ Dòng {row}: Lỗi bước ảnh: {e}")
        return get_active_driver(driver, wait)

    print(f"BẮT ĐẦU DÒNG {row}: {task['name']}")

    # 1) WORD
    if TEST_OVERWRITE_WORD or not is_word_ok(task["word_path"]):
        try:
            if progress:
                progress.update(row, "Đang viết và lưu file Word")
            task, driver, wait = write_article_if_needed(driver, wait, task)
        except Exception as e:
            write_retry_note(row, 9, "ARTICLE", "WORD_FINAL_ERROR", e)
            write_article_error(row, e)
            print(f"❌ Dòng {row}: Lỗi bước Word: {e}")
            return get_active_driver(driver, wait)
    elif task["article_status"].upper() != STATUS_OK:
        write_value(row, COL_ARTICLE_STATUS, STATUS_OK)

    task = read_task(row)
    if not is_word_ok(task["word_path"]):
        write_article_error(row, "Word không hợp lệ sau kiểm tra. Dừng dòng này.")
        return get_active_driver(driver, wait)

    # 2) BRIEF
    if not task["brief1"] or not task["brief2"]:
        try:
            if progress:
                progress.update(row, "Đang tạo brief ảnh")
            task, driver, wait = brief_if_needed(driver, wait, task)
        except Exception as e:
            write_retry_note(row, 9, "BRIEF", "BRIEF_FINAL_ERROR", e)
            write_value(row, COL_BRIEF_STATUS, f"{STATUS_ERROR_BRIEF}: {str(e)[:120]}")
            print(f"❌ Dòng {row}: Lỗi bước Brief: {e}")
            return get_active_driver(driver, wait)

    # 3) GEMINI ẢNH — dùng luôn cùng cửa sổ Edge
    task = read_task(row)
    if not file_exists(task["path_img1"]) or not file_exists(task["path_img2"]):
        try:
            if progress:
                progress.update(row, "Đang tạo ảnh 1/2 trên Gemini")
            images_if_needed(driver, wait, task)
        except Exception as e:
            write_image_final_error(row)
            write_value(row, COL_DONE, f"Lỗi ảnh: {str(e)[:120]}")
            print(f"❌ Dòng {row}: Lỗi bước ảnh: {e}")
            return get_active_driver(driver, wait)
    else:
        write_value(row, COL_DONE, STATUS_OK)
        print(f"Dòng {row}: Đã đủ ảnh. Ghi X = OK.")

    return get_active_driver(driver, wait)


class WriteProgressWindow:
    """Bảng theo dõi; Edge được ẩn bằng cách đưa ra ngoài màn hình."""

    def __init__(self, total_rows):
        self.total_rows = total_rows
        self.stop_after_row = threading.Event()
        self.messages = queue.Queue()
        threading.Thread(target=self._run, daemon=True).start()

    def _run(self):
        import tkinter as tk

        root = tk.Tk()
        root.title("Tiến độ viết bài — Edge đang chạy ẩn")
        root.geometry("590x175+30+80")
        root.resizable(False, False)
        root.attributes("-topmost", True)

        row_label = tk.Label(
            root,
            text="Đang chuẩn bị...",
            font=("Segoe UI", 13, "bold"),
        )
        row_label.pack(pady=(18, 7))
        step_label = tk.Label(
            root,
            text="Đang mở Edge chạy ẩn",
            font=("Segoe UI", 10),
        )
        step_label.pack(pady=(0, 5))
        initial_count = (
            f"Đã hoàn thành 0/{self.total_rows} dòng"
            if self.total_rows is not None
            else "Đã hoàn thành 0 dòng"
        )
        count_label = tk.Label(root, text=initial_count)
        count_label.pack(pady=(0, 12))

        button_frame = tk.Frame(root)
        button_frame.pack()

        def request_stop():
            self.stop_after_row.set()
            stop_button.config(
                text="Sẽ dừng sau dòng này",
                state="disabled",
            )
            step_label.config(text="Đã nhận lệnh dừng an toàn")

        tk.Button(
            button_frame,
            text="Ẩn bảng",
            width=12,
            command=root.withdraw,
        ).pack(side="left", padx=5)
        tk.Button(
            button_frame,
            text="Hiện Edge",
            width=12,
            command=lambda: set_current_edge_visible(True),
        ).pack(side="left", padx=5)
        tk.Button(
            button_frame,
            text="Ẩn Edge",
            width=12,
            command=lambda: set_current_edge_visible(False),
        ).pack(side="left", padx=5)
        stop_button = tk.Button(
            button_frame,
            text="Dừng sau dòng này",
            width=20,
            command=request_stop,
        )
        stop_button.pack(side="left", padx=5)

        # Dấu X chỉ ẩn bảng, không dừng code.
        root.protocol("WM_DELETE_WINDOW", root.withdraw)

        def poll():
            try:
                while True:
                    row, step, completed = self.messages.get_nowait()
                    row_label.config(
                        text=(
                            f"Dòng Excel đang chạy: {row}"
                            if row
                            else "Đã hoàn tất"
                        )
                    )
                    step_label.config(text=step)
                    count_text = (
                        f"Đã hoàn thành {completed}/{self.total_rows} dòng"
                        if self.total_rows is not None
                        else f"Đã hoàn thành {completed} dòng"
                    )
                    count_label.config(text=count_text)
            except queue.Empty:
                pass
            root.after(200, poll)

        poll()
        root.mainloop()

    def update(self, row, step, completed=0):
        self.messages.put((row, step, completed))

    def should_stop(self):
        return self.stop_after_row.is_set()

    def finish(self, completed, stopped=False):
        if stopped:
            message = "Đã dừng an toàn sau khi hoàn thành dòng hiện tại"
        else:
            message = "Đã quét xong toàn bộ danh sách"
        self.update(None, message, completed)
        try:
            winsound.MessageBeep(winsound.MB_ICONASTERISK)
        except Exception:
            pass

def pause_before_next_article(progress, completed_count):
    """Nghỉ có đếm ngược; nút Dừng vẫn phản hồi ngay trong thời gian chờ."""
    pause_seconds = max(0, int(SHORT_ARTICLE_PAUSE_MINUTES * 60))
    pause_end = time.time() + pause_seconds

    print(
        f"\n[TẠM NGHỈ] Có {SHORT_ARTICLE_STREAK_LIMIT} bài Word mới liên tiếp "
        f"dưới {SHORT_ARTICLE_WORD_LIMIT} từ. Nghỉ "
        f"{SHORT_ARTICLE_PAUSE_MINUTES} phút trước bài tiếp theo."
    )

    while True:
        remaining = max(0, int(pause_end - time.time()))
        if remaining <= 0:
            print("[TIẾP TỤC] Đã hết thời gian nghỉ. Bắt đầu bài tiếp theo.")
            return True

        minutes, seconds = divmod(remaining, 60)
        progress.update(
            None,
            f"Tạm nghỉ do 3 bài ngắn: còn {minutes:02d}:{seconds:02d}",
            completed_count,
        )

        if progress.stop_after_row.wait(timeout=min(30, remaining)):
            print("[ĐÃ DỪNG] Nhận lệnh dừng trong thời gian tạm nghỉ.")
            return False


def main():
    print("FILE TỔNG đang chạy. Hãy mở sẵn file Excel dữ liệu trước.")
    start_row, manual_mark_row = get_start_row_by_manual_mark()
    if manual_mark_row:
        print(f"Đã thấy mốc thủ công {MANUAL_MARK_TEXT} ở cột {COL_MANUAL_MARK}, dòng {manual_mark_row}.")
        print(f"Sẽ bắt đầu xử lý từ dòng {start_row}.")
    else:
        print(f"Không thấy mốc thủ công {MANUAL_MARK_TEXT} ở cột {COL_MANUAL_MARK}.")
        print(f"Sẽ bắt đầu xử lý từ dòng {START_ROW}.")

    if start_row > END_ROW:
        print("Mốc thủ công đã nằm ở dòng giới hạn cuối.")
        return

    # Chỉ kiểm tra dòng đầu tiên; không quét trước toàn bộ danh sách.
    if should_skip_empty_row(read_task(start_row)):
        print(f"Dòng {start_row} không có Từ khóa. Không có bài nào để chạy.")
        return

    # Không biết trước tổng số bài vì sẽ dừng ngay tại dòng trống đầu tiên.
    progress = WriteProgressWindow(None)
    completed_count = 0
    successful_rows_since_recycle = 0
    stopped_early = False
    stopped_at_blank = False
    short_article_streak = 0

    driver = None
    try:
        # Chỉ mở Edge đúng một lần cho toàn bộ danh sách.
        driver, wait = create_shared_driver()

        row = start_row
        while row <= END_ROW:
            task = read_task(row)
            if should_skip_empty_row(task):
                stopped_at_blank = True
                print(f"\n[XONG] Dòng {row} không có Từ khóa. Dừng danh sách.")
                break

            # Giữ nguyên nguyên tắc cũ: dòng đã hoàn tất thì bỏ qua.
            if task["done"] == STATUS_OK and not TEST_OVERWRITE_WORD:
                print(f"Dòng {row}: Trạng thái hoàn tất đã OK. Bỏ qua.")
                row += 1
                continue

            # Nếu bấm dừng ngay lúc chương trình còn chuẩn bị Edge thì không bắt đầu dòng đầu tiên.
            if progress.should_stop():
                stopped_early = True
                print("\n[ĐÃ DỪNG] Chưa bắt đầu dòng Excel tiếp theo.")
                break

            progress.update(row, "Đang kiểm tra dữ liệu của dòng", completed_count)
            word_existed_before = is_word_ok(task["word_path"])
            try:
                driver, wait = process_row(row, driver, wait, progress=progress)
                time.sleep(2)
            except Exception as e:
                print(f"❌ Lỗi ngoài dự kiến tại dòng {row}: {e}")
                try:
                    write_retry_note(row, 9, "TOTAL", "TOTAL_ERROR", e)
                    write_value(row, COL_DONE, f"Lỗi tổng: {str(e)[:120]}")
                except Exception:
                    pass
                time.sleep(2)
            finally:
                # Nếu Lớp 2 đã thay Edge rồi một bước sau đó lỗi, vẫn giữ phiên mới
                # cho dòng tiếp theo thay vì quay lại session cũ đã quit().
                driver, wait = get_active_driver(driver, wait)

            completed_count += 1
            progress.update(row, "Đã xử lý xong và lưu dòng", completed_count)

            # Chỉ tính bài Word vừa được tạo thành công trong phiên chạy này.
            # Các dòng bỏ qua hoặc chỉ xử lý brief/ảnh không ảnh hưởng bộ đếm.
            task_after = read_task(row)
            word_created_now = (
                (TEST_OVERWRITE_WORD or not word_existed_before)
                and is_word_ok(task_after["word_path"])
            )
            if word_created_now:
                saved_word_count = cell(row, COL_WORD_COUNT).value
                try:
                    saved_word_count = int(float(saved_word_count))
                except (TypeError, ValueError):
                    saved_word_count = None

                if saved_word_count is not None:
                    if saved_word_count < SHORT_ARTICLE_WORD_LIMIT:
                        short_article_streak += 1
                        print(
                            f"[BÀI NGẮN] Dòng {row}: {saved_word_count} từ. "
                            f"Chuỗi hiện tại: {short_article_streak}/"
                            f"{SHORT_ARTICLE_STREAK_LIMIT}."
                        )
                    else:
                        if short_article_streak:
                            print(
                                f"[ĐẠT] Dòng {row}: {saved_word_count} từ. "
                                "Đặt lại bộ đếm bài ngắn."
                            )
                        short_article_streak = 0

            # Chỉ dừng tại đây: toàn bộ xử lý và lưu Excel của dòng hiện tại đã kết thúc.
            if progress.should_stop():
                stopped_early = True
                print(f"\n[ĐÃ DỪNG] Hoàn thành dòng {row}, không chạy dòng kế tiếp.")
                break

            if short_article_streak >= SHORT_ARTICLE_STREAK_LIMIT:
                if not pause_before_next_article(progress, completed_count):
                    stopped_early = True
                    break
                short_article_streak = 0

            # Bộ đếm thuộc riêng tiến trình/worker hiện tại và chỉ nằm trong RAM.
            # Chỉ cộng khi chính dòng worker vừa xử lý chuyển từ chưa OK sang OK.
            if (
                task["done"] != STATUS_OK
                and task_after["done"] == STATUS_OK
            ):
                successful_rows_since_recycle += 1
                if successful_rows_since_recycle >= RECYCLE_EVERY_N_ROWS:
                    print(
                        f"\n[WORKER RECYCLING] Đã hoàn thành "
                        f"{successful_rows_since_recycle} bài. Khởi động lại Edge để xả RAM."
                    )
                    if driver:
                        try:
                            driver.quit()
                        except Exception:
                            pass
                        driver = None
                    time.sleep(3)
                    driver, wait = create_shared_driver()
                    successful_rows_since_recycle = 0

            row += 1

        progress.finish(completed_count, stopped=stopped_early)
        if stopped_early:
            print(f"\n[DỪNG AN TOÀN] Đã xử lý {completed_count} dòng.")
        elif stopped_at_blank:
            print(f"\n[XONG] Đã xử lý {completed_count} dòng rồi gặp dòng trống.")
        else:
            print(f"\n[XONG] Đã xử lý đến dòng giới hạn {END_ROW}.")

    finally:
        # Chỉ đóng Edge một lần sau khi toàn bộ danh sách đã chạy xong.
        if driver:
            try:
                driver.quit()
            except Exception:
                pass


if __name__ == "__main__":
    main()
