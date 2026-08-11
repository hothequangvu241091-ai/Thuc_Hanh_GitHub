# -*- coding: utf-8 -*-
"""
FILE TỔNG: VIẾT BÀI CHATGPT -> LƯU WORD -> XIN BRIEF ẢNH -> TẠO ẢNH GEMINI

Dùng lại file Excel cũ, không cần đổi cấu trúc cột.
Có thể chạy song song với 3 file lẻ cũ. Code tổng sẽ kiểm tra thiếu gì thì làm phần đó.

CỘT EXCEL ĐANG DÙNG:
B  = Web / thư mục website
C  = Tên bài / keyword / tên file
D  = Prompt viết bài
F  = Link GPT gốc
G  = Path Word
H  = URL chat bài viết
I  = Status bài viết
J  = Lỗi bài viết
K  = Số từ Word (chỉ lưu số nguyên)
O  = BRIEF_1
P  = BRIEF_2
Q  = Status brief
T  = Status ảnh 1
U  = Path ảnh 1
V  = Status ảnh 2
W  = Path ảnh 2
X  = Done tổng
Y  = URL Gemini ảnh 1
Z  = URL Gemini ảnh 2
AA = Mốc thủ công: nhập đúng "OK OK" để lần sau chạy từ dòng kế tiếp

AB = Retry Count / số lần bảo hiểm
AC = Retry Step / bước đang retry
AD = Retry Error / mã lỗi ngắn
AE = Retry Time / thời điểm ghi lỗi
"""

# =====================================================
# V2.15 CODEX: CHỌN PROFILE WORKER KHI KHỞI ĐỘNG + ĐA LUỒNG
# - Giữ nguyên luồng viết bài / Word / Brief / Gemini.
# - Bổ sung checkpoint và phân loại lỗi để dễ theo dõi.
# - Truyền Edge driver mới về process_row và main sau khi Lớp 2 reset Edge.
# - Giữ tham chiếu driver đang hoạt động kể cả khi lỗi phát sinh giữa chừng.
# =====================================================
VERSION = "03_viet_bai_tao_anh (engine V2.22)"

# =====================================================
# NHẬT KÝ GIẢM DELAY V2.10 (để có lỗi thì trả về đúng mức V2.9)
# =====================================================
# 01. Mở URL H cũ:                    6.0s -> 2.0s
# 02. Mở URL F tạo bài mới:           3.0s -> 1.0s
# 03. Sau gửi prompt bài viết:        3.0s -> 1.0s
# 04. Poll ChatGPT còn generate:       2.0s -> 0.3s
# 05. Đệm sau khi ChatGPT báo xong:    2.0s -> 0.3s
# 06. Nội dung bài đứng yên:           3.0s -> 1.0s
# 07. Snapshot giao Worker Word:       3.0s -> 0.3s
# 08. Mở lại URL trước khi xin Brief:  5.0s -> 2.0s
# 09. Sau gửi prompt Brief/prompt 2:   2.0s -> 0.5s
# 10. Nội dung Brief đứng yên:         3.0s -> 1.0s
# 11. Menu tải Gemini trước/sau click: 0.3s + 0.8s -> 0.1s + 0.3s
# 12. Reload bài viết bảo hiểm:       20.0s -> 8.0s
# 13. Reload tìm lại nút Send:        10.0s -> 5.0s
# 14. Reload Brief bảo hiểm:          10.0s -> 5.0s
# 15. Nghỉ trước retry Word lần 2:     2.0s -> 1.0s
# 16. Timeout tìm ô ChatGPT:          25.0s -> 10.0s
# 17. Timeout tìm nút Send:        20-35.0s -> 10.0s
# 18. Sau gửi prompt Gemini:          3.0s -> 1.0s
# 19. Bảo hiểm Lớp 2 mở Edge mới:     3.0s -> 1.0s
# 20. Poll ảnh Gemini mới:             2.0s -> 0.5s
# 21. Đệm mở viewer ảnh khi tải lỗi:   2.0s -> 0.5s
# Các lớp kiểm tra/bảo hiểm vẫn còn nguyên; V2.10 chỉ giảm thời gian.

ERROR_CODES = {
    "GPT": "GPT_ERROR",
    "WORD": "WORD_ERROR",
    "BRIEF": "BRIEF_ERROR",
    "IMAGE": "IMAGE_ERROR",
}


class ArticleTooShortError(Exception):
    """Bài đã kéo dài nhưng cuối cùng vẫn dưới MIN_WORDS."""


class DriverTransportError(Exception):
    """EdgeDriver/Edge của một Worker đã mất phản hồi hoặc mất session."""


class ChatGPTSendUnconfirmedError(Exception):
    """Đã bấm gửi nhưng không xác nhận được; dừng Worker để tránh gửi trùng."""


class ExcelWriterUnavailableError(RuntimeError):
    """Excel did not accept a command before the safety deadline."""


class WorkbookIdentityChangedError(ExcelWriterUnavailableError):
    """The writer is no longer attached to the selected workbook path."""


class RowIdentityChangedError(Exception):
    """Dòng Excel đã đổi sau khi nạp RAM; không được ghi vào dòng đó."""


import os
import json
import ctypes
import queue
import re
import threading
import time
import winsound
import shutil
import tempfile
import base64
import psutil
from dataclasses import dataclass
from ctypes import wintypes
from html.parser import HTMLParser
import xlwings as xw
from PIL import Image
from docx import Document

from selenium import webdriver
from selenium.webdriver.edge.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import StaleElementReferenceException

# Các HWND Edge thuộc lần chạy hiện tại, dùng cho nút Hiện/Ẩn trên bảng tiến độ.
_HIDDEN_EDGE_HANDLES = set()

# Khóa cấp phát tên ảnh: tránh hai luồng cùng chọn một đường dẫn khi tên bài trùng.
_IMAGE_PATH_LOCK = threading.RLock()
_RESERVED_IMAGE_BASES = set()
_WORKER_EDGE_HANDLES = {}
_EDGE_HANDLE_LOCK = threading.RLock()
_WORKER_CURRENT_NAMES = {}

# =====================================================
# CẤU HÌNH EXCEL / FILE
# =====================================================
WORD_MACRO = "FullProcess_AllSteps"
SHEET_NAME = "VIET_BAI"
START_ROW = 2
END_ROW = 10000
PROJECT_ROOT = os.environ.get(
    "HOTKEYVIP_RUNTIME_ROOT",
    r"D:\CodexProjects\Hotkeyvip",
)
ROOT_OUTPUT = os.path.join(
    PROJECT_ROOT,
    "07_ket_qua",
    "bai_viet",
)
WORKER_DATA_ROOT = r"D:\CodexProjects\Hotkeyvip\02_viet_bai\du_lieu_3_workers"
WORKER_PROFILE_ROOT = os.path.join(WORKER_DATA_ROOT, "profiles")
WORKER_DOWNLOAD_ROOT = os.path.join(WORKER_DATA_ROOT, "downloads")
TEMP_DOWNLOAD_DIR = WORKER_DOWNLOAD_ROOT

COL_WEB = "B"
COL_NAME = "C"
COL_PROMPT = "D"
COL_GPT_URL = "F"
COL_WORD_PATH = "G"
COL_CHAT_URL = "H"
COL_ARTICLE_STATUS = "I"
COL_ARTICLE_ERROR = "J"
COL_WORD_COUNT = "K"

COL_BRIEF_1 = "O"
COL_BRIEF_2 = "P"
COL_BRIEF_STATUS = "Q"

COL_STATUS_IMG1 = "T"
COL_PATH_IMG1 = "U"
COL_STATUS_IMG2 = "V"
COL_PATH_IMG2 = "W"
COL_DONE = "X"
COL_GEMINI_URL_IMG1 = "Y"
COL_GEMINI_URL_IMG2 = "Z"
COL_MANUAL_MARK = "AA"
MANUAL_MARK_TEXT = "OK OK"

# Cột ghi chú lỗi phụ, không ảnh hưởng logic code cũ.
# Không ghi lỗi vào G/O/P/U/W để code cũ vẫn tự chạy tiếp được.
COL_RETRY_COUNT = "AB"
COL_RETRY_STEP = "AC"
COL_RETRY_ERROR = "AD"
COL_RETRY_TIME = "AE"

# Chỉ ánh xạ Excel theo tên tiêu đề. Phần xử lý bài viết giữ nguyên code cũ.
COLUMN_HEADER_BINDINGS = {
    "COL_WEB": "Tên Miền",
    "COL_NAME": "Từ khóa",
    "COL_PROMPT": "Prompt viết bài",
    "COL_GPT_URL": "URL GPT gốc",
    "COL_WORD_PATH": "Đường dẫn Word",
    "COL_CHAT_URL": "URL ChatGPT",
    "COL_ARTICLE_STATUS": "Trạng thái viết",
    "COL_ARTICLE_ERROR": "Lỗi viết",
    "COL_WORD_COUNT": "Số từ Word",
    "COL_BRIEF_1": "Brief ảnh 1",
    "COL_BRIEF_2": "Brief ảnh 2",
    "COL_BRIEF_STATUS": "Trạng thái brief",
    "COL_STATUS_IMG1": "Trạng thái ảnh 1",
    "COL_PATH_IMG1": "Đường dẫn ảnh 1",
    "COL_STATUS_IMG2": "Trạng thái ảnh 2",
    "COL_PATH_IMG2": "Đường dẫn ảnh 2",
    "COL_DONE": "Trạng thái hoàn tất",
    "COL_GEMINI_URL_IMG1": "URL Gemini ảnh 1",
    "COL_GEMINI_URL_IMG2": "URL Gemini ảnh 2",
    "COL_MANUAL_MARK": "Mốc bắt đầu",
    "COL_RETRY_COUNT": "Số lần thử lại",
    "COL_RETRY_STEP": "Bước thử lại",
    "COL_RETRY_ERROR": "Lỗi thử lại",
    "COL_RETRY_TIME": "Thời gian thử lại",
}

_RESOLVED_WORKBOOK = None
_LAST_STABLE_ARTICLE = None
_ACTIVE_DRIVER = None
_ACTIVE_WAIT = None


def remember_active_driver(driver, wait):
    """Lưu phiên Edge mới nhất để không mất session khi Lớp 2 reset Edge."""
    global _ACTIVE_DRIVER, _ACTIVE_WAIT
    if getattr(_THREAD_CONTEXT, "is_worker", False):
        _THREAD_CONTEXT.active_driver = driver
        _THREAD_CONTEXT.active_wait = wait
        return driver, wait
    _ACTIVE_DRIVER = driver
    _ACTIVE_WAIT = wait
    return driver, wait


def get_active_driver(driver=None, wait=None):
    """Lấy phiên Edge còn hoạt động mới nhất, có fallback về đối số hiện tại."""
    if getattr(_THREAD_CONTEXT, "is_worker", False):
        return (
            getattr(_THREAD_CONTEXT, "active_driver", None) or driver,
            getattr(_THREAD_CONTEXT, "active_wait", None) or wait,
        )
    return _ACTIVE_DRIVER or driver, _ACTIVE_WAIT or wait

# Thời gian bảo hiểm theo yêu cầu
ARTICLE_WAIT_SECONDS = 600
ARTICLE_RELOAD_WAIT_SECONDS = 8
BRIEF_WAIT_SECONDS = 40
BRIEF_RELOAD_WAIT_SECONDS = 5
GEMINI_IMAGE_WAIT_SECONDS = 120
GEMINI_RELOAD_WAIT_SECONDS = 20

# V2.16: nhịp chờ riêng cho ChatGPT. Đây là chờ trạng thái thật của giao diện,
# không phải ép nút hoặc tự gửi lại khi chưa xác nhận được lần gửi đầu.
CHATGPT_DOM_POLL_SECONDS = 1.75
CHATGPT_AFTER_FILL_SECONDS = 3.0
CHATGPT_SEND_CONFIRM_SECONDS = 12
CHATGPT_CONTENT_STABLE_SECONDS = 4.0

STATUS_OK = "OK"
STATUS_RUNNING = "RUNNING"
STATUS_ERROR = "ERROR"
STATUS_WORD_QUEUED = "WORD_QUEUED"
STATUS_WORD_ERROR = "WORD_ERROR"
STATUS_OK_BRIEF = "OK BRIEF ẢNH"
STATUS_RUNNING_BRIEF = "ĐANG TẠO BRIEF ẢNH"
STATUS_ERROR_BRIEF = "LỖI BRIEF ẢNH"
STATUS_SENT_IMG1 = "Đã gửi Gemini ảnh 1"
STATUS_SAVED_IMG1 = "Đã lưu ảnh 1 Gemini"
STATUS_SENT_IMG2 = "Đã gửi Gemini ảnh 2"
STATUS_SAVED_IMG2 = "Đã lưu ảnh 2 Gemini"

MIN_WORDS = 700

# Nếu đủ số bài Word mới liên tiếp có số từ dưới ngưỡng này,
# chương trình sẽ nghỉ trước khi bắt đầu bài tiếp theo.
SHORT_ARTICLE_WORD_LIMIT = 1300
SHORT_ARTICLE_PAUSE_MINUTES = 15
SHORT_ARTICLE_STREAK_LIMIT = 2

# Worker Recycling: Tắt/mở lại Edge sau bao nhiêu bài hoàn thành để xả RAM
RECYCLE_EVERY_N_ROWS = 30

# V2: sửa đúng một số này để tăng/giảm số Edge chạy độc lập.
NUM_WORKERS = 2
# Danh sách ID profile được chọn khi khởi động, ví dụ [2, 3].
SELECTED_WORKER_IDS = [1, 2]
WORKER_DEBUG_PORT_BASE = 9222
WRITER_SAVE_INTERVAL_SECONDS = 1.0
EXCEL_RETRY_WARNING_AFTER = 3
EXCEL_RETRY_PAUSE_AFTER = 8
EXCEL_RETRY_MAX_DELAY_SECONDS = 5.0
EXCEL_RETRY_TIMEOUT_SECONDS = 30.0

# Hạ tầng RAM dùng chung; Worker tuyệt đối không giữ Excel COM object thật.
TASK_QUEUE = queue.Queue()
RESULT_QUEUE = queue.Queue()
WORD_QUEUE = queue.Queue()
UI_QUEUE = queue.Queue()
STOP_EVENT = threading.Event()
# Dừng mềm: Worker không lấy thêm dòng mới, nhưng Word/Excel phải xử lý hết
# những việc đã được giao rồi lưu workbook trước khi kết thúc.
SOFT_STOP_EVENT = threading.Event()
RUN_EVENT = threading.Event()
RUN_EVENT.set()
SKIP_PAUSE_EVENTS = {}
_THREAD_CONTEXT = threading.local()
_MEMORY_LOCK = threading.RLock()
_WORD_PENDING_LOCK = threading.RLock()
_WORD_PENDING_ROWS = set()
_WORD_CURRENT_ROW = None
_WORD_CURRENT_STATUS = "Đang chờ bài"
_EXCEL_WRITER_STATUS = "Đang chờ lệnh"
_EXCEL_WRITER_RETRY_COUNT = 0
_EXCEL_WRITER_LAST_ERROR = ""
_EXCEL_WRITER_ALIVE = threading.Event()
_EXCEL_WRITER_READY = threading.Event()
_EXCEL_WRITER_FAILED = threading.Event()

MAX_IMAGE_SIZE = 800
GEMINI_URL = "https://gemini.google.com/app"

# Bản thử một dòng: luôn tạo lại đúng tên Word cũ, không sinh thêm file đánh số.
TEST_OVERWRITE_WORD = False

# Lần đầu để False nhằm quan sát selector menu tải ảnh. Sau khi test ổn có thể
# đổi thành True để toàn bộ Edge Selenium chạy ẩn.
RUN_HEADLESS = False

SECOND_PROMPT = """Các file ĐẦU VÀO đã được cung cấp đầy đủ trong PROMT của tôi
=> Hãy kiểm tra lại thật kỹ để thực hiện đúng"""

ASK_BRIEF_PROMPT = """
Nhiệm vụ:

Đọc toàn bộ bài viết trong cuộc trò chuyện hiện tại.

Không tạo ảnh.

Hãy tạo 2 mô tả cảnh dùng để tạo ảnh minh họa cho bài viết.

Cách lấy nội dung:

Mô tả cảnh 1:
- Dựa trên toàn bộ phần từ H1 đến ngay trước H2 đầu tiên.
- Đây là ảnh mở đầu của bài viết.

Mô tả cảnh 2:
- Dựa trên toàn bộ phần H2 đầu tiên đến ngay trước H2 tiếp theo.
- Đây là ảnh minh họa cho kiến thức trong H2 đầu tiên.

Yêu cầu:

Không mô tả lại tiêu đề.

Hãy đọc toàn bộ nội dung của từng section rồi xác định ý chính mà người đọc cần hình dung.

Từ ý chính đó, xây dựng một concept ảnh phù hợp.

Nếu hai section có nội dung gần giống nhau thì phải tạo hai concept khác nhau để thể hiện hai góc nhìn khác nhau của cùng chủ đề.

Không được chỉ thay đổi:

- Góc chụp
- Khoảng cách camera
- Bố cục

mà phải thay đổi cách truyền tải nội dung.

Mỗi mô tả cảnh cần có:

- Chủ thể hoặc đối tượng trung tâm
- Bối cảnh phù hợp
- Hành động hoặc trạng thái
- Không khí hoặc cảm xúc
- Các chi tiết nổi bật giúp người xem hiểu đúng nội dung section

Viết tự nhiên bằng tiếng Việt.

Không viết theo dạng checklist.

Không thêm giải thích.

Chỉ trả về:

===BRIEF_1===
...

===BRIEF_2===
...
""".strip()

# =====================================================
# EXCEL UTILS
# =====================================================
def get_real_sheet():
    global _RESOLVED_WORKBOOK
    selected_path = os.environ.get("HOTKEYVIP_SELECTED_EXCEL", "").strip()
    wb = None
    app = None
    if selected_path:
        expected = os.path.normcase(os.path.abspath(selected_path))
        try:
            for candidate_app in xw.apps:
                for candidate_book in candidate_app.books:
                    actual = os.path.normcase(os.path.abspath(candidate_book.fullname))
                    if actual == expected:
                        app = candidate_app
                        wb = candidate_book
                        break
                if wb is not None:
                    break
        except Exception:
            wb = None
    if wb is None:
        app = xw.apps.active
        if app:
            candidate_book = app.books.active
            if candidate_book and (
                not selected_path
                or os.path.normcase(os.path.abspath(candidate_book.fullname))
                == os.path.normcase(os.path.abspath(selected_path))
            ):
                wb = candidate_book
    if wb is None:
        raise Exception(
            f"Không tìm thấy đúng workbook Excel theo đường dẫn: {selected_path or '(chưa có đường dẫn)'}"
        )
    sh = wb.sheets[SHEET_NAME]
    workbook_key = (wb.fullname, sh.name)
    if _RESOLVED_WORKBOOK != workbook_key:
        resolve_columns_by_header(sh)
        _RESOLVED_WORKBOOK = workbook_key
    return wb, sh


class _MemoryApi:
    WrapText = False


class _MemoryCell:
    """Ô Excel giả trong RAM. Ghi ô = gửi lệnh sang Excel writer."""
    def __init__(self, address):
        match = re.fullmatch(r"([A-Z]+)(\d+)", address.upper())
        if not match:
            raise ValueError(f"Địa chỉ ô RAM không hợp lệ: {address}")
        self.col, row = match.groups()
        self.row = int(row)
        self.api = _MemoryApi()

    @property
    def value(self):
        with _MEMORY_LOCK:
            return _THREAD_CONTEXT.rows[self.row].get(self.col)

    @value.setter
    def value(self, value):
        with _MEMORY_LOCK:
            _THREAD_CONTEXT.rows[self.row][self.col] = value
        if _EXCEL_WRITER_FAILED.is_set():
            return
        RESULT_QUEUE.put(("WRITE", self.row, self.col, value))


class _MemorySheet:
    def range(self, address):
        return _MemoryCell(address)


class _MemoryWorkbook:
    def save(self):
        if _EXCEL_WRITER_FAILED.is_set():
            return
        RESULT_QUEUE.put(("SAVE",))


def get_sheet():
    """Worker nhận proxy RAM; chỉ writer/main được phép nhận Excel thật."""
    if getattr(_THREAD_CONTEXT, "is_worker", False):
        return _MemoryWorkbook(), _MemorySheet()
    return get_real_sheet()


def column_letter(number):
    result = ""
    while number:
        number, remainder = divmod(number - 1, 26)
        result = chr(65 + remainder) + result
    return result


def resolve_columns_by_header(sh):
    headers = {}
    last_col = sh.used_range.last_cell.column
    for col in range(1, last_col + 1):
        value = sh.range((1, col)).value
        if value is not None:
            headers[str(value).strip().casefold()] = column_letter(col)

    missing = []
    for global_name, header_name in COLUMN_HEADER_BINDINGS.items():
        letter = headers.get(header_name.casefold())
        if letter:
            globals()[global_name] = letter
        else:
            missing.append(header_name)

    if missing:
        raise Exception(
            "Sheet VIET_BAI thiếu tiêu đề: " + ", ".join(missing)
        )


def is_blank(value):
    return value is None or str(value).strip() == ""


def cell(row, col):
    _, sh = get_sheet()
    return sh.range(f"{col}{row}")


def safe_filename(name):
    name = str(name or "").strip()
    name = re.sub(r'[\\/:*?"<>|]', "-", name)
    name = re.sub(r"\s+", " ", name)
    return name[:160]


def make_output_path(web, file_name):
    folder = os.path.join(ROOT_OUTPUT, safe_filename(web))
    os.makedirs(folder, exist_ok=True)

    base_name = safe_filename(file_name)

    # Ưu tiên tên gốc
    path = os.path.join(folder, base_name + ".docx")
    if TEST_OVERWRITE_WORD:
        return path
    if not os.path.exists(path):
        return path

    # Nếu đã tồn tại thì thêm số phía sau
    i = 1
    while True:
        path = os.path.join(folder, f"{base_name} {i}.docx")

        if not os.path.exists(path):
            return path

        i += 1


def file_exists(path):
    return bool(path) and os.path.exists(str(path).strip())


def read_task(row):
    _, sh = get_sheet()
    return {
        "row": row,
        "web": str(sh.range(f"{COL_WEB}{row}").value or "").strip(),
        "name": str(sh.range(f"{COL_NAME}{row}").value or "").strip(),
        "prompt": str(sh.range(f"{COL_PROMPT}{row}").value or ""),
        "gpt_url": str(sh.range(f"{COL_GPT_URL}{row}").value or "").strip(),
        "word_path": str(sh.range(f"{COL_WORD_PATH}{row}").value or "").strip(),
        "chat_url": str(sh.range(f"{COL_CHAT_URL}{row}").value or "").strip(),
        "article_status": str(sh.range(f"{COL_ARTICLE_STATUS}{row}").value or "").strip(),
        "brief1": str(sh.range(f"{COL_BRIEF_1}{row}").value or "").strip(),
        "brief2": str(sh.range(f"{COL_BRIEF_2}{row}").value or "").strip(),
        "path_img1": str(sh.range(f"{COL_PATH_IMG1}{row}").value or "").strip(),
        "path_img2": str(sh.range(f"{COL_PATH_IMG2}{row}").value or "").strip(),
        "done": str(sh.range(f"{COL_DONE}{row}").value or "").strip().upper(),
    }


def get_last_row():
    _, sh = get_sheet()
    last_row = sh.range(f"{COL_NAME}{sh.cells.last_cell.row}").end("up").row
    return min(max(last_row, START_ROW), END_ROW)


def get_start_row_by_manual_mark():
    """
    Cột AA là mốc thủ công do người dùng tự nhập.
    Nếu có ô AA ghi đúng "OK OK", lấy dòng "OK OK" cuối cùng từ dưới lên
    rồi bắt đầu xử lý từ dòng kế tiếp. Code không tự ghi vào cột AA.
    Nếu không có mốc, giữ nguyên cách chạy cũ từ START_ROW.
    """
    _, sh = get_sheet()
    values = sh.range(
        f"{COL_MANUAL_MARK}{START_ROW}:{COL_MANUAL_MARK}{END_ROW}"
    ).options(ndim=1).value
    for index in range(len(values) - 1, -1, -1):
        row = START_ROW + index
        value = str(values[index] or "").strip()
        if value == MANUAL_MARK_TEXT:
            return row + 1, row
    return START_ROW, None


def write_value(row, col, value, save=True):
    wb, sh = get_sheet()
    target = sh.range(f"{col}{row}")
    target.value = value
    if col in {
        COL_ARTICLE_STATUS,
        COL_ARTICLE_ERROR,
        COL_BRIEF_STATUS,
        COL_DONE,
        COL_RETRY_COUNT,
        COL_RETRY_STEP,
        COL_RETRY_ERROR,
        COL_RETRY_TIME,
    }:
        target.api.WrapText = False
    if save:
        wb.save()


def write_retry_note(row, retry_count, step, error_code, detail=""):
    """
    Ghi chú lỗi phụ ở AB/AC/AD/AE.
    Không ghi vào các cột dữ liệu thật G/O/P/U/W để code cũ vẫn resume bình thường.
    """
    try:
        wb, sh = get_sheet()
        sh.range(f"{COL_RETRY_COUNT}{row}").value = retry_count
        sh.range(f"{COL_RETRY_STEP}{row}").value = step
        sh.range(f"{COL_RETRY_ERROR}{row}").value = f"{error_code}: {str(detail)[:180]}" if detail else error_code
        sh.range(f"{COL_RETRY_TIME}{row}").value = time.strftime("%Y-%m-%d %H:%M:%S")
        for col in (
            COL_RETRY_COUNT,
            COL_RETRY_STEP,
            COL_RETRY_ERROR,
            COL_RETRY_TIME,
        ):
            sh.range(f"{col}{row}").api.WrapText = False
        wb.save()
    except Exception as e:
        print(f"⚠️ Không ghi được retry note dòng {row}: {e}")


def write_image_final_error(row):
    """Ghi mã lỗi ảnh cuối cùng dạng ngắn; không chụp hoặc lưu ảnh màn hình lỗi."""
    try:
        _wb, sh = get_sheet()
        saved_error = str(
            sh.range(f"{COL_RETRY_ERROR}{row}").value or ""
        ).strip()
        if saved_error.startswith("GEMINI_"):
            return
    except Exception:
        pass
    write_retry_note(row, 9, "IMAGE", "IMAGE_FINAL_ERROR")


def reload_current_url(driver, wait_seconds, label=""):
    """
    Không bấm phím F5 thật. Load lại URL hiện tại bằng Selenium để ổn định hơn.
    """
    try:
        current_url = driver.current_url
        print(f"-> Load lại URL hiện tại {label}: {current_url}")
        driver.get(current_url)
        time.sleep(wait_seconds)
        return True
    except Exception as e:
        print(f"⚠️ Không load lại được URL hiện tại {label}: {e}")
        return False


class ArticleHTMLTextExtractor(HTMLParser):
    """Tách chữ để kiểm tra từ HTML gốc dùng cho Word, không sửa HTML."""

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.parts = []
        self.ignored_depth = 0

    def handle_starttag(self, tag, attrs):
        if tag in {"script", "style", "svg"}:
            self.ignored_depth += 1
        elif not self.ignored_depth and tag in {
            "p", "div", "br", "li", "tr", "td", "th",
            "h1", "h2", "h3", "h4", "h5", "h6",
        }:
            self.parts.append(" ")

    def handle_endtag(self, tag):
        if tag in {"script", "style", "svg"} and self.ignored_depth:
            self.ignored_depth -= 1
        elif not self.ignored_depth and tag in {
            "p", "div", "li", "tr", "td", "th",
            "h1", "h2", "h3", "h4", "h5", "h6",
        }:
            self.parts.append(" ")

    def handle_data(self, data):
        if not self.ignored_depth and data:
            self.parts.append(data)

    def text(self):
        return re.sub(r"\s+", " ", " ".join(self.parts)).strip()


def get_word_source_text(article_snapshot):
    """
    Lấy text từ chính HTML sẽ được copy_and_save_perfect() đưa vào Word.
    Chỉ tạo bản text trong bộ nhớ để kiểm tra; HTML gốc được giữ nguyên.
    """
    html_content = str((article_snapshot or {}).get("html") or "").strip()
    if not html_content:
        return ""
    parser = ArticleHTMLTextExtractor()
    parser.feed(html_content)
    parser.close()
    return parser.text()


def get_gpt_content_after_wait(
    driver,
    timeout_seconds,
    label="",
    stable_seconds=CHATGPT_CONTENT_STABLE_SECONDS,
):
    """
    Chờ ChatGPT xong trong thời gian giới hạn rồi lấy markdown cuối.
    Nếu timeout/lag thì trả None để lớp bảo hiểm xử lý tiếp.
    """
    try:
        wait_for_gpt_done(driver, max_timeout_seconds=timeout_seconds)
    except Exception as e:
        print(f"⚠️ {label} chờ GPT quá lâu hoặc lỗi: {e}")
    try:
        # Mỗi Worker giữ snapshot riêng; không dùng biến global chung giữa 1-5 Edge.
        _THREAD_CONTEXT.last_stable_article = capture_stable_assistant_article(
            driver,
            stable_seconds=stable_seconds,
            max_timeout_seconds=max(15, min(45, timeout_seconds)),
        )
        content = get_word_source_text(_THREAD_CONTEXT.last_stable_article)
        if content and str(content).strip():
            return content
    except Exception as e:
        print(f"⚠️ {label} không lấy được nội dung GPT: {e}")
    return None


def write_article_running(row):
    wb, sh = get_sheet()
    sh.range(f"{COL_ARTICLE_STATUS}{row}").value = STATUS_RUNNING
    sh.range(f"{COL_ARTICLE_ERROR}{row}").value = ""
    sh.range(f"{COL_ARTICLE_STATUS}{row}").api.WrapText = False
    sh.range(f"{COL_ARTICLE_ERROR}{row}").api.WrapText = False


def write_article_error(row, text):
    wb, sh = get_sheet()
    sh.range(f"{COL_ARTICLE_STATUS}{row}").value = STATUS_ERROR
    sh.range(f"{COL_ARTICLE_ERROR}{row}").value = str(text)[:500]
    sh.range(f"{COL_ARTICLE_STATUS}{row}").api.WrapText = False
    sh.range(f"{COL_ARTICLE_ERROR}{row}").api.WrapText = False
    wb.save()


def write_article_success(row, word_path, chat_url, word_count):
    wb, sh = get_sheet()
    sh.range(f"{COL_WORD_PATH}{row}").value = word_path
    sh.range(f"{COL_CHAT_URL}{row}").value = chat_url
    sh.range(f"{COL_ARTICLE_STATUS}{row}").value = STATUS_OK
    sh.range(f"{COL_ARTICLE_ERROR}{row}").value = ""
    sh.range(f"{COL_WORD_COUNT}{row}").value = int(word_count)
    sh.range(f"{COL_ARTICLE_STATUS}{row}").api.WrapText = False
    sh.range(f"{COL_ARTICLE_ERROR}{row}").api.WrapText = False
    wb.save()


def write_article_queued(row, word_path, chat_url, word_count):
    """Ghi nhận bài đã giao cho Worker Word; chưa coi là Word thành công."""
    wb, sh = get_sheet()
    sh.range(f"{COL_WORD_PATH}{row}").value = word_path
    sh.range(f"{COL_CHAT_URL}{row}").value = chat_url
    sh.range(f"{COL_ARTICLE_STATUS}{row}").value = STATUS_WORD_QUEUED
    sh.range(f"{COL_ARTICLE_ERROR}{row}").value = ""
    sh.range(f"{COL_WORD_COUNT}{row}").value = int(word_count)
    sh.range(f"{COL_ARTICLE_STATUS}{row}").api.WrapText = False
    sh.range(f"{COL_ARTICLE_ERROR}{row}").api.WrapText = False
    wb.save()


def write_word_worker_error(row, text):
    wb, sh = get_sheet()
    sh.range(f"{COL_ARTICLE_STATUS}{row}").value = STATUS_WORD_ERROR
    sh.range(f"{COL_ARTICLE_ERROR}{row}").value = str(text)[:500]
    sh.range(f"{COL_ARTICLE_STATUS}{row}").api.WrapText = False
    sh.range(f"{COL_ARTICLE_ERROR}{row}").api.WrapText = False
    wb.save()


def word_job_is_pending(row):
    with _WORD_PENDING_LOCK:
        return row in _WORD_PENDING_ROWS


def word_ready_or_pending(task):
    """Dùng để quyết định có cần tạo/giao lại Word hay không."""
    return is_word_ok(task.get("word_path")) or word_job_is_pending(task["row"])


def article_ready_for_downstream(task):
    """Brief/Gemini được chạy sau khi HTML đã giao Word, không chờ file .docx."""
    status = str(task.get("article_status") or "").strip().upper()
    return (
        is_word_ok(task.get("word_path"))
        or word_job_is_pending(task["row"])
        or status in {STATUS_WORD_QUEUED, STATUS_WORD_ERROR}
    )


def mark_done_if_complete(row):
    task = read_task(row)
    if is_word_ok(task["word_path"]) and task["brief1"] and task["brief2"] and file_exists(task["path_img1"]) and file_exists(task["path_img2"]):
        write_value(row, COL_DONE, STATUS_OK)
        return True
    return False

# =====================================================
# KIỂM TRA WORD
# =====================================================
def count_words(text_content):
    if not text_content:
        return 0
    return len(re.findall(r"\S+", text_content.strip()))


def read_word_text(path):
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell_obj in row.cells:
                if cell_obj.text:
                    parts.append(cell_obj.text)
    return "\n".join(parts)


def is_word_ok(path, min_words=80):
    try:
        if not file_exists(path):
            return False
        if os.path.getsize(path) < 5 * 1024:
            return False
        text = read_word_text(path)
        return count_words(text) >= min_words
    except Exception:
        return False


def check_keyword_exists(text_content, keyword):
    return str(keyword).lower() in str(text_content).lower()

# =====================================================
# DRIVER
# =====================================================
def is_driver_transport_error(error):
    """Phân biệt lỗi EdgeDriver chết/treo với lỗi nội dung hoặc selector bình thường."""
    parts = []
    current = error
    seen = set()
    while current is not None and id(current) not in seen:
        seen.add(id(current))
        parts.append(str(current))
        current = current.__cause__ or current.__context__
    message = " | ".join(parts).casefold()
    signatures = (
        "httpconnectionpool(host='localhost'",
        'httpconnectionpool(host="localhost"',
        "read timed out",
        "max retries exceeded with url",
        "invalid session id",
        "session deleted because of page crash",
        "disconnected: not connected to devtools",
        "chrome not reachable",
        "no such window: target window already closed",
        "connection refused",
        "remote end closed connection",
    )
    return any(signature in message for signature in signatures)


def raise_if_driver_transport_error(error):
    if is_driver_transport_error(error):
        raise DriverTransportError(str(error)) from error


def ensure_temp_download_dir():
    os.makedirs(get_temp_download_dir(), exist_ok=True)


def get_temp_download_dir():
    worker_id = getattr(_THREAD_CONTEXT, "worker_id", None)
    if worker_id is None:
        return WORKER_DOWNLOAD_ROOT
    return os.path.join(WORKER_DOWNLOAD_ROOT, f"worker_{worker_id}")


def add_background_running_options(options):
    """Giữ Chromium tiếp tục xử lý trang khi cửa sổ bị minimize/che khuất."""
    options.add_argument("--disable-background-timer-throttling")
    options.add_argument("--disable-backgrounding-occluded-windows")
    options.add_argument("--disable-renderer-backgrounding")
    options.add_argument("--disable-features=CalculateNativeWinOcclusion")


def keep_page_lifecycle_active(driver, quiet=False):
    """
    Yêu cầu Chromium giữ trang ở lifecycle active và giả lập trang vẫn có
    focus. Không gọi Page.bringToFront nên không tự mở/khôi phục cửa sổ Edge.
    """
    errors = []
    try:
        driver.execute_cdp_cmd(
            "Page.setWebLifecycleState",
            {"state": "active"},
        )
    except Exception as e:
        errors.append(f"lifecycle: {e}")

    try:
        driver.execute_cdp_cmd(
            "Emulation.setFocusEmulationEnabled",
            {"enabled": True},
        )
    except Exception as e:
        errors.append(f"focus: {e}")

    if errors and not quiet:
        print(
            "⚠️ Không giữ được đầy đủ trạng thái active/focus, "
            f"tiếp tục bình thường: {'; '.join(errors)}"
        )
    return not errors


def hide_edge_window_by_driver(driver, timeout=10):
    """Đưa Edge ra ngoài màn hình; không dùng trạng thái Windows SW_HIDE."""
    global _HIDDEN_EDGE_HANDLES
    service_process = getattr(
        getattr(driver, "service", None),
        "process",
        None,
    )
    if service_process is None:
        raise Exception("Không lấy được PID của EdgeDriver để chạy SW_HIDE.")

    root_pid = int(service_process.pid)
    process_ids = {root_pid}
    user32 = ctypes.windll.user32
    hidden_handles = set()
    deadline = time.time() + timeout
    callback_type = ctypes.WINFUNCTYPE(
        wintypes.BOOL,
        wintypes.HWND,
        wintypes.LPARAM,
    )

    while time.time() < deadline and not hidden_handles:
        candidates = []
        try:
            root_process = psutil.Process(root_pid)
            process_ids.update(
                child.pid for child in root_process.children(recursive=True)
            )
        except (psutil.Error, OSError):
            pass

        @callback_type
        def enum_callback(hwnd, _lparam):
            window_pid = wintypes.DWORD()
            user32.GetWindowThreadProcessId(
                hwnd,
                ctypes.byref(window_pid),
            )
            if int(window_pid.value) not in process_ids:
                return True

            class_buffer = ctypes.create_unicode_buffer(256)
            user32.GetClassNameW(hwnd, class_buffer, len(class_buffer))
            if class_buffer.value.startswith("Chrome_WidgetWin"):
                title_length = user32.GetWindowTextLengthW(hwnd)
                title_buffer = ctypes.create_unicode_buffer(
                    max(1, title_length + 1)
                )
                user32.GetWindowTextW(
                    hwnd,
                    title_buffer,
                    len(title_buffer),
                )
                rect = wintypes.RECT()
                user32.GetWindowRect(hwnd, ctypes.byref(rect))
                area = max(0, rect.right - rect.left) * max(
                    0,
                    rect.bottom - rect.top,
                )
                candidates.append(
                    {
                        "hwnd": int(hwnd),
                        "title": title_buffer.value.strip(),
                        "visible": bool(user32.IsWindowVisible(hwnd)),
                        "area": area,
                    }
                )
            return True

        user32.EnumWindows(enum_callback, 0)

        # Cửa sổ trình duyệt chính có tiêu đề trang. Các Chrome_WidgetWin
        # không tiêu đề là cửa sổ phụ nội bộ, không được hiện/ẩn cùng nút.
        main_candidates = [
            item for item in candidates
            if item["title"] and item["visible"] and item["area"] > 0
        ]
        if not main_candidates:
            main_candidates = [
                item for item in candidates
                if item["visible"] and item["area"] > 0
            ]

        if main_candidates:
            main_window = max(
                main_candidates,
                key=lambda item: item["area"],
            )
            virtual_right = (
                user32.GetSystemMetrics(76)
                + user32.GetSystemMetrics(78)
            )
            user32.ShowWindow(main_window["hwnd"], 9)  # SW_RESTORE
            user32.SetWindowPos(
                main_window["hwnd"],
                0,
                virtual_right + 100,
                0,
                0,
                0,
                0x0001 | 0x0004 | 0x0010,
            )
            hidden_handles.add(main_window["hwnd"])
        else:
            time.sleep(0.25)

    if not hidden_handles:
        raise Exception("Không tìm thấy cửa sổ Edge Selenium để đưa ra ngoài màn hình.")

    worker_id = int(getattr(_THREAD_CONTEXT, "worker_id", 0))
    with _EDGE_HANDLE_LOCK:
        _HIDDEN_EDGE_HANDLES.update(hidden_handles)
        if worker_id:
            _WORKER_EDGE_HANDLES[worker_id] = set(hidden_handles)
    print(
        "-> Đã đưa Edge Selenium ra ngoài phạm vi màn hình."
    )


def set_worker_edge_visible(worker_id, visible):
    """Hiện/ẩn đúng cửa sổ Edge của một Worker mà không đổi Selenium session."""
    user32 = ctypes.windll.user32
    with _EDGE_HANDLE_LOCK:
        handles = list(_WORKER_EDGE_HANDLES.get(int(worker_id), set()))
    handles = [hwnd for hwnd in handles if user32.IsWindow(hwnd)]
    if not handles:
        UI_QUEUE.put(("STATUS", int(worker_id), None, "Chưa tìm thấy cửa sổ Edge", 0))
        return False

    if visible:
        screen_width = max(900, user32.GetSystemMetrics(0))
        screen_height = max(700, user32.GetSystemMetrics(1))
        width = max(700, screen_width // 2)
        height = max(600, screen_height - 120)
        x = 30 + ((int(worker_id) - 1) % 2) * max(40, screen_width // 2)
        y = 50
        for hwnd in handles:
            user32.ShowWindow(hwnd, 9)  # SW_RESTORE
            user32.SetWindowPos(hwnd, 0, x, y, width, height, 0x0004)
        try:
            user32.SetForegroundWindow(handles[0])
        except Exception:
            pass
    else:
        virtual_right = user32.GetSystemMetrics(76) + user32.GetSystemMetrics(78)
        for hwnd in handles:
            user32.ShowWindow(hwnd, 9)
            user32.SetWindowPos(
                hwnd, 0, virtual_right + 100, 0, 0, 0,
                0x0001 | 0x0004 | 0x0010,
            )
    return True


def set_all_worker_edges_visible(visible):
    """Hiện/ẩn toàn bộ Edge Worker hiện còn sống."""
    changed = False
    for worker_id in SELECTED_WORKER_IDS:
        changed = set_worker_edge_visible(worker_id, visible) or changed
    return changed


def set_current_edge_visible(visible):
    """Hiện hoặc ẩn lại các cửa sổ Edge Selenium của lần chạy hiện tại."""
    user32 = ctypes.windll.user32
    valid_handles = [
        hwnd for hwnd in _HIDDEN_EDGE_HANDLES
        if user32.IsWindow(hwnd)
    ]
    if not valid_handles:
        print("⚠️ Chưa tìm thấy cửa sổ Edge của lần chạy hiện tại.")
        return False

    for hwnd in valid_handles:
        user32.ShowWindow(hwnd, 9)  # SW_RESTORE
        if visible:
            user32.SetWindowPos(
                hwnd, 0, 60, 80, 0, 0, 0x0001 | 0x0004
            )
        else:
            virtual_right = (
                user32.GetSystemMetrics(76)
                + user32.GetSystemMetrics(78)
            )
            user32.SetWindowPos(
                hwnd,
                0,
                virtual_right + 100,
                0,
                0,
                0,
                0x0001 | 0x0004 | 0x0010,
            )

    if visible:
        try:
            user32.SetForegroundWindow(valid_handles[0])
        except Exception:
            pass
        print("-> Đã hiện lại Edge Selenium.")
    else:
        print("-> Đã đưa Edge Selenium ra ngoài phạm vi màn hình.")
    return True


def close_existing_edge_for_profile(profile_dir, timeout=8):
    """
    Đóng Edge Selenium còn sót đang dùng đúng profile của bản CHAY_HIDE.
    Không tác động Edge cá nhân hoặc worker dùng profile khác.
    """
    normalized_profile = os.path.normcase(
        os.path.abspath(profile_dir)
    )
    browser_roots = []

    for process in psutil.process_iter(["pid", "name", "cmdline"]):
        try:
            if (process.info.get("name") or "").lower() != "msedge.exe":
                continue

            command_parts = process.info.get("cmdline") or []
            command_line = " ".join(command_parts)
            normalized_command = os.path.normcase(command_line)
            uses_profile = (
                f"--user-data-dir={normalized_profile}" in normalized_command
                or f'--user-data-dir="{normalized_profile}"'
                in normalized_command
            )
            is_browser_root = not any(
                str(part).startswith("--type=")
                for part in command_parts
            )

            if uses_profile and is_browser_root:
                browser_roots.append(process)
        except (psutil.Error, OSError):
            continue

    if not browser_roots:
        return 0

    targets = []
    for root_process in browser_roots:
        try:
            targets.extend(root_process.children(recursive=True))
        except psutil.Error:
            pass
        targets.append(root_process)

    unique_targets = {
        process.pid: process for process in targets
    }
    for process in unique_targets.values():
        try:
            process.terminate()
        except psutil.Error:
            pass

    _, alive = psutil.wait_procs(
        list(unique_targets.values()),
        timeout=timeout,
    )
    for process in alive:
        try:
            process.kill()
        except psutil.Error:
            pass

    if alive:
        psutil.wait_procs(alive, timeout=3)

    print(
        "-> Đã đóng Edge Selenium cũ dùng profile CHAY_HIDE: "
        f"{len(unique_targets)} tiến trình."
    )
    return len(unique_targets)


def create_chatgpt_driver():
    options = Options()
    options.add_argument(r"--user-data-dir=D:\autodangky\hotkey\SeleniumData")
    options.add_argument("--remote-debugging-port=9222")
    options.add_argument("--disable-notifications")
    options.add_argument("--start-maximized")
    add_background_running_options(options)
    driver = webdriver.Edge(options=options)
    wait = WebDriverWait(driver, 45)
    return driver, wait


def create_gemini_driver():
    ensure_temp_download_dir()
    options = Options()
    options.add_argument(r"--user-data-dir=D:\autodangky\hotkey\SeleniumData")
    options.add_argument("--disable-notifications")
    options.add_argument("--start-maximized")
    add_background_running_options(options)
    prefs = {
        "download.default_directory": get_temp_download_dir(),
        "download.prompt_for_download": False,
        "download.directory_upgrade": True,
        "safebrowsing.enabled": True,
        "profile.default_content_setting_values.automatic_downloads": 1,
    }
    options.add_experimental_option("prefs", prefs)
    driver = webdriver.Edge(options=options)
    wait = WebDriverWait(driver, 45)
    return driver, wait


def create_shared_driver():
    """
    Mở duy nhất 1 cửa sổ Edge cho toàn bộ quá trình.
    Driver này dùng chung cho cả ChatGPT và Gemini để tránh đóng/mở trình duyệt
    ở mỗi dòng hoặc khi chuyển từ bước viết bài sang bước tạo ảnh.
    """
    ensure_temp_download_dir()
    worker_id = int(getattr(_THREAD_CONTEXT, "worker_id", 1))
    selenium_profile = os.path.join(
        WORKER_PROFILE_ROOT, f"worker_{worker_id}"
    )
    debug_port = WORKER_DEBUG_PORT_BASE + worker_id - 1
    close_existing_edge_for_profile(selenium_profile)

    options = Options()
    options.add_argument(f"--user-data-dir={selenium_profile}")
    options.add_argument(f"--remote-debugging-port={debug_port}")
    options.add_argument("--disable-notifications")
    options.add_argument("--start-maximized")
    add_background_running_options(options)
    if RUN_HEADLESS:
        options.add_argument("--headless=new")
        options.add_argument("--window-size=1920,1080")
    prefs = {
        "download.default_directory": get_temp_download_dir(),
        "download.prompt_for_download": False,
        "download.directory_upgrade": True,
        "safebrowsing.enabled": True,
        "profile.default_content_setting_values.automatic_downloads": 1,
    }
    options.add_experimental_option("prefs", prefs)
    driver = webdriver.Edge(options=options)
    if not RUN_HEADLESS:
        hide_edge_window_by_driver(driver)
    wait = WebDriverWait(driver, 45)
    return remember_active_driver(driver, wait)

# =====================================================
# CHATGPT: GỬI PROMPT / LẤY NỘI DUNG / LƯU WORD
# =====================================================
def wait_chatgpt_page_ready(driver, wait, timeout=45):
    """Chờ document và composer ChatGPT thật sự sẵn sàng trước khi điền prompt."""
    deadline = time.time() + timeout
    last_error = None
    while time.time() < deadline:
        try:
            document_ready = driver.execute_script(
                "return document.readyState === 'complete' || document.readyState === 'interactive';"
            )
            if document_ready:
                box = get_chatgpt_input_box(
                    driver,
                    wait,
                    timeout=min(3, max(1, int(deadline - time.time()))),
                )
                if box and box.is_displayed() and box.is_enabled():
                    return box
        except Exception as exc:
            last_error = exc
        time.sleep(CHATGPT_DOM_POLL_SECONDS)
    raise Exception(
        f"ChatGPT chưa sẵn sàng sau {timeout}s. Lỗi gần nhất: {last_error}"
    )


def get_chatgpt_input_box(driver, wait, timeout=10):
    """
    Lấy đúng ô nhập prompt thật của ChatGPT ở dưới cùng.
    Tránh dán/gửi nhầm vào Article Detail / Edit / Canvas vì các vùng đó cũng có contenteditable/ProseMirror.
    """
    start = time.time()
    last_error = None

    while time.time() - start < timeout:
        try:
            # Chỉ dùng id prompt-textarea. Không dùng XPath contenteditable rộng.
            candidates = driver.find_elements(By.XPATH, "//*[@id='prompt-textarea']")
            valid_boxes = []

            win_h = driver.execute_script("return window.innerHeight || 0;")

            for el in candidates:
                try:
                    if not el.is_displayed():
                        continue

                    rect = driver.execute_script("""
                        const r = arguments[0].getBoundingClientRect();
                        return {
                            top: r.top,
                            bottom: r.bottom,
                            left: r.left,
                            right: r.right,
                            width: r.width,
                            height: r.height
                        };
                    """, el)

                    # Bỏ qua editor ở phần nội dung bài viết nếu nó nằm quá cao.
                    if rect["bottom"] < win_h * 0.50:
                        continue

                    # Bỏ qua phần tử quá nhỏ / không phải ô nhập.
                    if rect["width"] < 200 or rect["height"] < 15:
                        continue

                    valid_boxes.append((rect["top"], el))
                except Exception as e:
                    last_error = e
                    continue

            if valid_boxes:
                # Nếu có nhiều phần tử cùng id, lấy phần tử thấp nhất trên màn hình.
                valid_boxes.sort(key=lambda x: x[0], reverse=True)
                return valid_boxes[0][1]

        except Exception as e:
            last_error = e

        time.sleep(CHATGPT_DOM_POLL_SECONDS)

    raise Exception(f"Không tìm thấy ô nhập prompt thật ở dưới cùng. Lỗi gần nhất: {last_error}")


def focus_chatgpt_input_safely(driver, wait, timeout=10):
    """
    Lấy đúng ô nhập ChatGPT thật ở dưới cùng. Nội dung được gán trực tiếp
    bằng JavaScript nên không bắt buộc document.activeElement phải là hộp.
    get_chatgpt_input_box vẫn kiểm tra selector, hiển thị, kích thước và vị trí.
    """
    chatbox = get_chatgpt_input_box(driver, wait, timeout=timeout)

    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", chatbox)
    time.sleep(0.2)

    # Focus chỉ để giao diện cập nhật thuận lợi, không dùng làm điều kiện lỗi.
    driver.execute_script("arguments[0].focus();", chatbox)
    time.sleep(0.1)

    return chatbox


def get_chatgpt_send_button(driver, wait, timeout=35):
    """
    Tìm nút gửi prompt ChatGPT có bảo hiểm riêng:
    - đợi UI render
    - xử lý khi GPT đang generate
    - thử selector mới/cũ
    - refresh URL hiện tại 1 lần nếu không thấy nút gửi
    """

    def find_button_once():
        return driver.execute_script("""
            const input = document.activeElement && document.activeElement.id === 'prompt-textarea'
                ? document.activeElement
                : document.querySelector('#prompt-textarea');

            if (!input) return null;

            // Nếu ChatGPT còn đang trả lời thì chưa tìm nút Send
            const stopBtn = document.querySelector(
                'button[aria-label="Stop generating"], button[data-testid="stop-button"]'
            );
            if (stopBtn) return "GPT_STILL_GENERATING";

            const form = input.closest('form');

            const selectors = [
                'button[data-testid="send-button"]',
                'button[aria-label="Send prompt"]',
                'button[aria-label="Send message"]',
                'button[type="submit"]',
                'button svg[data-testid="send-button"]'
            ];

            if (form) {
                for (const sel of selectors) {
                    let b = form.querySelector(sel);
                    if (b) {
                        if (b.tagName.toLowerCase() === 'svg') b = b.closest('button');
                        const r = b.getBoundingClientRect();
                        const st = getComputedStyle(b);
                        if (r.width > 0 && r.height > 0 && st.display !== 'none' && st.visibility !== 'hidden') {
                            return b;
                        }
                    }
                }
            }

            // Fallback: lấy button thấp nhất gần composer
            const buttons = Array.from(document.querySelectorAll('button')).filter(b => {
                const r = b.getBoundingClientRect();
                const st = getComputedStyle(b);
                const html = (b.outerHTML || '').toLowerCase();
                const label = (
                    b.getAttribute('aria-label') ||
                    b.getAttribute('data-testid') ||
                    b.innerText ||
                    ''
                ).toLowerCase();

                const looksSend =
                    label.includes('send') ||
                    label.includes('gửi') ||
                    html.includes('send-button') ||
                    html.includes('arrow-up') ||
                    html.includes('submit');

                return looksSend &&
                    r.width > 0 &&
                    r.height > 0 &&
                    st.display !== 'none' &&
                    st.visibility !== 'hidden';
            });

            if (!buttons.length) return null;

            buttons.sort((a, b) => b.getBoundingClientRect().top - a.getBoundingClientRect().top);
            return buttons[0];
        """)

    def wait_find(label, seconds):
        end_time = time.time() + seconds
        last_state = None

        while time.time() < end_time:
            try:
                btn = find_button_once()

                if btn == "GPT_STILL_GENERATING":
                    last_state = "GPT vẫn đang generate"
                    time.sleep(0.5)
                    continue

                if btn and btn.is_enabled():
                    aria_disabled = str(btn.get_attribute("aria-disabled") or "").lower()
                    disabled = btn.get_attribute("disabled")
                    if aria_disabled == "true" or disabled is not None:
                        last_state = "Nút Send chưa được bật"
                        time.sleep(CHATGPT_DOM_POLL_SECONDS)
                        continue
                    return btn

                last_state = "Chưa thấy nút Send"
            except Exception as e:
                last_state = str(e)

            time.sleep(CHATGPT_DOM_POLL_SECONDS)

        raise Exception(f"{label}: Không tìm thấy nút gửi prompt của ChatGPT. Trạng thái cuối: {last_state}")

    try:
        return wait_find("Lần 1", timeout)
    except Exception as first_error:
        print(f"⚠️ Chưa tìm thấy nút Send: {first_error}. Reload và thử lại...")
        try:
            current_url = driver.current_url
            driver.get(current_url)
            time.sleep(10)
            focus_chatgpt_input_safely(driver, wait, timeout=25)
            return wait_find("Sau reload", 25)
        except Exception as second_error:
            raise Exception(
                "Không tìm thấy nút gửi prompt sau khi reload. "
                f"Lỗi đầu: {first_error} | Lỗi sau reload: {second_error}"
            )


def send_prompt_by_js(driver, wait, text_to_send):
    """
    Gửi prompt bằng JS nhưng vẫn bắt buộc chọn đúng ô nhập ChatGPT.
    Đã bỏ selector contenteditable/ProseMirror rộng để tránh dán nhầm vào Article Detail/Edit.
    """
    wait_chatgpt_page_ready(driver, wait, timeout=45)
    chatbox = focus_chatgpt_input_safely(driver, wait, timeout=10)

    tag_name = (chatbox.tag_name or "").lower()
    if tag_name == "textarea":
        driver.execute_script("""
            arguments[0].value = arguments[1];
            arguments[0].dispatchEvent(new Event('input', { bubbles: true }));
            arguments[0].dispatchEvent(new Event('change', { bubbles: true }));
        """, chatbox, text_to_send)
    else:
        # Với div/p contenteditable id=prompt-textarea.
        driver.execute_script("""
            arguments[0].focus();
            arguments[0].innerText = arguments[1];
            arguments[0].dispatchEvent(new InputEvent('input', {
                bubbles: true,
                inputType: 'insertText',
                data: arguments[1]
            }));
        """, chatbox, text_to_send)

    time.sleep(CHATGPT_AFTER_FILL_SECONDS)

    send_button = get_chatgpt_send_button(driver, wait, timeout=35)
    driver.execute_script("arguments[0].removeAttribute('disabled');", send_button)
    driver.execute_script("arguments[0].click();", send_button)
    print("-> Đã gửi prompt bằng JS vào đúng ô nhập ChatGPT.")
    time.sleep(1)


def send_prompt_by_real_paste(driver, wait, text_to_send):
    """
    ĐÃ ĐỔI SANG JS THUẦN (không còn dùng Clipboard hệ điều hành, không còn
    giả lập Ctrl+V/Enter thật). Giữ nguyên tên hàm để không phải sửa nơi gọi.

    Lý do đổi: pyperclip dùng chung 1 Clipboard cho toàn máy. Khi chạy nhiều
    worker song song, worker này pyperclip.copy() có thể đè nội dung ngay
    lúc worker khác chuẩn bị Ctrl+V, gây dán nhầm bài. Cách cũ cũng bắt buộc
    cửa sổ Edge phải đang có focus thật tại đúng thời điểm bấm phím.
    """
    wait_chatgpt_page_ready(driver, wait, timeout=45)
    chatbox = focus_chatgpt_input_safely(driver, wait, timeout=10)

    tag_name = (chatbox.tag_name or "").lower()
    if tag_name == "textarea":
        driver.execute_script("""
            arguments[0].value = arguments[1];
            arguments[0].dispatchEvent(new Event('input', { bubbles: true }));
            arguments[0].dispatchEvent(new Event('change', { bubbles: true }));
        """, chatbox, text_to_send)
    else:
        driver.execute_script("""
            arguments[0].focus();
            arguments[0].innerText = arguments[1];
            arguments[0].dispatchEvent(new InputEvent('input', {
                bubbles: true,
                inputType: 'insertText',
                data: arguments[1]
            }));
        """, chatbox, text_to_send)
    print("-> Đã gõ prompt bằng JS vào đúng ô nhập ChatGPT dưới cùng.")

    time.sleep(CHATGPT_AFTER_FILL_SECONDS)

    send_button = get_chatgpt_send_button(driver, wait, timeout=35)
    driver.execute_script("arguments[0].removeAttribute('disabled');", send_button)
    driver.execute_script("arguments[0].click();", send_button)
    print("-> Đã bấm gửi prompt.")
    time.sleep(0.5)


def normalize_prompt_for_compare(text):
    return re.sub(r"\s+", " ", str(text or "")).strip().casefold()


def prompt_already_sent(driver, text_to_send):
    """Kiểm tra các tin nhắn user gần nhất trước khi cho phép gửi lại."""
    expected = normalize_prompt_for_compare(text_to_send)
    if not expected:
        return False
    try:
        messages = driver.execute_script("""
            const direct = Array.from(
                document.querySelectorAll('[data-message-author-role="user"]')
            );
            const fromTurns = Array.from(
                document.querySelectorAll('div[data-testid^="conversation-turn-"]')
            ).map(turn => turn.querySelector('[data-message-author-role="user"]'))
             .filter(Boolean);
            const unique = Array.from(new Set([...direct, ...fromTurns]));
            return unique.slice(-8).map(
                el => (el.innerText || el.textContent || '').trim()
            );
        """) or []
    except Exception:
        return False

    for message in messages:
        actual = normalize_prompt_for_compare(message)
        if not actual:
            continue
        if actual == expected:
            return True
        if len(expected) >= 40 and expected in actual:
            return True
        if len(expected) >= 240 and len(actual) >= 240:
            if expected[:120] == actual[:120] and expected[-120:] == actual[-120:]:
                return True
    return False


def wait_prompt_send_signal(driver, text_to_send, timeout=10):
    """Quan sát dấu hiệu gửi thành công nhưng không dùng nó để tự gửi lại."""
    deadline = time.time() + timeout
    while time.time() < deadline:
        if prompt_already_sent(driver, text_to_send):
            return "prompt đã xuất hiện trong lịch sử chat"
        try:
            state = driver.execute_script("""
                const box = document.querySelector('#prompt-textarea');
                const boxText = box
                    ? ((box.value || box.innerText || box.textContent || '').trim())
                    : '';
                const stop = !!document.querySelector(
                    'button[aria-label="Stop generating"], '
                    + 'button[data-testid="stop-button"], .result-streaming'
                );
                return { boxEmpty: !!box && !boxText, generating: stop };
            """) or {}
            if state.get("generating"):
                return "ChatGPT đã bắt đầu trả lời"
            if state.get("boxEmpty"):
                return "ô nhập đã trống sau khi bấm gửi"
        except Exception:
            pass
        time.sleep(CHATGPT_DOM_POLL_SECONDS)
    return ""


def send_once_unless_present(driver, wait, text_to_send, send_func, label):
    """Đã click Send thì không tự gửi lại chỉ vì DOM chưa xác nhận rõ."""
    if prompt_already_sent(driver, text_to_send):
        print(f"-> {label}: Prompt đã có trong cuộc trò chuyện. Không gửi trùng.")
        return

    send_func(driver, wait, text_to_send)
    signal = wait_prompt_send_signal(
        driver, text_to_send, timeout=CHATGPT_SEND_CONFIRM_SECONDS
    )
    if signal:
        print(f"-> {label}: Xác nhận gửi thành công ({signal}).")
    else:
        raise ChatGPTSendUnconfirmedError(
            f"{label}: Đã click gửi nhưng không xác nhận được trong "
            f"{CHATGPT_SEND_CONFIRM_SECONDS}s; dừng Worker để tránh gửi trùng."
        )


def send_prompt_with_3_layers(driver, wait, text_to_send, send_func, task=None):
    """
    V2.17: chỉ gửi đúng một lần. Không reload/reset Edge rồi tự gửi lại khi
    lần đầu lỗi hoặc không xác nhận được, nhằm tránh prompt trùng. Dòng lỗi
    chỉ bị bỏ qua trong lượt hiện tại; Worker vẫn xử lý dòng kế tiếp.
    Giữ tên hàm cũ để không phải thay đổi các nơi gọi.
    """
    try:
        send_once_unless_present(driver, wait, text_to_send, send_func, "GỬI 1 LẦN")
        return driver, wait
    except Exception as e:
        row = (task or {}).get("row", 0)
        write_retry_note(
            row,
            9,
            "FINAL_ERROR",
            "CHATGPT_SEND_ONCE_FAILED",
            str(e),
        )
        raise ChatGPTSendUnconfirmedError(
            f"Gửi ChatGPT lần đầu thất bại; bỏ qua dòng trong lượt này: {e}"
        ) from e


def wait_for_gpt_done(driver, max_timeout_seconds=300):
    start_time = time.time()
    while True:
        if time.time() - start_time > max_timeout_seconds:
            raise Exception(f"ChatGPT phản hồi quá lâu hoặc bị treo (> {max_timeout_seconds}s).")
        # Giữ renderer ChatGPT tiếp tục cập nhật DOM khi Edge bị minimize.
        # Không mở lại cửa sổ và không chiếm chuột.
        keep_page_lifecycle_active(driver, quiet=True)
        is_writing = driver.execute_script("""
            let stopBtn = document.querySelector('button[aria-label="Stop generating"], button[data-testid="stop-button"]');
            let streaming = document.querySelector('.result-streaming');
            return (stopBtn !== null || streaming !== null);
        """)
        if not is_writing:
            break
        time.sleep(CHATGPT_DOM_POLL_SECONDS)
    # Chờ thêm để phản hồi cuối và các thành phần React ổn định trước khi đọc.
    time.sleep(CHATGPT_CONTENT_STABLE_SECONDS)


def capture_stable_assistant_article(
    driver,
    stable_seconds=CHATGPT_CONTENT_STABLE_SECONDS,
    max_timeout_seconds=45,
):
    """
    Chụp text + HTML của assistant cuối sau khi cả hai đã đứng yên.
    Dùng textContent thay cho innerText để không phụ thuộc cửa sổ ChatGPT
    đang được render/hiển thị hay đang minimize.
    """
    script = """
        const turns = Array.from(
            document.querySelectorAll('div[data-testid^="conversation-turn-"]')
        );
        const candidates = [];
        for (const turn of turns) {
            const assistant = turn.querySelector('[data-message-author-role="assistant"]');
            if (!assistant) continue;
            const markdown = assistant.querySelector('.markdown') || turn.querySelector('.markdown');
            if (markdown) candidates.push(markdown);
        }
        if (!candidates.length) {
            document.querySelectorAll(
                '[data-message-author-role="assistant"] .markdown, div.agent-turn .markdown'
            ).forEach(node => candidates.push(node));
        }
        if (!candidates.length) return null;

        const source = candidates[candidates.length - 1];
        const clone = source.cloneNode(true);
        clone.querySelectorAll(
            'button, svg, [data-testid*="copy"], [aria-label*="Copy"], [aria-label*="Sao chép"]'
        ).forEach(el => el.remove());
        const originalText = (clone.textContent || source.textContent || '').trim();
        return {text: originalText, html: clone.innerHTML || ''};
    """

    deadline = time.time() + max_timeout_seconds
    last_text = None
    last_html = None
    stable_since = None
    next_cdp_wakeup = 0.0

    while time.time() < deadline:
        # Nhắc lại active/focus định kỳ để React tiếp tục đưa nội dung vào DOM
        # ngay cả khi cửa sổ Edge đang minimize.
        now = time.time()
        if now >= next_cdp_wakeup:
            keep_page_lifecycle_active(driver, quiet=True)
            next_cdp_wakeup = now + 2.0

        article = driver.execute_script(script)
        text_value = str((article or {}).get("text") or "").strip()
        html_value = str((article or {}).get("html") or "").strip()

        if not text_value or not html_value:
            last_text = None
            last_html = None
            stable_since = None
            time.sleep(0.5)
            continue

        if text_value == last_text and html_value == last_html:
            if stable_since is None:
                stable_since = time.time()
            elif time.time() - stable_since >= stable_seconds:
                print(
                    f"-> Nội dung assistant đã ổn định {stable_seconds:.0f}s "
                    f"({count_words(text_value)} từ)."
                )
                return {"text": text_value, "html": html_value}
        else:
            last_text = text_value
            last_html = html_value
            stable_since = None

        time.sleep(0.5)

    raise Exception(
        f"Nội dung assistant chưa ổn định sau {max_timeout_seconds}s."
    )


def extract_content_by_js(driver):
    return driver.execute_script("""
        let blocks = document.querySelectorAll('div[data-testid*="-turn-assistant"] .markdown, div.agent-turn .markdown, .markdown');
        if (blocks.length > 0) return blocks[blocks.length - 1].textContent;
        return null;
    """)


def copy_and_save_snapshot(article, file_path):
    """Worker Word dùng snapshot HTML riêng; tuyệt đối không truy cập Selenium."""
    import win32com.client as win32

    word_app = None
    doc = None
    html_path = None
    temp_docx = file_path + f".wordtmp_{threading.get_ident()}.docx"
    try:
        if not article or not article.get("html"):
            raise Exception("Gói Word không có snapshot HTML assistant.")
        if count_words(article.get("text")) < 80:
            raise Exception("Snapshot HTML có quá ít nội dung.")

        html_document = f"""<!doctype html>
<html><head><meta charset="utf-8">
<style>
body {{ font-family: Arial, sans-serif; font-size: 11pt; line-height: 1.45; }}
h1 {{ font-size: 20pt; }} h2 {{ font-size: 16pt; }} h3 {{ font-size: 13pt; }}
p {{ margin: 0 0 6pt 0; }} li {{ margin-bottom: 3pt; }}
table {{ border-collapse: collapse; width: 100%; }}
th, td {{ border: 1px solid #999; padding: 5px; }}
</style></head><body>{article['html']}</body></html>"""

        output_folder = os.path.dirname(os.path.abspath(file_path))
        os.makedirs(output_folder, exist_ok=True)
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".html", prefix="chatgpt_word_queue_",
            dir=output_folder, encoding="utf-8", delete=False
        ) as temp_html:
            temp_html.write(html_document)
            html_path = temp_html.name

        if os.path.exists(temp_docx):
            os.remove(temp_docx)

        # Chỉ thread Word gọi COM. DispatchEx tạo một Word instance riêng cho hàng chờ.
        word_app = win32.DispatchEx("Word.Application")
        word_app.Visible = False
        word_app.DisplayAlerts = 0
        doc = word_app.Documents.Open(
            html_path, ConfirmConversions=False, ReadOnly=False,
            AddToRecentFiles=False, Encoding=65001
        )
        for para in doc.Paragraphs:
            para.SpaceAfter = 4
            para.SpaceBefore = 0

        word_app.Run(WORD_MACRO)
        doc.SaveAs2(temp_docx, FileFormat=16)
        doc.Close(False)
        doc = None

        if not is_word_ok(temp_docx):
            raise Exception("File Word tạm sau VBA không đạt kiểm tra nội dung.")
        os.replace(temp_docx, file_path)
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if word_app is not None:
            try:
                word_app.Quit()
            except Exception:
                pass
        for temp_path in (html_path, temp_docx):
            if temp_path and os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                except Exception:
                    pass


def enqueue_word_job(driver, task, word_path, word_count):
    """Chụp HTML đúng từ Edge hiện tại và giao cho Worker Word."""
    row = task["row"]
    with _WORD_PENDING_LOCK:
        if row in _WORD_PENDING_ROWS:
            print(f"-> Dòng {row}: Word đã có trong hàng chờ, không giao trùng.")
            return
        _WORD_PENDING_ROWS.add(row)

    try:
        # Snapshot chính đã được xác nhận trước đó. Vẫn kiểm tra lại để bảo hiểm,
        # nhưng chỉ cần đứng yên 0.5s thay vì lặp nguyên chu kỳ 3s.
        snapshot = capture_stable_assistant_article(
            driver, stable_seconds=0.3, max_timeout_seconds=15
        )
        if not snapshot or count_words(snapshot.get("text")) < 80:
            raise Exception("Không chụp được snapshot HTML hợp lệ để giao Worker Word.")
        chat_url = driver.current_url
        write_article_queued(row, word_path, chat_url, word_count)
        WORD_QUEUE.put({
            "row": row,
            "job_id": f"WORD_{row}_{int(time.time() * 1000)}",
            "name": task.get("name", ""),
            "word_path": word_path,
            "chat_url": chat_url,
            "word_count": int(word_count),
            "article": snapshot,
        })
        print(f"-> Dòng {row}: Đã giao Worker Word. Hàng chờ hiện có {WORD_QUEUE.qsize()} bài.")
    except Exception:
        with _WORD_PENDING_LOCK:
            _WORD_PENDING_ROWS.discard(row)
        raise


def copy_and_save_perfect(driver, file_path):
    import win32com.client as win32

    def copy_save_once(article_snapshot):
        word_app = None
        doc = None
        html_path = None

        try:
            article = article_snapshot
            if not article or not article.get("html"):
                raise Exception("Không có snapshot HTML assistant đã ổn định.")
            if count_words(article.get("text")) < 80:
                raise Exception("HTML lấy được có quá ít nội dung.")

            html_document = f"""<!doctype html>
<html><head><meta charset="utf-8">
<style>
body {{ font-family: Arial, sans-serif; font-size: 11pt; line-height: 1.45; }}
h1 {{ font-size: 20pt; }} h2 {{ font-size: 16pt; }} h3 {{ font-size: 13pt; }}
p {{ margin: 0 0 6pt 0; }} li {{ margin-bottom: 3pt; }}
table {{ border-collapse: collapse; width: 100%; }}
th, td {{ border: 1px solid #999; padding: 5px; }}
</style></head><body>{article['html']}</body></html>"""

            output_folder = os.path.dirname(os.path.abspath(file_path))
            os.makedirs(output_folder, exist_ok=True)
            with tempfile.NamedTemporaryFile(
                mode="w", suffix=".html", prefix="chatgpt_",
                dir=output_folder, encoding="utf-8", delete=False
            ) as temp_html:
                temp_html.write(html_document)
                html_path = temp_html.name

            word_app = win32.gencache.EnsureDispatch('Word.Application')
            word_app.Visible = False
            word_app.DisplayAlerts = 0

            doc = word_app.Documents.Open(
                html_path, ConfirmConversions=False, ReadOnly=False,
                AddToRecentFiles=False, Encoding=65001
            )

            for para in doc.Paragraphs:
                para.SpaceAfter = 4
                para.SpaceBefore = 0

            word_app.Run(WORD_MACRO)

            if TEST_OVERWRITE_WORD and os.path.exists(file_path):
                os.remove(file_path)
            doc.SaveAs2(file_path, FileFormat=16)

            doc.Close()
            doc = None

        finally:
            if doc is not None:
                try:
                    doc.Close(False)
                except Exception:
                    pass

            if word_app is not None:
                try:
                    word_app.Quit()
                except Exception:
                    pass

            if html_path and os.path.exists(html_path):
                try:
                    os.remove(html_path)
                except Exception:
                    pass

    try:
        global _LAST_STABLE_ARTICLE
        if not _LAST_STABLE_ARTICLE:
            _LAST_STABLE_ARTICLE = capture_stable_assistant_article(driver)
        copy_save_once(_LAST_STABLE_ARTICLE)

    except Exception as first_error:
        print(f"⚠️ Copy/lưu Word lần 1 lỗi: {first_error}")
        print("-> Bảo hiểm: load lại URL hiện tại rồi thử copy/lưu Word thêm 1 lần.")

        try:
            current_url = driver.current_url
            driver.get(current_url)
            time.sleep(ARTICLE_RELOAD_WAIT_SECONDS)

            # Chờ ChatGPT ổn định lại rồi thử copy lần 2
            get_gpt_content_after_wait(driver, 10, "Copy Word sau reload")

            copy_save_once(_LAST_STABLE_ARTICLE)

        except Exception as second_error:
            raise Exception(
                "Copy/lưu Word lỗi sau khi đã reload URL và thử lại 1 lần. "
                f"Lỗi lần 1: {first_error} | Lỗi lần 2: {second_error}"
            )


def parse_briefs_logic(text):
    text = str(text or "").strip()
    m = re.search(r"===BRIEF_1===\s*(.*?)\s*===BRIEF_2===\s*(.*)", text, flags=re.I | re.S)
    if not m:
        raise Exception("Không tìm thấy đúng định dạng ===BRIEF_1=== và ===BRIEF_2===")
    brief1 = re.sub(r"^```.*?\n|\n```$", "", m.group(1).strip(), flags=re.S).strip()
    brief2 = re.sub(r"^```.*?\n|\n```$", "", m.group(2).strip(), flags=re.S).strip()
    if not brief1 or not brief2:
        raise Exception("Brief 1 hoặc Brief 2 bị rỗng.")
    return brief1, brief2

def save_word_from_current_chat(driver, task, label=""):
    row = task["row"]

    content_text = get_gpt_content_after_wait(driver, ARTICLE_WAIT_SECONDS, label)

    if not content_text:
        write_retry_note(row, 1, "ARTICLE_RESUME", "NO_CONTENT", "Không lấy được nội dung từ URL H.")
        reload_current_url(driver, ARTICLE_RELOAD_WAIT_SECONDS, "Resume URL H")
        content_text = get_gpt_content_after_wait(driver, 10, "Resume sau reload")

    if not content_text:
        return None, None, "Không lấy được nội dung từ URL hiện tại."

    word_count = count_words(content_text)
    print(f"-> Số từ lấy từ chat hiện tại: {word_count}")

    if word_count < MIN_WORDS:
        print(f"-> Nội dung dưới {MIN_WORDS} từ. Gửi prompt bổ sung trong chính URL hiện tại.")
        send_prompt_by_real_paste(driver, WebDriverWait(driver, 45), SECOND_PROMPT)
        content_text = get_gpt_content_after_wait(driver, ARTICLE_WAIT_SECONDS, "Kéo dài bài trong URL H")

        if not content_text:
            return None, None, "Không lấy được nội dung sau prompt bổ sung trong URL H."

        word_count = count_words(content_text)
        print(f"-> Số từ sau bổ sung: {word_count}")

        if word_count < MIN_WORDS:
            return None, None, f"Bài vẫn dưới {MIN_WORDS} từ: {word_count}"

    if not check_keyword_exists(content_text, task["name"]):
        return None, None, f"Nội dung không chứa từ khóa/tên file: {task['name']}"

    word_path = make_output_path(task["web"], task["name"])
    print("-> Đang giao snapshot HTML cho Worker Word chuyên dụng...")
    enqueue_word_job(driver, task, word_path, word_count)

    return word_path, word_count, None

def write_article_if_needed(driver, wait, task):
    row = task["row"]

    if is_word_ok(task["word_path"]) and not TEST_OVERWRITE_WORD:
        print(f"Dòng {row}: Word đã có và đọc được. Bỏ qua viết bài.")
        if task["article_status"].upper() != STATUS_OK:
            write_value(row, COL_ARTICLE_STATUS, STATUS_OK)
        return read_task(row), driver, wait

    if not task["web"] or not task["name"] or not task["prompt"]:
        raise Exception("Thiếu B/C/D nên không thể viết bài.")

    print(f"\n[DÒNG {row}] Bắt đầu xử lý tạo Word...")

    write_article_running(row)

    # =====================================================
    # ƯU TIÊN 1: Nếu cột H có URL thì mở H trước
    # =====================================================
    if task["chat_url"]:
        print(f"-> Cột H có URL. Ưu tiên mở lại chat cũ: {task['chat_url']}")

        try:
            driver.get(task["chat_url"])
            time.sleep(2)

            word_path, word_count, err = save_word_from_current_chat(driver, task, "Resume từ URL H")

            if word_path:
                print(f"-> Dòng {row}: Đã xếp Word từ URL H vào hàng chờ.")
                return read_task(row), driver, wait

            print(f"⚠️ Không lưu được Word từ URL H: {err}")
            write_retry_note(row, 1, "ARTICLE_RESUME", "RESUME_H_FAILED", err)

        except Exception as e:
            print(f"⚠️ URL H lỗi, sẽ xét tạo mới từ F: {e}")
            write_retry_note(row, 2, "ARTICLE_RESUME", "URL_H_ERROR", e)

    # =====================================================
    # ƯU TIÊN 2: Chỉ tạo mới từ F khi H trống hoặc H lỗi
    # =====================================================
    if not task["gpt_url"]:
        raise Exception("Không có URL H dùng được và thiếu F nên không thể tạo mới.")

    print("-> Không dùng được URL H. Bắt đầu tạo mới từ URL F...")

    driver.get(task["gpt_url"])
    time.sleep(1)

    driver, wait = send_prompt_with_3_layers(driver, wait, task["prompt"], send_prompt_by_js, task=task)
    write_value(row, COL_CHAT_URL, driver.current_url)

    content_text = get_gpt_content_after_wait(driver, ARTICLE_WAIT_SECONDS, "Viết bài lần 1")

    if not content_text:
        raise Exception(
            "Không lấy được nội dung sau lần gửi đầu; không reload hoặc gửi lại."
        )

    word_count = count_words(content_text)
    print(f"-> Số từ ước tính bài viết: {word_count}")

    if word_count < MIN_WORDS:
        print(f"-> Bài dưới {MIN_WORDS} từ. Gửi prompt kéo dài bài.")
        send_prompt_by_real_paste(driver, wait, SECOND_PROMPT)
        content_text = get_gpt_content_after_wait(driver, ARTICLE_WAIT_SECONDS, "Viết bài prompt 2")

        if not content_text:
            raise Exception("Không lấy được nội dung sau prompt kéo dài.")

        word_count = count_words(content_text)
        print(f"-> Số từ sau kéo dài: {word_count}")

        if word_count < MIN_WORDS:
            raise ArticleTooShortError(
                f"Bài viết vẫn dưới {MIN_WORDS} từ sau prompt kéo dài: {word_count} từ."
            )

    if not check_keyword_exists(content_text, task["name"]):
        raise Exception(f"Bài viết không chứa từ khóa/tên file: {task['name']}")

    word_path = make_output_path(task["web"], task["name"])
    print("-> Đang giao snapshot HTML cho Worker Word chuyên dụng...")
    enqueue_word_job(driver, task, word_path, word_count)
    print(f"-> Dòng {row}: Đã xếp Word từ URL mới vào hàng chờ.")
    return read_task(row), driver, wait

def try_get_briefs_from_current_answer(driver):
    raw_answer = extract_content_by_js(driver)
    if not raw_answer:
        raise Exception("Không lấy được phản hồi Brief.")

    brief1, brief2 = parse_briefs_logic(raw_answer)

    def invalid(text):
        text = str(text or "").strip()

        if not text:
            return True

        # chỉ toàn dấu chấm: ., .., ..., .....
        if text.replace(".", "").strip() == "":
            return True

        # quá ngắn thì coi như lỗi
        if len(text) < 20:
            return True

        return False

    if invalid(brief1):
        raise Exception("Brief 1 không hợp lệ.")

    if invalid(brief2):
        raise Exception("Brief 2 không hợp lệ.")

    return brief1, brief2


def brief_if_needed(driver, wait, task):
    row = task["row"]
    if task["brief1"] and task["brief2"]:
        print(f"Dòng {row}: Brief đã có. Bỏ qua xin Brief.")
        if str(cell(row, COL_BRIEF_STATUS).value or "").strip().upper() != STATUS_OK_BRIEF:
            write_value(row, COL_BRIEF_STATUS, STATUS_OK_BRIEF)
        return read_task(row), driver, wait

    if not article_ready_for_downstream(task):
        raise Exception("Bài viết chưa được giao cho Worker Word nên không xin Brief.")

    print(f"\n[DÒNG {row}] Bắt đầu xin Brief ảnh...")
    write_value(row, COL_BRIEF_STATUS, STATUS_RUNNING_BRIEF, save=False)

    # Ưu tiên dùng tab ChatGPT hiện tại nếu vừa viết xong. Nếu đang resume thì mở lại URL H.
    current_url = driver.current_url or ""
    target_url = task["chat_url"]
    if target_url and target_url not in current_url:
        print("-> Đang mở lại URL chat cũ ở cột H để xin Brief.")
        driver.get(target_url)
        time.sleep(2)

    # LẦN 1: gửi prompt xin Brief bằng Bảo Hiểm 3 Lớp
    driver, wait = send_prompt_with_3_layers(driver, wait, ASK_BRIEF_PROMPT, send_prompt_by_real_paste, task=task)
    get_gpt_content_after_wait(driver, BRIEF_WAIT_SECONDS, "Brief lần 1")

    brief1 = brief2 = None
    first_error = None
    try:
        brief1, brief2 = try_get_briefs_from_current_answer(driver)
    except Exception as e:
        first_error = e
        print(f"⚠️ Brief lần 1 chưa đúng: {e}")

    if not brief1 or not brief2:
        write_retry_note(
            row,
            9,
            "BRIEF",
            "BRIEF_FIRST_RESPONSE_INVALID",
            first_error or "Brief rỗng/sai cấu trúc sau lần gửi đầu.",
        )
        raise Exception(
            f"Brief lần đầu không hợp lệ; không reload hoặc gửi lại: {first_error}"
        )

    wb, sh = get_sheet()
    sh.range(f"{COL_BRIEF_1}{row}").value = brief1
    sh.range(f"{COL_BRIEF_2}{row}").value = brief2
    sh.range(f"{COL_BRIEF_STATUS}{row}").value = STATUS_OK_BRIEF
    sh.range(f"{COL_BRIEF_1}{row}").api.WrapText = False
    sh.range(f"{COL_BRIEF_2}{row}").api.WrapText = False
    wb.save()
    print(f"-> Dòng {row}: Đã lưu Brief vào O/P.")
    return read_task(row), driver, wait


# =====================================================
# GEMINI: TẠO ẢNH / DOWNLOAD
# =====================================================
def build_gemini_prompt(brief):
    brief = str(brief or "").strip()
    return f"""
Hãy tạo ra chính xác 1 hình minh họa chất lượng cao.

=========================
BRIEF ẢNH CHI TIẾT
=========================

{brief}

=========================
CÁCH TRÌNH BÀY HÌNH ẢNH
=========================

Thể hiện Brief bằng ngôn ngữ editorial tối giản dành cho website cao cấp.

Brief là nguồn nội dung chính. Cách trình bày hình ảnh chỉ điều chỉnh cách tổ chức, nhấn mạnh và thể hiện Brief; không được làm sai thông điệp hoặc bản chất sự việc trong Brief.

=========================
YÊU CẦU BẮT BUỘC VỀ ẢNH
=========================

- Tạo đúng 1 ảnh duy nhất.
- Ảnh ngang, tỷ lệ 4:3.
- Giữ đúng nội dung chính của Brief.
- Mọi chi tiết phải đúng logic thực tế.

Chỉ xuất ra ảnh, không giải thích gì thêm.
""".strip()


def get_output_folder(word_path):
    folder = os.path.dirname(str(word_path).strip())
    if not folder:
        raise Exception("Không lấy được thư mục từ cột G.")
    os.makedirs(folder, exist_ok=True)
    return folder


def get_unique_image_path(folder, base_name):
    """
    Cấp phát đường dẫn ảnh không trùng tên.

    Ví dụ:
    - Tên bài 1.png
    - Tên bài 1 2.png
    - Tên bài 1 3.png

    Kiểm tra đồng thời PNG/JPG/JPEG/WEBP và khóa trong RAM để các worker
    không thể cùng chọn một tên trước khi file ảnh được ghi xuống ổ đĩa.
    """
    folder = os.path.abspath(folder)
    base_name = safe_filename(base_name)
    supported_exts = (".png", ".jpg", ".jpeg", ".webp")

    with _IMAGE_PATH_LOCK:
        def name_is_available(candidate_base):
            reserved_key = os.path.normcase(
                os.path.join(folder, candidate_base)
            )
            if reserved_key in _RESERVED_IMAGE_BASES:
                return False
            return not any(
                os.path.exists(os.path.join(folder, candidate_base + ext))
                for ext in supported_exts
            )

        candidate_base = base_name
        if not name_is_available(candidate_base):
            i = 2
            while True:
                candidate_base = f"{base_name} {i}"
                if name_is_available(candidate_base):
                    break
                i += 1

        reserved_key = os.path.normcase(
            os.path.join(folder, candidate_base)
        )
        _RESERVED_IMAGE_BASES.add(reserved_key)
        return os.path.join(folder, candidate_base + ".png")

def list_download_files():
    ensure_temp_download_dir()
    files = set()
    download_dir = get_temp_download_dir()
    for name in os.listdir(download_dir):
        path = os.path.join(download_dir, name)
        if not os.path.isfile(path):
            continue
        if name.endswith(".crdownload") or name.endswith(".tmp"):
            continue
        files.add(path)
    return files


def wait_new_download(before_files, timeout=180):
    print("-> Đang chờ file tải về...")
    start = time.time()
    while time.time() - start < timeout:
        now_files = list_download_files()
        candidates = list(now_files - before_files)
        if candidates:
            newest = max(candidates, key=os.path.getmtime)
            size1 = os.path.getsize(newest)
            time.sleep(1)
            size2 = os.path.getsize(newest)
            if size1 == size2 and size2 > 10 * 1024:
                return newest
        time.sleep(1)
    raise Exception("Không thấy file ảnh mới trong thư mục tải xuống.")


def resize_image_max_800(image_path, max_size=MAX_IMAGE_SIZE):
    try:
        with Image.open(image_path) as img:
            img.load()
            w, h = img.size
            longest = max(w, h)
            if longest <= max_size:
                return image_path
            ratio = max_size / float(longest)
            new_w = int(w * ratio)
            new_h = int(h * ratio)
            img = img.resize((new_w, new_h), Image.LANCZOS)
            ext = os.path.splitext(image_path)[1].lower()
            if ext in [".jpg", ".jpeg"]:
                if img.mode in ("RGBA", "P"):
                    img = img.convert("RGB")
                img.save(image_path, quality=92, optimize=True)
            else:
                img.save(image_path, optimize=True)
            print(f"-> Đã resize ảnh về tối đa {max_size}px: {new_w}x{new_h}")
            return image_path
    except Exception as e:
        print(f"⚠️ Không resize được ảnh, giữ file gốc: {e}")
        return image_path


def move_downloaded_image(downloaded_path, final_path):
    ext = os.path.splitext(downloaded_path)[1].lower()
    if ext not in [".png", ".jpg", ".jpeg", ".webp"]:
        ext = ".png"
    final_path = os.path.splitext(final_path)[0] + ext
    if os.path.exists(final_path):
        os.remove(final_path)
    shutil.move(downloaded_path, final_path)
    resize_image_max_800(final_path)
    return final_path


def save_image_directly_from_element(driver, img_element, final_path):
    """Đọc blob ảnh ngay trong trang Gemini; không phụ thuộc nút Download."""
    result = driver.execute_async_script("""
        const img = arguments[0];
        const done = arguments[arguments.length - 1];
        const src = img.currentSrc || img.src || '';
        if (!src) { done({ok:false, error:'Ảnh không có src'}); return; }
        if (!img.complete || !img.naturalWidth || !img.naturalHeight) {
            done({
                ok:false,
                error:'Thẻ ảnh chưa load/decode xong',
                complete:Boolean(img.complete),
                naturalWidth:img.naturalWidth || 0,
                naturalHeight:img.naturalHeight || 0
            });
            return;
        }

        // Đọc pixel đã decode bằng canvas. Cách này không cần fetch lại blob
        // URL vốn đã được kiểm chứng là có thể bị Gemini thu hồi.
        try {
            const canvas = document.createElement('canvas');
            canvas.width = img.naturalWidth;
            canvas.height = img.naturalHeight;
            const context = canvas.getContext('2d');
            context.drawImage(img, 0, 0);
            const data = canvas.toDataURL('image/png');
            if (data && data.length > 20000) {
                done({ok:true, data:data, type:'image/png', method:'canvas'});
                return;
            }
        } catch (canvasError) {
            // Nếu canvas bị giới hạn bảo mật thì thử fetch blob bên dưới.
        }

        fetch(src, {credentials:'include'})
            .then(response => {
                if (!response.ok) throw new Error('HTTP ' + response.status);
                return response.blob();
            })
            .then(blob => {
                const reader = new FileReader();
                reader.onloadend = () => done({
                    ok:true,
                    data:reader.result,
                    type:blob.type || ''
                });
                reader.onerror = () => done({ok:false, error:'FileReader lỗi'});
                reader.readAsDataURL(blob);
            })
            .catch(error => done({ok:false, error:String(error)}));
    """, img_element)

    if not result or not result.get("ok"):
        raise Exception("Không đọc trực tiếp được blob Gemini: " + str(result))

    data_url = result.get("data") or ""
    if "," not in data_url:
        raise Exception("Gemini trả về data URL không hợp lệ.")
    header, encoded = data_url.split(",", 1)
    mime = (result.get("type") or header).lower()
    if "jpeg" in mime or "jpg" in mime:
        ext = ".jpg"
    elif "webp" in mime:
        ext = ".webp"
    else:
        ext = ".png"

    final_path = os.path.splitext(final_path)[0] + ext
    image_bytes = base64.b64decode(encoded)
    if len(image_bytes) < 10 * 1024:
        raise Exception(f"Blob ảnh quá nhỏ: {len(image_bytes)} bytes.")

    with open(final_path, "wb") as image_file:
        image_file.write(image_bytes)

    # Mở thử bằng PIL để chắc chắn bytes vừa lưu là ảnh thật.
    try:
        with Image.open(final_path) as image:
            image.verify()
    except Exception:
        try:
            os.remove(final_path)
        except Exception:
            pass
        raise

    resize_image_max_800(final_path)
    print("-> Đã lưu ảnh trực tiếp từ blob Gemini, không cần mở viewer.")
    return final_path


def send_prompt_to_gemini(driver, wait, text_to_send):
    """
    ĐÃ ĐỔI SANG JS THUẦN: không còn pyperclip (Clipboard hệ điều hành dùng
    chung cho cả máy) và không còn send_keys Ctrl+A/Backspace/Ctrl+V/Enter
    thật (cần cửa sổ đang focus). Toàn bộ thao tác gõ + gửi đều qua
    execute_script nên chạy được cả khi cửa sổ không active và cả headless.
    """
    chatbox_xpath = (
        "//div[@contenteditable='true' and @role='textbox']"
        "|//rich-textarea//div[@contenteditable='true']"
        "|//div[contains(@class,'textarea') and @contenteditable='true']"
    )

    for attempt in range(1, 4):
        enter_started = False
        try:
            # Gemini có thể vừa hiện ô nhập đã render lại ngay. Chờ ngắn rồi
            # tìm lại để luôn thao tác trên phần tử mới nhất.
            wait.until(EC.element_to_be_clickable((By.XPATH, chatbox_xpath)))
            time.sleep(0.8)
            chatbox = wait.until(
                EC.element_to_be_clickable((By.XPATH, chatbox_xpath))
            )
            driver.execute_script(
                "arguments[0].scrollIntoView({block:'center'});", chatbox
            )

            # Gõ nội dung bằng JS thay vì Ctrl+A/Backspace/Ctrl+V thật.
            driver.execute_script("""
                const el = arguments[0];
                const text = arguments[1];
                el.focus();
                el.innerText = text;
                el.dispatchEvent(new InputEvent('input', {
                    bubbles: true,
                    inputType: 'insertText',
                    data: text
                }));
            """, chatbox, text_to_send)
            time.sleep(0.5)

            # Từ thời điểm gọi gửi, không tự gửi lại nếu trạng thái không rõ,
            # tránh Gemini nhận hai prompt giống nhau.
            enter_started = True

            # Ưu tiên bấm nút gửi thật bằng JS; chỉ dùng phím Enter giả lập
            # qua JS làm phương án dự phòng nếu không thấy nút gửi.
            send_clicked = driver.execute_script("""
                const el = arguments[0];
                const form = el.closest('form') || el.closest('rich-textarea') || document;
                const btn = form.querySelector(
                    'button[aria-label*="Send" i], button[aria-label*="Gửi" i], ' +
                    'button[aria-label*="gửi" i], button[type="submit"]'
                );
                if (btn && !btn.disabled) { btn.click(); return true; }
                return false;
            """, chatbox)

            if not send_clicked:
                driver.execute_script("""
                    const el = arguments[0];
                    const opts = {
                        key: 'Enter', code: 'Enter', keyCode: 13, which: 13,
                        bubbles: true, cancelable: true
                    };
                    el.dispatchEvent(new KeyboardEvent('keydown', opts));
                    el.dispatchEvent(new KeyboardEvent('keyup', opts));
                """, chatbox)

            print("-> Đã gửi prompt sang Gemini bằng JS.")
            time.sleep(1)
            return
        except StaleElementReferenceException:
            if enter_started:
                raise Exception(
                    "Ô nhập Gemini đổi đúng lúc gửi; dừng để tránh gửi trùng prompt."
                )
            print(
                f"⚠️ Ô nhập Gemini vừa được render lại, "
                f"đang tìm lại ({attempt}/3)..."
            )
            time.sleep(0.5)

    raise Exception("Ô nhập Gemini liên tục thay đổi sau 3 lần thử.")


def open_gemini_and_wait_ready(driver, timeout=30):
    """Mở Gemini; tự tìm lại ô nhập nếu DOM render lại trước lúc gửi prompt."""
    started = time.time()
    driver.get(GEMINI_URL)
    keep_page_lifecycle_active(driver)
    chatbox_xpath = (
        "//div[@contenteditable='true' and @role='textbox']"
        "|//rich-textarea//div[@contenteditable='true']"
        "|//div[contains(@class,'textarea') and @contenteditable='true']"
    )

    def latest_chatbox_is_ready(current_driver):
        try:
            elements = current_driver.find_elements(By.XPATH, chatbox_xpath)
            for element in reversed(elements):
                try:
                    if element.is_displayed() and element.is_enabled():
                        return True
                except StaleElementReferenceException:
                    continue
            return False
        except StaleElementReferenceException:
            return False

    last_error = None
    for attempt in range(1, 6):
        try:
            WebDriverWait(
                driver,
                timeout if attempt == 1 else 5,
                poll_frequency=0.25,
                ignored_exceptions=(StaleElementReferenceException,),
            ).until(latest_chatbox_is_ready)
            # Gemini hay thay ô nhập ngay sau lần xuất hiện đầu tiên. Kiểm tra lại
            # một nhịp nữa nhưng không giữ WebElement cũ.
            time.sleep(0.5)
            if latest_chatbox_is_ready(driver):
                print(
                    f"-> Gemini đã sẵn sàng sau {time.time() - started:.1f}s "
                    f"(kiểm tra DOM lần {attempt}/5)."
                )
                return
        except Exception as exc:
            raise_if_driver_transport_error(exc)
            last_error = exc
        print(
            f"⚠️ Ô nhập Gemini vừa render lại; đang tìm phần tử mới "
            f"({attempt}/5)..."
        )
        time.sleep(0.5)

    raise Exception(
        "Gemini liên tục thay ô nhập trước khi gửi prompt sau 5 lần thử. "
        f"Lỗi cuối: {last_error}"
    )


def get_current_large_image_srcs(driver):
    srcs = set()
    imgs = driver.find_elements(By.XPATH, "//img")
    for img in imgs:
        try:
            src = img.get_attribute("src") or ""
            w = driver.execute_script("return arguments[0].naturalWidth || 0;", img)
            h = driver.execute_script("return arguments[0].naturalHeight || 0;", img)
            if src and w >= 250 and h >= 250:
                srcs.add(src)
        except Exception:
            pass
    return srcs


def get_current_gemini_blob_srcs(driver):
    """Lấy URL blob ảnh Gemini hiện có, không phụ thuộc ảnh đã render hay chưa."""
    return set(driver.execute_script("""
        return Array.from(document.querySelectorAll('img'))
            .map(img => img.currentSrc || img.src || '')
            .filter(src => src.startsWith(
                'blob:https://gemini.google.com/'
            ));
    """))


def wait_new_gemini_completed_image(driver, old_srcs, timeout=240):
    """
    V2.22: chờ ảnh Gemini theo 3 lớp, vẫn chống lấy nhầm blob cũ.

    Lớp 1: nút Download + blob mới trong đúng response (logic V2.21).
    Lớp 2: nếu nút Download chưa xuất hiện, nhận blob mới đã load hoàn chỉnh
            và có kích thước ổn định qua 3 lần poll liên tiếp.
    Lớp 3: fallback quét response/model output mới nhất, chọn blob mới có
            kích thước thực lớn nhất và đã load hoàn chỉnh.

    Code không click nút Download; ảnh vẫn được lưu bằng canvas ở bước sau.
    """
    print(
        "-> Đang chờ ảnh Gemini theo 3 lớp: "
        "Download+blob -> blob mới ổn định -> response mới nhất..."
    )
    deadline = time.time() + timeout
    old_srcs = set(old_srcs or [])
    next_cdp_wakeup = 0.0
    stable_src = None
    stable_size = None
    stable_count = 0

    while time.time() < deadline:
        now = time.time()
        if now >= next_cdp_wakeup:
            keep_page_lifecycle_active(driver, quiet=True)
            next_cdp_wakeup = now + 2.0

        # LỚP 1: giữ nguyên tiêu chí chắc nhất của V2.21.
        result = driver.execute_script("""
            const oldSrcs = new Set(arguments[0]);
            const buttons = Array.from(document.querySelectorAll(
                'button[aria-label]'
            )).filter(button => {
                const label = (button.getAttribute('aria-label') || '')
                    .toLocaleLowerCase();
                return (
                    label.includes('tải hình ảnh có kích thước đầy đủ') ||
                    label.includes('download full-size image') ||
                    label.includes('download full size image')
                );
            });

            for (const button of buttons.reverse()) {
                const response =
                    button.closest('message-content') ||
                    button.closest('model-response') ||
                    button.closest('response-element') ||
                    button.closest('generated-image');
                if (!response) continue;

                const fresh = Array.from(response.querySelectorAll('img'))
                    .filter(img => {
                        const src = img.currentSrc || img.src || '';
                        return src.startsWith(
                            'blob:https://gemini.google.com/'
                        ) && !oldSrcs.has(src);
                    });
                if (fresh.length) return fresh[fresh.length - 1];
            }
            return null;
        """, list(old_srcs))

        if result is not None:
            src = result.get_attribute("src") or ""
            print(
                "-> Gemini LỚP 1: có nút Download + blob mới: "
                f"{src[:100]}"
            )
            return result

        # LỚP 2: blob mới đã load xong và kích thước ổn định 3 nhịp.
        candidate = driver.execute_script("""
            const oldSrcs = new Set(arguments[0]);
            const fresh = Array.from(document.querySelectorAll('img'))
                .filter(img => {
                    const src = img.currentSrc || img.src || '';
                    return src.startsWith('blob:https://gemini.google.com/') &&
                           !oldSrcs.has(src) &&
                           img.complete &&
                           img.naturalWidth >= 250 &&
                           img.naturalHeight >= 250;
                });
            if (!fresh.length) return null;
            return fresh[fresh.length - 1];
        """, list(old_srcs))

        if candidate is not None:
            info = driver.execute_script("""
                const img = arguments[0];
                return {
                    src: img.currentSrc || img.src || '',
                    w: img.naturalWidth || 0,
                    h: img.naturalHeight || 0,
                    complete: Boolean(img.complete)
                };
            """, candidate)
            key = info.get('src') or ''
            size = (int(info.get('w') or 0), int(info.get('h') or 0))
            if key == stable_src and size == stable_size:
                stable_count += 1
            else:
                stable_src = key
                stable_size = size
                stable_count = 1

            if stable_count >= 3:
                print(
                    "-> Gemini LỚP 2: blob mới đã load và ổn định "
                    f"3 nhịp ({size[0]}x{size[1]}): {key[:100]}"
                )
                return candidate
        else:
            stable_src = None
            stable_size = None
            stable_count = 0

        # LỚP 3: fallback trong response/model output mới nhất. Chỉ nhận blob
        # mới, đã complete và >=250px; chọn ảnh có diện tích lớn nhất.
        fallback = driver.execute_script("""
            const oldSrcs = new Set(arguments[0]);
            const selectors = [
                'message-content',
                'model-response',
                'response-element',
                'generated-image'
            ];
            const responses = Array.from(document.querySelectorAll(
                selectors.join(',')
            ));
            for (const response of responses.reverse()) {
                const fresh = Array.from(response.querySelectorAll('img'))
                    .filter(img => {
                        const src = img.currentSrc || img.src || '';
                        return src.startsWith(
                            'blob:https://gemini.google.com/'
                        ) && !oldSrcs.has(src) && img.complete &&
                        img.naturalWidth >= 250 && img.naturalHeight >= 250;
                    })
                    .sort((a, b) =>
                        (b.naturalWidth * b.naturalHeight) -
                        (a.naturalWidth * a.naturalHeight)
                    );
                if (fresh.length) return fresh[0];
            }
            return null;
        """, list(old_srcs))

        if fallback is not None:
            # Không lấy ngay ở nhịp đầu để tránh ảnh preview đang thay đổi.
            fb = driver.execute_script("""
                const img = arguments[0];
                return {
                    src: img.currentSrc || img.src || '',
                    w: img.naturalWidth || 0,
                    h: img.naturalHeight || 0
                };
            """, fallback)
            fb_key = fb.get('src') or ''
            fb_size = (int(fb.get('w') or 0), int(fb.get('h') or 0))
            if fb_key == stable_src and fb_size == stable_size and stable_count >= 2:
                print(
                    "-> Gemini LỚP 3: nhận ảnh từ response mới nhất "
                    f"({fb_size[0]}x{fb_size[1]}): {fb_key[:100]}"
                )
                return fallback

        time.sleep(0.5)

    raise Exception(
        "Hết thời gian chờ Gemini: cả 3 lớp đều không xác nhận được ảnh mới."
    )

def force_gemini_image_load(driver, img_element, timeout=45):
    """
    Bỏ lazy-load và chủ động yêu cầu Edge load/decode ảnh. Cơ chế này đã
    được kiểm chứng khi cửa sổ Edge đang minimize.
    """
    keep_page_lifecycle_active(driver)
    driver.set_script_timeout(timeout + 5)
    result = driver.execute_async_script("""
        const img = arguments[0];
        const timeoutMs = arguments[1] * 1000;
        const done = arguments[arguments.length - 1];
        let finished = false;

        function finish(value) {
            if (finished) return;
            finished = true;
            done({
                ...value,
                complete: Boolean(img.complete),
                naturalWidth: img.naturalWidth || 0,
                naturalHeight: img.naturalHeight || 0,
                src: img.currentSrc || img.src || ''
            });
        }

        // Không chờ trình duyệt tự xử lý loading="lazy" khi minimize.
        img.loading = 'eager';
        img.removeAttribute('loading');
        try { img.fetchPriority = 'high'; } catch (_) {}
        img.style.contentVisibility = 'visible';
        img.style.visibility = 'visible';
        img.style.display = 'block';
        img.scrollIntoView({block:'center', inline:'center'});

        if (img.complete && img.naturalWidth > 0 && img.naturalHeight > 0) {
            finish({ok:true, method:'already_complete'});
            return;
        }

        const timer = setTimeout(() => {
            finish({ok:false, method:'decode_timeout'});
        }, timeoutMs);

        img.addEventListener('load', () => {
            clearTimeout(timer);
            finish({ok:true, method:'load_event'});
        }, {once:true});

        img.addEventListener('error', () => {
            clearTimeout(timer);
            finish({ok:false, method:'error_event'});
        }, {once:true});

        if (typeof img.decode === 'function') {
            img.decode().then(() => {
                clearTimeout(timer);
                finish({ok:true, method:'img_decode'});
            }).catch(() => {
                // Vẫn chờ sự kiện load cho tới timeout.
            });
        }
    """, img_element, timeout)

    width = int(result.get("naturalWidth") or 0) if result else 0
    height = int(result.get("naturalHeight") or 0) if result else 0
    if not result or not result.get("ok") or width <= 0 or height <= 0:
        raise Exception(f"Không ép tải/decode được ảnh Gemini: {result}")

    print(
        f"-> Đã ép tải ảnh Gemini thành công bằng "
        f"{result.get('method')}: {width}x{height}"
    )
    return result


def wait_new_large_image(driver, old_srcs, timeout=240):
    print("-> Đang chờ ảnh mới Gemini xuất hiện...")
    start = time.time()
    while time.time() - start < timeout:
        imgs = driver.find_elements(By.XPATH, "//img")
        for img in reversed(imgs):
            try:
                src = img.get_attribute("src") or ""
                w = driver.execute_script("return arguments[0].naturalWidth || 0;", img)
                h = driver.execute_script("return arguments[0].naturalHeight || 0;", img)
                if src and w >= 250 and h >= 250 and src not in old_srcs:
                    print(f"-> Đã tìm thấy ảnh mới: {w}x{h}")
                    return img
            except Exception:
                pass
        time.sleep(0.5)
    raise Exception("Không tìm thấy ảnh mới sau khi gửi prompt.")


def debug_images_on_page(driver):
    imgs = driver.find_elements(By.XPATH, "//img")
    print(f"\nTổng số img trên trang: {len(imgs)}")
    for i, img in enumerate(imgs):
        try:
            src = img.get_attribute("src") or ""
            w = driver.execute_script("return arguments[0].naturalWidth || 0;", img)
            h = driver.execute_script("return arguments[0].naturalHeight || 0;", img)
            rect = driver.execute_script("""
                const r = arguments[0].getBoundingClientRect();
                return {
                    x: Math.round(r.x),
                    y: Math.round(r.y),
                    width: Math.round(r.width),
                    height: Math.round(r.height)
                };
            """, img)
            if w >= 200 and h >= 200:
                print(f"[IMG {i}] size={w}x{h} rect={rect} src={src[:100]}")
        except Exception as e:
            print(f"[IMG {i}] lỗi đọc: {e}")


def open_viewer_and_download(driver, img_element, timeout=30):
    print("-> Click ảnh mới để mở viewer...")
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", img_element)
    time.sleep(1)

    driver.execute_script("""
        const img = arguments[0];
        img.scrollIntoView({block:'center'});
        img.click();
    """, img_element)

    download_xpath = (
        "(//button[contains(@aria-label,'Tải xuống') "
        "or contains(@aria-label,'Download') "
        "or .//mat-icon[contains(text(),'download')] "
        "or .//mat-icon[contains(@data-mat-icon-name,'download')]])[last()]"
    )

    for attempt in range(1, 4):
        print(f"-> Tìm nút tải xuống lần {attempt}...")
        try:
            btn = WebDriverWait(driver, timeout).until(
                EC.element_to_be_clickable((By.XPATH, download_xpath))
            )
            driver.execute_script("arguments[0].click();", btn)
            print(f"-> Đã click nút tải xuống lần {attempt}.")
            return True
        except Exception as e:
            print(f"⚠️ Không tìm thấy/click được nút tải lần {attempt}: {e}")
            try:
                driver.execute_script("""
                    const img = arguments[0];
                    img.scrollIntoView({block:'center'});
                    img.click();
                """, img_element)
                time.sleep(0.5)
            except Exception:
                pass

    raise Exception("Không tìm được nút tải xuống sau 3 lần thử.")


def click_more_button_near_image(driver, img_element):
    print("-> Đang tìm nút ... gần ảnh mới...")
    button = driver.execute_script("""
        const img = arguments[0];
        const ir = img.getBoundingClientRect();
        const cx = ir.left + ir.width / 2;
        const cy = ir.top + ir.height / 2;
        function visible(el){
            const r = el.getBoundingClientRect();
            const st = getComputedStyle(el);
            return r.width > 0 && r.height > 0 && st.visibility !== 'hidden' && st.display !== 'none';
        }
        const candidates = Array.from(document.querySelectorAll('button, gem-icon-button'))
            .filter(visible)
            .map(el => {
                const txt = (el.innerText || el.getAttribute('aria-label') || el.getAttribute('arialabel') || '').toLowerCase();
                const html = (el.outerHTML || '').toLowerCase();
                const r = el.getBoundingClientRect();
                const bx = r.left + r.width / 2;
                const by = r.top + r.height / 2;
                const dist = Math.abs(bx - cx) + Math.abs(by - cy);
                const isMore = txt.includes('thêm') || txt.includes('more') || txt.includes('tuỳ chọn') || txt.includes('tùy chọn') || html.includes('more_vert');
                return {el, txt, x:r.left, y:r.top, dist, isMore};
            })
            .filter(o => o.isMore)
            .sort((a,b) => a.dist - b.dist);
        if (!candidates.length) return null;
        return candidates[0].el;
    """, img_element)
    if button is None:
        raise Exception("Không tìm thấy nút dấu ba chấm gần ảnh mới.")
    driver.execute_script("arguments[0].scrollIntoView({block:'center'});", button)
    time.sleep(0.1)
    driver.execute_script("arguments[0].click();", button)
    print("-> Đã mở menu dấu ba chấm gần ảnh mới.")
    time.sleep(0.3)


def click_download_image_menu_item(driver, timeout=10):
    menu_xpath = (
        "//*[@role='menuitem' and (contains(normalize-space(.),'Tải hình ảnh xuống') "
        "or contains(translate(normalize-space(.),'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'download image'))]"
        "|//button[contains(normalize-space(.),'Tải hình ảnh xuống') "
        "or contains(translate(normalize-space(.),'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'download image')]"
    )
    item = WebDriverWait(driver, timeout).until(
        EC.element_to_be_clickable((By.XPATH, menu_xpath))
    )
    driver.execute_script("arguments[0].click();", item)
    print("-> Đã click 'Tải hình ảnh xuống' trong menu Gemini.")


def generate_and_download_image(driver, wait, prompt, final_save_path, row=None, step="IMAGE"):
    # Luồng chính:
    # 1) lưu danh sách blob ảnh đang có
    # 2) gửi prompt
    # 3) chờ nút tải ảnh kích thước đầy đủ trong đúng phản hồi mới
    # 4) nếu lần đầu lỗi, ghi mã lỗi ngắn rồi mở chat Gemini mới và thử lại
    # 5) nếu lần bảo hiểm cũng lỗi, ghi mã lỗi ngắn và dừng bước ảnh
    # 6) lấy thẻ img blob trong cùng phản hồi (không click nút)
    # 7) bỏ lazy-load, ép Edge load/decode cả khi minimize
    # 8) lưu pixel trực tiếp bằng canvas
    #
    # Bảo hiểm này chỉ chạy khi giai đoạn gửi/chờ ảnh gặp lỗi.
    img = None

    for image_attempt in (1, 2):
        print(
            "\n===== DEBUG TRƯỚC KHI GỬI PROMPT"
            + (" (LẦN BẢO HIỂM)" if image_attempt == 2 else "")
            + " ====="
        )
        debug_images_on_page(driver)

        try:
            old_srcs = get_current_gemini_blob_srcs(driver)
            send_prompt_to_gemini(driver, wait, prompt)
            img = wait_new_gemini_completed_image(
                driver,
                old_srcs,
                timeout=GEMINI_IMAGE_WAIT_SECONDS,
            )
            break
        except Exception as image_error:
            raise_if_driver_transport_error(image_error)
            if image_attempt == 1:
                if row is not None:
                    write_retry_note(
                        row,
                        1,
                        step,
                        "GEMINI_IMAGE_ATTEMPT_1_FAILED",
                    )
                print(
                    "⚠️ Lần tạo ảnh đầu tiên thất bại. "
                    "Sẽ thử lại đúng 1 lần bằng chat Gemini mới."
                )
                try:
                    open_gemini_and_wait_ready(driver)
                except Exception:
                    if row is not None:
                        write_retry_note(
                            row,
                            2,
                            step,
                            "GEMINI_RETRY_OPEN_FAILED",
                        )
                    raise
                continue

            if row is not None:
                write_retry_note(
                    row,
                    2,
                    step,
                    "GEMINI_IMAGE_ATTEMPT_2_FAILED",
                )
            raise

    print("\n===== PHẢN HỒI ẢNH GEMINI ĐÃ HOÀN TẤT =====")
    debug_images_on_page(driver)

    try:
        force_gemini_image_load(driver, img, timeout=45)
        return save_image_directly_from_element(
            driver,
            img,
            final_save_path,
        )
    except Exception as direct_error:
        if row is not None:
            write_retry_note(
                row,
                2,
                step,
                "GEMINI_CANVAS_FAILED",
                direct_error,
            )
        raise Exception(
            "Gemini đã báo hoàn tất ảnh nhưng không decode/lưu canvas được: "
            f"{direct_error}"
        ) from direct_error

def images_if_needed(driver, wait, task):
    row = task["row"]
    if not article_ready_for_downstream(task):
        raise Exception("Bài viết chưa được giao cho Worker Word nên không tạo ảnh.")
    if not task["brief1"] or not task["brief2"]:
        raise Exception("Chưa có đủ Brief O/P nên không tạo ảnh.")

    output_folder = get_output_folder(task["word_path"])

    if not file_exists(task["path_img1"]):
        print(f"\n[DÒNG {row}] Bắt đầu tạo ảnh 1 Gemini...")
        write_value(row, COL_STATUS_IMG1, "Đang sinh ảnh 1...", save=False)
        open_gemini_and_wait_ready(driver)
        prompt1 = build_gemini_prompt(task["brief1"])
        write_value(row, COL_STATUS_IMG1, STATUS_SENT_IMG1, save=False)
        save_path1 = get_unique_image_path(output_folder, f"{task['name']} 1")
        final_path1 = generate_and_download_image(driver, wait, prompt1, save_path1, row=row, step="IMG1")
        wb, sh = get_sheet()
        sh.range(f"{COL_PATH_IMG1}{row}").value = final_path1
        sh.range(f"{COL_STATUS_IMG1}{row}").value = STATUS_SAVED_IMG1
        sh.range(f"{COL_GEMINI_URL_IMG1}{row}").value = driver.current_url
        wb.save()
        task = read_task(row)
    else:
        print(f"Dòng {row}: Ảnh 1 đã có. Bỏ qua.")
        write_value(row, COL_STATUS_IMG1, STATUS_SAVED_IMG1, save=False)

    if not file_exists(task["path_img2"]):
        print(f"\n[DÒNG {row}] Bắt đầu tạo ảnh 2 Gemini...")
        write_value(row, COL_STATUS_IMG2, "Đang sinh ảnh 2...", save=False)
        open_gemini_and_wait_ready(driver)
        prompt2 = build_gemini_prompt(task["brief2"])
        write_value(row, COL_STATUS_IMG2, STATUS_SENT_IMG2, save=False)
        save_path2 = get_unique_image_path(output_folder, f"{task['name']} 2")
        final_path2 = generate_and_download_image(driver, wait, prompt2, save_path2, row=row, step="IMG2")
        wb, sh = get_sheet()
        sh.range(f"{COL_PATH_IMG2}{row}").value = final_path2
        sh.range(f"{COL_STATUS_IMG2}{row}").value = STATUS_SAVED_IMG2
        sh.range(f"{COL_GEMINI_URL_IMG2}{row}").value = driver.current_url
        wb.save()
    else:
        print(f"Dòng {row}: Ảnh 2 đã có. Bỏ qua.")
        write_value(row, COL_STATUS_IMG2, STATUS_SAVED_IMG2, save=False)

    if mark_done_if_complete(row):
        print(f"=====> HOÀN THÀNH HOÀN TOÀN DÒNG {row} <=====")
    return read_task(row)

# =====================================================
# MAIN TỔNG
# =====================================================
def should_skip_empty_row(task):
    # Dòng trống hoặc thiếu tối thiểu C thì bỏ qua.
    return not task["name"]




def get_resume_checkpoint(task):
    """
    Xác định bước còn thiếu để chạy tiếp.
    Không tạo lại phần đã hoàn thành.
    """
    if task.get("done") == STATUS_OK and not TEST_OVERWRITE_WORD:
        return "DONE"
    if not is_word_ok(task.get("word_path")):
        return "WORD"
    if not task.get("brief1") or not task.get("brief2"):
        return "BRIEF"
    if not file_exists(task.get("path_img1")):
        return "IMG1"
    if not file_exists(task.get("path_img2")):
        return "IMG2"
    return "DONE"


def process_row(row, driver, wait, progress=None):
    """
    Dùng chung một Edge driver đã được mở từ main().
    Không tự quit driver sau từng dòng nên trình duyệt không còn đóng rồi mở lại.
    """
    task = read_task(row)

    if should_skip_empty_row(task):
        return get_active_driver(driver, wait)

    print("\n" + "=" * 70)
    print(f"KIỂM TRA DÒNG {row}: {task['name']} | CHECKPOINT={get_resume_checkpoint(task)}")

    # 1) Nếu X đã OK thì bỏ qua tuyệt đối
    if task["done"] == STATUS_OK and not TEST_OVERWRITE_WORD:
        print(f"Dòng {row}: Cột X đã OK. Bỏ qua toàn dòng.")
        return get_active_driver(driver, wait)

    # 2) Chỉ coi là hoàn tất khi cả hai file ảnh tồn tại thật.
    image1_ready = file_exists(task["path_img1"])
    image2_ready = file_exists(task["path_img2"])
    if image2_ready and not TEST_OVERWRITE_WORD:
        if image1_ready and image2_ready and is_word_ok(task["word_path"]):
            print(f"Dòng {row}: Đã đủ 2 file ảnh và Word hợp lệ. Ghi X = OK.")
            write_value(row, COL_DONE, STATUS_OK)
            return get_active_driver(driver, wait)
        if word_job_is_pending(row):
            print(f"Dòng {row}: Ảnh đã đủ; đang chờ Worker Word hoàn tất.")
            return get_active_driver(driver, wait)

    # 3) Nếu đã có ảnh 1 thì chỉ tạo tiếp ảnh 2
    if task["path_img1"] and word_ready_or_pending(task) and not TEST_OVERWRITE_WORD:
        print(f"Dòng {row}: Đã có ảnh 1. Bỏ qua Word/Brief/Ảnh 1, tạo tiếp ảnh 2.")
        if not task["brief2"]:
            write_value(row, COL_DONE, "Lỗi: Có ảnh 1 nhưng thiếu Brief 2")
            print(f"❌ Dòng {row}: Có ảnh 1 nhưng thiếu Brief 2.")
            return get_active_driver(driver, wait)
        try:
            if progress:
                progress.update(row, "Đang tạo tiếp ảnh 2 trên Gemini")
            images_if_needed(driver, wait, task)
        except Exception as e:
            raise_if_driver_transport_error(e)
            write_image_final_error(row)
            write_value(row, COL_DONE, f"Lỗi ảnh: {str(e)[:120]}")
            print(f"❌ Dòng {row}: Lỗi bước ảnh: {e}")
        return get_active_driver(driver, wait)

    # 4) Nếu đã có đủ Brief thì chỉ tạo ảnh
    if task["brief1"] and task["brief2"] and word_ready_or_pending(task) and not TEST_OVERWRITE_WORD:
        print(f"Dòng {row}: Đã có Brief. Bỏ qua Word, tạo ảnh.")
        try:
            if progress:
                progress.update(row, "Đang tạo ảnh 1/2 trên Gemini")
            images_if_needed(driver, wait, task)
        except Exception as e:
            raise_if_driver_transport_error(e)
            write_image_final_error(row)
            write_value(row, COL_DONE, f"Lỗi ảnh: {str(e)[:120]}")
            print(f"❌ Dòng {row}: Lỗi bước ảnh: {e}")
        return get_active_driver(driver, wait)

    print(f"BẮT ĐẦU DÒNG {row}: {task['name']}")

    # 1) WORD
    if TEST_OVERWRITE_WORD or not word_ready_or_pending(task):
        try:
            if progress:
                progress.update(row, "Đang viết bài và giao Worker Word")
            task, driver, wait = write_article_if_needed(driver, wait, task)
        except ArticleTooShortError as e:
            _THREAD_CONTEXT.article_too_short_failed = True
            write_retry_note(row, 9, "ARTICLE", "ARTICLE_TOO_SHORT_FINAL", e)
            write_article_error(row, e)
            print(f"❌ Dòng {row}: Bài cuối cùng vẫn dưới {MIN_WORDS} từ: {e}")
            return get_active_driver(driver, wait)
        except Exception as e:
            raise_if_driver_transport_error(e)
            write_retry_note(row, 9, "ARTICLE", "WORD_FINAL_ERROR", e)
            write_article_error(row, e)
            print(f"❌ Dòng {row}: Lỗi bước Word: {e}")
            return get_active_driver(driver, wait)
    elif task["article_status"].upper() != STATUS_OK:
        write_value(row, COL_ARTICLE_STATUS, STATUS_OK)

    task = read_task(row)
    if not article_ready_for_downstream(task):
        write_article_error(row, "Chưa giao được bài cho Worker Word. Dừng dòng này.")
        return get_active_driver(driver, wait)

    # 2) BRIEF
    if not task["brief1"] or not task["brief2"]:
        try:
            if progress:
                progress.update(row, "Đang tạo brief ảnh")
            task, driver, wait = brief_if_needed(driver, wait, task)
        except Exception as e:
            raise_if_driver_transport_error(e)
            write_retry_note(row, 9, "BRIEF", "BRIEF_FINAL_ERROR", e)
            write_value(row, COL_BRIEF_STATUS, f"{STATUS_ERROR_BRIEF}: {str(e)[:120]}")
            print(f"❌ Dòng {row}: Lỗi bước Brief: {e}")
            return get_active_driver(driver, wait)

    # 3) GEMINI ẢNH — dùng luôn cùng cửa sổ Edge
    task = read_task(row)
    if not file_exists(task["path_img1"]) or not file_exists(task["path_img2"]):
        try:
            if progress:
                progress.update(row, "Đang tạo ảnh 1/2 trên Gemini")
            images_if_needed(driver, wait, task)
        except Exception as e:
            raise_if_driver_transport_error(e)
            write_image_final_error(row)
            write_value(row, COL_DONE, f"Lỗi ảnh: {str(e)[:120]}")
            print(f"❌ Dòng {row}: Lỗi bước ảnh: {e}")
            return get_active_driver(driver, wait)
    else:
        if mark_done_if_complete(row):
            print(f"Dòng {row}: Đã đủ Word và ảnh. Ghi X = OK.")
        else:
            print(f"Dòng {row}: Đã đủ ảnh, đang chờ Worker Word.")

    return get_active_driver(driver, wait)


class WriteProgressWindow:
    """Bảng theo dõi; Edge được ẩn bằng cách đưa ra ngoài màn hình."""

    def __init__(self, total_rows):
        self.total_rows = total_rows
        self.stop_after_row = threading.Event()
        self.messages = queue.Queue()
        threading.Thread(target=self._run, daemon=True).start()

    def _run(self):
        import tkinter as tk

        root = tk.Tk()
        root.title("Tiến độ viết bài — Edge đang chạy ẩn")
        root.geometry("590x175+30+80")
        root.resizable(False, False)
        root.attributes("-topmost", True)

        row_label = tk.Label(
            root,
            text="Đang chuẩn bị...",
            font=("Segoe UI", 13, "bold"),
        )
        row_label.pack(pady=(18, 7))
        step_label = tk.Label(
            root,
            text="Đang mở Edge chạy ẩn",
            font=("Segoe UI", 10),
        )
        step_label.pack(pady=(0, 5))
        initial_count = (
            f"Đã hoàn thành 0/{self.total_rows} dòng"
            if self.total_rows is not None
            else "Đã hoàn thành 0 dòng"
        )
        count_label = tk.Label(root, text=initial_count)
        count_label.pack(pady=(0, 12))

        button_frame = tk.Frame(root)
        button_frame.pack()

        def request_stop():
            self.stop_after_row.set()
            stop_button.config(
                text="Sẽ dừng sau dòng này",
                state="disabled",
            )
            step_label.config(text="Đã nhận lệnh dừng an toàn")

        tk.Button(
            button_frame,
            text="Ẩn bảng",
            width=12,
            command=root.withdraw,
        ).pack(side="left", padx=5)
        tk.Button(
            button_frame,
            text="Hiện Edge",
            width=12,
            command=lambda: set_current_edge_visible(True),
        ).pack(side="left", padx=5)
        tk.Button(
            button_frame,
            text="Ẩn Edge",
            width=12,
            command=lambda: set_current_edge_visible(False),
        ).pack(side="left", padx=5)
        stop_button = tk.Button(
            button_frame,
            text="Dừng sau dòng này",
            width=20,
            command=request_stop,
        )
        stop_button.pack(side="left", padx=5)

        # Dấu X chỉ ẩn bảng, không dừng code.
        root.protocol("WM_DELETE_WINDOW", root.withdraw)

        def poll():
            try:
                while True:
                    row, step, completed = self.messages.get_nowait()
                    row_label.config(
                        text=(
                            f"Dòng Excel đang chạy: {row}"
                            if row
                            else "Đã hoàn tất"
                        )
                    )
                    step_label.config(text=step)
                    count_text = (
                        f"Đã hoàn thành {completed}/{self.total_rows} dòng"
                        if self.total_rows is not None
                        else f"Đã hoàn thành {completed} dòng"
                    )
                    count_label.config(text=count_text)
            except queue.Empty:
                pass
            root.after(200, poll)

        poll()
        root.mainloop()

    def update(self, row, step, completed=0):
        self.messages.put((row, step, completed))

    def should_stop(self):
        return self.stop_after_row.is_set()

    def finish(self, completed, stopped=False):
        if stopped:
            message = "Đã dừng an toàn sau khi hoàn thành dòng hiện tại"
        else:
            message = "Đã quét xong toàn bộ danh sách"
        self.update(None, message, completed)
        try:
            winsound.MessageBeep(winsound.MB_ICONASTERISK)
        except Exception:
            pass

def pause_before_next_article(progress, completed_count):
    """Nghỉ có đếm ngược; nút Dừng vẫn phản hồi ngay trong thời gian chờ."""
    pause_seconds = max(0, int(SHORT_ARTICLE_PAUSE_MINUTES * 60))
    pause_end = time.time() + pause_seconds

    print(
        f"\n[TẠM NGHỈ] Có {SHORT_ARTICLE_STREAK_LIMIT} bài Word mới liên tiếp "
        f"dưới {SHORT_ARTICLE_WORD_LIMIT} từ. Nghỉ "
        f"{SHORT_ARTICLE_PAUSE_MINUTES} phút trước bài tiếp theo."
    )

    while True:
        remaining = max(0, int(pause_end - time.time()))
        if remaining <= 0:
            print("[TIẾP TỤC] Đã hết thời gian nghỉ. Bắt đầu bài tiếp theo.")
            return True

        minutes, seconds = divmod(remaining, 60)
        progress.update(
            None,
            f"Tạm nghỉ do 3 bài ngắn: còn {minutes:02d}:{seconds:02d}",
            completed_count,
        )

        if progress.stop_after_row.wait(timeout=min(30, remaining)):
            print("[ĐÃ DỪNG] Nhận lệnh dừng trong thời gian tạm nghỉ.")
            return False


def single_thread_main_legacy():
    print("FILE TỔNG đang chạy. Hãy mở sẵn file Excel dữ liệu trước.")
    start_row, manual_mark_row = get_start_row_by_manual_mark()
    if manual_mark_row:
        print(f"Đã thấy mốc thủ công {MANUAL_MARK_TEXT} ở cột {COL_MANUAL_MARK}, dòng {manual_mark_row}.")
        print(f"Sẽ bắt đầu xử lý từ dòng {start_row}.")
    else:
        print(f"Không thấy mốc thủ công {MANUAL_MARK_TEXT} ở cột {COL_MANUAL_MARK}.")
        print(f"Sẽ bắt đầu xử lý từ dòng {START_ROW}.")

    if start_row > END_ROW:
        print("Mốc thủ công đã nằm ở dòng giới hạn cuối.")
        return

    # Chỉ kiểm tra dòng đầu tiên; không quét trước toàn bộ danh sách.
    if should_skip_empty_row(read_task(start_row)):
        print(f"Dòng {start_row} không có Từ khóa. Không có bài nào để chạy.")
        return

    # Không biết trước tổng số bài vì sẽ dừng ngay tại dòng trống đầu tiên.
    progress = WriteProgressWindow(None)
    completed_count = 0
    successful_rows_since_recycle = 0
    stopped_early = False
    stopped_at_blank = False
    short_article_streak = 0

    driver = None
    try:
        # Chỉ mở Edge đúng một lần cho toàn bộ danh sách.
        driver, wait = create_shared_driver()

        row = start_row
        while row <= END_ROW:
            task = read_task(row)
            if should_skip_empty_row(task):
                stopped_at_blank = True
                print(f"\n[XONG] Dòng {row} không có Từ khóa. Dừng danh sách.")
                break

            # Giữ nguyên nguyên tắc cũ: dòng đã hoàn tất thì bỏ qua.
            if task["done"] == STATUS_OK and not TEST_OVERWRITE_WORD:
                print(f"Dòng {row}: Trạng thái hoàn tất đã OK. Bỏ qua.")
                row += 1
                continue

            # Nếu bấm dừng ngay lúc chương trình còn chuẩn bị Edge thì không bắt đầu dòng đầu tiên.
            if progress.should_stop():
                stopped_early = True
                print("\n[ĐÃ DỪNG] Chưa bắt đầu dòng Excel tiếp theo.")
                break

            progress.update(row, "Đang kiểm tra dữ liệu của dòng", completed_count)
            word_existed_before = is_word_ok(task["word_path"])
            try:
                driver, wait = process_row(row, driver, wait, progress=progress)
                time.sleep(2)
            except Exception as e:
                print(f"❌ Lỗi ngoài dự kiến tại dòng {row}: {e}")
                try:
                    write_retry_note(row, 9, "TOTAL", "TOTAL_ERROR", e)
                    write_value(row, COL_DONE, f"Lỗi tổng: {str(e)[:120]}")
                except Exception:
                    pass
                time.sleep(2)
            finally:
                # Nếu Lớp 2 đã thay Edge rồi một bước sau đó lỗi, vẫn giữ phiên mới
                # cho dòng tiếp theo thay vì quay lại session cũ đã quit().
                driver, wait = get_active_driver(driver, wait)

            completed_count += 1
            progress.update(row, "Đã xử lý xong và lưu dòng", completed_count)

            # Chỉ tính bài Word vừa được tạo thành công trong phiên chạy này.
            # Các dòng bỏ qua hoặc chỉ xử lý brief/ảnh không ảnh hưởng bộ đếm.
            task_after = read_task(row)
            word_created_now = (
                (TEST_OVERWRITE_WORD or not word_existed_before)
                and is_word_ok(task_after["word_path"])
            )
            if word_created_now:
                saved_word_count = cell(row, COL_WORD_COUNT).value
                try:
                    saved_word_count = int(float(saved_word_count))
                except (TypeError, ValueError):
                    saved_word_count = None

                if saved_word_count is not None:
                    if saved_word_count < SHORT_ARTICLE_WORD_LIMIT:
                        short_article_streak += 1
                        print(
                            f"[BÀI NGẮN] Dòng {row}: {saved_word_count} từ. "
                            f"Chuỗi hiện tại: {short_article_streak}/"
                            f"{SHORT_ARTICLE_STREAK_LIMIT}."
                        )
                    else:
                        if short_article_streak:
                            print(
                                f"[ĐẠT] Dòng {row}: {saved_word_count} từ. "
                                "Đặt lại bộ đếm bài ngắn."
                            )
                        short_article_streak = 0

            # Chỉ dừng tại đây: toàn bộ xử lý và lưu Excel của dòng hiện tại đã kết thúc.
            if progress.should_stop():
                stopped_early = True
                print(f"\n[ĐÃ DỪNG] Hoàn thành dòng {row}, không chạy dòng kế tiếp.")
                break

            if short_article_streak >= SHORT_ARTICLE_STREAK_LIMIT:
                if not pause_before_next_article(progress, completed_count):
                    stopped_early = True
                    break
                short_article_streak = 0

            # Bộ đếm thuộc riêng tiến trình/worker hiện tại và chỉ nằm trong RAM.
            # Chỉ cộng khi chính dòng worker vừa xử lý chuyển từ chưa OK sang OK.
            if (
                task["done"] != STATUS_OK
                and task_after["done"] == STATUS_OK
            ):
                successful_rows_since_recycle += 1
                if successful_rows_since_recycle >= RECYCLE_EVERY_N_ROWS:
                    print(
                        f"\n[WORKER RECYCLING] Đã hoàn thành "
                        f"{successful_rows_since_recycle} bài. Khởi động lại Edge để xả RAM."
                    )
                    if driver:
                        try:
                            driver.quit()
                        except Exception:
                            pass
                        driver = None
                    time.sleep(3)
                    driver, wait = create_shared_driver()
                    successful_rows_since_recycle = 0

            row += 1

        progress.finish(completed_count, stopped=stopped_early)
        if stopped_early:
            print(f"\n[DỪNG AN TOÀN] Đã xử lý {completed_count} dòng.")
        elif stopped_at_blank:
            print(f"\n[XONG] Đã xử lý {completed_count} dòng rồi gặp dòng trống.")
        else:
            print(f"\n[XONG] Đã xử lý đến dòng giới hạn {END_ROW}.")

    finally:
        # Chỉ đóng Edge một lần sau khi toàn bộ danh sách đã chạy xong.
        if driver:
            try:
                driver.quit()
            except Exception:
                pass


@dataclass
class WorkerProgress:
    worker_id: int

    def update(self, row, step, completed=0):
        wait_until_system_resumed(self.worker_id, row)
        name = _WORKER_CURRENT_NAMES.get(self.worker_id, "")
        row_text = f"Dòng Excel {row}" if row else "Không có dòng"
        name_text = f" — {name}" if name else ""
        full_step = f"{row_text}{name_text} — {step}"
        message = f"[WORKER {self.worker_id}] {full_step}"
        print(message)
        UI_QUEUE.put(("STATUS", self.worker_id, row, full_step, completed))

    def should_stop(self):
        return STOP_EVENT.is_set()


def wait_until_system_resumed(worker_id, row=None):
    """Chặn Worker tại điểm an toàn, giữ nguyên Edge/session để người dùng đăng nhập."""
    announced = False
    while not RUN_EVENT.is_set() and not STOP_EVENT.is_set():
        if not announced:
            name = _WORKER_CURRENT_NAMES.get(worker_id, "")
            detail = f" — {name}" if name else ""
            UI_QUEUE.put((
                "PAUSE", worker_id, row,
                f"Dòng Excel {row or '-'}{detail} — ĐÃ TẠM DỪNG, chờ đăng nhập",
                0,
            ))
            announced = True
        RUN_EVENT.wait(timeout=0.5)


def pause_all_workers():
    """Yêu cầu tất cả Worker dừng tại checkpoint gần nhất và hiện Edge để đăng nhập."""
    RUN_EVENT.clear()
    set_all_worker_edges_visible(True)


def resume_all_workers():
    """Cho tất cả Worker tiếp tục bằng chính driver/session đang mở."""
    RUN_EVENT.set()


class MultiWorkerMonitor:
    """Bảng trạng thái từng Worker; UI chỉ nhận message, không chạm Excel."""
    def __init__(self, worker_ids):
        self.worker_ids = list(worker_ids)
        self.worker_count = len(self.worker_ids)
        self.thread = threading.Thread(target=self._run, daemon=True)
        self.thread.start()

    def _run(self):
        import tkinter as tk
        root = tk.Tk()
        worker_text = ", ".join(map(str, self.worker_ids))
        root.title(f"{VERSION} - WORKER {worker_text}")
        root.geometry("980x380+30+60")
        root.attributes("-topmost", True)
        paused_workers = set()

        overall_status = tk.Label(
            root,
            text="● ĐANG CHẠY",
            fg="#137333",
            bg="#d9f2df",
            font=("Segoe UI", 11, "bold"),
            pady=7,
        )
        overall_status.pack(fill="x", padx=8, pady=(8, 4))
        labels = {}
        for worker_id in self.worker_ids:
            frame = tk.Frame(root, padx=8, pady=6, relief="groove", borderwidth=1)
            frame.pack(fill="x", padx=8, pady=3)
            name = tk.Label(frame, text=f"WORKER {worker_id}", width=12, anchor="w")
            name.pack(side="left")
            status = tk.Label(frame, text="Đang chờ", anchor="w", fg="#1464a0")
            status.pack(side="left", fill="x", expand=True)
            tk.Button(
                frame,
                text="Xem Edge",
                command=lambda wid=worker_id: set_worker_edge_visible(wid, True),
            ).pack(side="right", padx=3)
            tk.Button(
                frame,
                text="Ẩn Edge",
                command=lambda wid=worker_id: set_worker_edge_visible(wid, False),
            ).pack(side="right", padx=3)
            skip = tk.Button(
                frame,
                text="Bỏ qua nghỉ",
                command=lambda wid=worker_id: SKIP_PAUSE_EVENTS[wid].set(),
            )
            skip.pack(side="right", padx=3)
            labels[worker_id] = status

        word_frame = tk.Frame(root, padx=8, pady=6, relief="groove", borderwidth=1)
        word_frame.pack(fill="x", padx=8, pady=3)
        tk.Label(word_frame, text="WORD + VBA", width=12, anchor="w").pack(side="left")
        word_status_label = tk.Label(
            word_frame,
            text="Đang chờ bài | Hàng chờ: 0",
            anchor="w",
            fg="#7d3c98",
        )
        word_status_label.pack(side="left", fill="x", expand=True)

        excel_frame = tk.Frame(root, padx=8, pady=6, relief="groove", borderwidth=1)
        excel_frame.pack(fill="x", padx=8, pady=3)
        tk.Label(excel_frame, text="EXCEL WRITER", width=12, anchor="w").pack(side="left")
        excel_status_label = tk.Label(
            excel_frame,
            text="Đang khởi động | Hàng chờ: 0",
            anchor="w",
            fg="#137333",
        )
        excel_status_label.pack(side="left", fill="x", expand=True)

        def stop_soft():
            # STOP_EVENT chỉ dành cho dừng khẩn hoặc lỗi thật. Nếu dùng nó ở
            # đây, Excel Writer sẽ hủy lệnh ghi đang xử lý và báo lỗi giả.
            SOFT_STOP_EVENT.set()
            # Nhả trạng thái tạm dừng (nếu có) để Worker thấy yêu cầu dừng mềm
            # và thoát mà không lấy thêm dòng mới.
            RUN_EVENT.set()
            overall_status.config(
                text="\u25cf \u0110ANG D\u1eeaNG AN TO\u00c0N \u2014 \u0111ang ho\u00e0n t\u1ea5t h\u00e0ng ch\u1edd Word/Excel",
                fg="#9a5b00",
                bg="#fff0cc",
            )
            for label in labels.values():
                label.config(text="Đã yêu cầu dừng sau dòng hiện tại", fg="#c0392b")

        controls = tk.Frame(root)
        controls.pack(pady=8)
        pause_button = tk.Button(
            controls,
            text="Tạm dừng tất cả",
            bg="#f39c12",
            activebackground="#e67e22",
        )
        pause_button.pack(side="left", padx=5)
        resume_button = tk.Button(
            controls,
            text="Chạy tiếp",
            bg="#d5d8dc",
            state="disabled",
        )
        resume_button.pack(side="left", padx=5)

        def request_pause_all():
            paused_workers.clear()
            pause_all_workers()
            overall_status.config(
                text="● ĐANG YÊU CẦU TẠM DỪNG — chờ Worker tới điểm an toàn",
                fg="#9a5b00",
                bg="#fff0cc",
            )
            pause_button.config(
                text="Đang tạm dừng...",
                bg="#e74c3c",
                activebackground="#c0392b",
                state="disabled",
                relief="sunken",
            )
            resume_button.config(
                text="Chạy tiếp",
                bg="#7bd389",
                activebackground="#58bd6b",
                state="normal",
                relief="raised",
            )

        def request_resume_all():
            resume_all_workers()
            paused_workers.clear()
            overall_status.config(
                text="● ĐANG CHẠY — đã phát lệnh tiếp tục",
                fg="#137333",
                bg="#d9f2df",
            )
            pause_button.config(
                text="Tạm dừng tất cả",
                bg="#f39c12",
                activebackground="#e67e22",
                state="normal",
                relief="raised",
            )
            resume_button.config(
                text="Đã chạy tiếp",
                bg="#d5d8dc",
                state="disabled",
                relief="raised",
            )
            for label in labels.values():
                if "TẠM DỪNG" in str(label.cget("text")).upper():
                    label.config(text="Đã nhận lệnh chạy tiếp...", fg="#137333")

        pause_button.config(command=request_pause_all)
        resume_button.config(command=request_resume_all)
        tk.Button(
            controls,
            text="Xem tất cả Edge",
            command=lambda: set_all_worker_edges_visible(True),
        ).pack(side="left", padx=5)
        tk.Button(
            controls,
            text="Ẩn tất cả Edge",
            command=lambda: set_all_worker_edges_visible(False),
        ).pack(side="left", padx=5)
        tk.Button(controls, text="Dừng an toàn", command=stop_soft).pack(side="left", padx=5)

        def poll():
            word_status_label.config(
                text=f"{_WORD_CURRENT_STATUS} | Hàng chờ: {WORD_QUEUE.qsize()}"
            )
            excel_color = "#c0392b" if _EXCEL_WRITER_FAILED.is_set() else "#137333"
            excel_alive = "Sống" if _EXCEL_WRITER_ALIVE.is_set() else "Đã dừng"
            excel_status_label.config(
                text=(
                    f"{excel_alive} | {_EXCEL_WRITER_STATUS} | "
                    f"Hàng chờ: {RESULT_QUEUE.qsize()}"
                ),
                fg=excel_color,
            )
            try:
                while True:
                    kind, worker_id, row, step, _completed = UI_QUEUE.get_nowait()
                    color = "#d06b00" if kind == "PAUSE" else "#1464a0"
                    labels[worker_id].config(text=step, fg=color)
                    if kind == "PAUSE" and not RUN_EVENT.is_set():
                        paused_workers.add(worker_id)
                        if len(paused_workers) >= self.worker_count:
                            overall_status.config(
                                text="■ ĐÃ TẠM DỪNG TẤT CẢ — có thể đăng nhập/chỉnh tài khoản",
                                fg="#a61b1b",
                                bg="#ffd8d8",
                            )
                            pause_button.config(
                                text="Đã tạm dừng tất cả",
                                bg="#c0392b",
                            )
                    elif RUN_EVENT.is_set():
                        paused_workers.discard(worker_id)
            except queue.Empty:
                pass
            root.after(200, poll)

        poll()
        root.mainloop()


def _queue_text(value):
    return " ".join(str(value or "").split()).casefold()


def _load_write_plan():
    raw = os.environ.get("HOTKEYVIP_WRITE_PLAN", "").strip()
    if not raw:
        return None
    try:
        plan = json.loads(raw)
    except (TypeError, ValueError) as exc:
        print(f"⚠️ Kế hoạch ưu tiên không hợp lệ; chạy theo cơ chế cũ: {exc}")
        return None
    if not isinstance(plan, dict):
        print("⚠️ Kế hoạch ưu tiên không phải object; chạy theo cơ chế cũ.")
        return None
    return plan


def _row_is_open(data):
    return bool(str(data.get(COL_NAME) or "").strip()) and (
        str(data.get(COL_DONE) or "").strip().upper() != STATUS_OK
    )


def _row_has_retryable_error(data):
    if not _row_is_open(data):
        return False
    article_status = _queue_text(data.get(COL_ARTICLE_STATUS))
    brief_status = _queue_text(data.get(COL_BRIEF_STATUS))
    done_status = _queue_text(data.get(COL_DONE))
    return bool(
        str(data.get(COL_ARTICLE_ERROR) or "").strip()
        or str(data.get(COL_RETRY_ERROR) or "").strip()
        or article_status in {"error", "word_error"}
        or "lỗi" in brief_status
        or "error" in brief_status
        or "lỗi" in done_status
        or "error" in done_status
        or "tạm bỏ lượt" in done_status
    )


def _build_task_rows(rows, plan):
    """Xếp lỗi -> domain ưu tiên -> luồng thường, không di chuyển dòng Excel."""
    manual_rows = [
        row for row, data in rows.items()
        if str(data.get(COL_MANUAL_MARK) or "").strip() == MANUAL_MARK_TEXT
    ]
    start_row = (max(manual_rows) + 1) if manual_rows else START_ROW

    # Không có kế hoạch từ app: giữ nguyên tuyệt đối hành vi cũ.
    if plan is None:
        tasks = [
            row for row in range(start_row, END_ROW + 1)
            if _row_is_open(rows[row])
        ]
        return tasks, {
            "start_row": start_row,
            "error_count": 0,
            "priority_count": 0,
            "normal_count": len(tasks),
            "legacy": True,
        }

    selected = set()
    tasks = []

    def add_rows(candidates):
        added = 0
        for row in candidates:
            if row in selected:
                continue
            selected.add(row)
            tasks.append(row)
            added += 1
        return added

    error_count = 0
    if bool(plan.get("retry_errors_first", True)):
        error_count = add_rows(
            row for row in range(START_ROW, END_ROW + 1)
            if _row_has_retryable_error(rows[row])
        )

    priority_domain = _queue_text(plan.get("priority_domain"))
    try:
        requested_count = max(0, int(plan.get("priority_count", 0)))
    except (TypeError, ValueError):
        requested_count = 0
    priority_candidates = []
    if priority_domain and requested_count:
        for row in range(START_ROW, END_ROW + 1):
            data = rows[row]
            if row in selected or not _row_is_open(data):
                continue
            if _queue_text(data.get(COL_WEB)) == priority_domain:
                priority_candidates.append(row)
                if len(priority_candidates) >= requested_count:
                    break
    priority_count = add_rows(priority_candidates)

    normal_count = 0
    if bool(plan.get("continue_normal", True)):
        normal_count = add_rows(
            row for row in range(start_row, END_ROW + 1)
            if _row_is_open(rows[row])
        )

    return tasks, {
        "start_row": start_row,
        "error_count": error_count,
        "priority_count": priority_count,
        "priority_requested": requested_count,
        "priority_domain": str(plan.get("priority_domain") or "").strip(),
        "normal_count": normal_count,
        "legacy": False,
    }


def load_excel_tasks_once():
    """Đọc B2:AE10000 một lần, rồi tạo hàng chờ an toàn trong RAM."""
    _wb, sh = get_real_sheet()
    values = sh.range(f"B{START_ROW}:AE{END_ROW}").value
    columns = [column_letter(number) for number in range(2, 32)]
    rows = {}
    for offset, raw_values in enumerate(values):
        row = START_ROW + offset
        rows[row] = dict(zip(columns, raw_values))

    tasks, summary = _build_task_rows(rows, _load_write_plan())
    if summary["legacy"]:
        print(
            f"[HÀNG CHỜ] Chế độ cũ: {len(tasks)} bài từ dòng {summary['start_row']} sau OK OK."
        )
    else:
        domain = summary.get("priority_domain") or "không chọn"
        print(
            f"[HÀNG CHỜ] Lỗi: {summary['error_count']} | "
            f"Ưu tiên {domain}: {summary['priority_count']}/{summary['priority_requested']} | "
            f"Bình thường sau OK OK (dòng {summary['start_row']}): {summary['normal_count']} | "
            f"Tổng không trùng: {len(tasks)}"
        )
    return rows, tasks


def _excel_retry_delay(attempt):
    """Chờ tăng dần nhưng không quá lâu để Excel kịp thoát Copy/Paste/Edit mode."""
    return min(EXCEL_RETRY_MAX_DELAY_SECONDS, 0.5 * (2 ** min(attempt - 1, 4)))


def _assert_selected_workbook(wb):
    """Refuse to write after Excel has redirected the book through Save As."""
    selected_path = os.environ.get("HOTKEYVIP_SELECTED_EXCEL", "").strip()
    if not selected_path:
        return
    expected = os.path.normcase(os.path.abspath(selected_path))
    try:
        actual = os.path.normcase(os.path.abspath(str(wb.fullname)))
    except Exception as exc:
        raise ExcelWriterUnavailableError(
            f"Không đọc được đường dẫn workbook đang ghi: {exc}"
        ) from exc
    if actual != expected:
        raise WorkbookIdentityChangedError(
            "Excel đã chuyển sang workbook khác. "
            f"Cần: {selected_path}; đang bám: {wb.fullname}. "
            "Dừng để không ghi dữ liệu vào file Save As tên ngẫu nhiên."
        )


def _run_excel_with_retry(label, operation):
    """Giữ nguyên lệnh và thử đến khi Excel nhận; không làm chết Excel Writer."""
    global _EXCEL_WRITER_STATUS, _EXCEL_WRITER_RETRY_COUNT
    global _EXCEL_WRITER_LAST_ERROR

    attempt = 0
    started_at = time.monotonic()
    paused_by_writer = False
    while True:
        if STOP_EVENT.is_set():
            raise ExcelWriterUnavailableError(
                f"Đã nhận lệnh dừng trong khi chờ Excel thực hiện: {label}."
            )
        try:
            result = operation()
            if attempt:
                print(f"✅ [EXCEL WRITER] {label} đã phục hồi sau {attempt} lần thử lại.")
            _EXCEL_WRITER_RETRY_COUNT = 0
            _EXCEL_WRITER_LAST_ERROR = ""
            _EXCEL_WRITER_FAILED.clear()
            _EXCEL_WRITER_STATUS = "Đang ghi bình thường"
            if paused_by_writer and not STOP_EVENT.is_set():
                RUN_EVENT.set()
                print("▶ [EXCEL WRITER] Excel đã hoạt động lại; tiếp tục các Worker.")
            return result
        except WorkbookIdentityChangedError:
            raise
        except Exception as exc:
            attempt += 1
            _EXCEL_WRITER_RETRY_COUNT = attempt
            _EXCEL_WRITER_LAST_ERROR = str(exc)[:240]
            elapsed = time.monotonic() - started_at
            if elapsed >= EXCEL_RETRY_TIMEOUT_SECONDS:
                _EXCEL_WRITER_FAILED.set()
                STOP_EVENT.set()
                # Wake workers/Word that may currently be paused on RUN_EVENT.
                RUN_EVENT.set()
                raise ExcelWriterUnavailableError(
                    f"Excel không nhận lệnh '{label}' sau {elapsed:.1f}s "
                    f"({attempt} lần thử). Có thể Excel đang bị hộp thoại Save As chặn."
                ) from exc
            delay = _excel_retry_delay(attempt)
            _EXCEL_WRITER_STATUS = (
                f"Excel bận; thử lại lần {attempt} sau {delay:.1f}s"
            )
            if attempt == 1 or attempt % EXCEL_RETRY_WARNING_AFTER == 0:
                print(
                    f"⚠️ [EXCEL WRITER] {label} bị Excel từ chối "
                    f"(lần {attempt}): {exc}. Sẽ thử lại sau {delay:.1f}s."
                )
            if attempt >= EXCEL_RETRY_PAUSE_AFTER and RUN_EVENT.is_set():
                RUN_EVENT.clear()
                paused_by_writer = True
                _EXCEL_WRITER_FAILED.set()
                print(
                    "⏸ [EXCEL WRITER] Excel bận kéo dài; đã tạm dừng Worker. "
                    "Hãy bấm Esc/đóng hộp thoại trong Excel. Lệnh ghi không bị mất."
                )
            time.sleep(delay)


def verify_row_identity(row):
    """Nhờ Excel Writer xác nhận domain + từ khóa vẫn ở đúng dòng trước khi chạy."""
    if _EXCEL_WRITER_FAILED.is_set():
        raise ExcelWriterUnavailableError(
            "Excel Writer đã dừng; không giao thêm dòng mới."
        )
    data = _THREAD_CONTEXT.rows[row]
    result = {}
    completed = threading.Event()
    RESULT_QUEUE.put(
        (
            "VERIFY",
            row,
            str(data.get(COL_WEB) or "").strip(),
            str(data.get(COL_NAME) or "").strip(),
            result,
            completed,
        )
    )
    if not completed.wait(timeout=30):
        raise RowIdentityChangedError(
            f"Không xác minh được dòng {row} trong 30 giây; đã dừng để tránh ghi nhầm."
        )
    if result.get("error"):
        raise RowIdentityChangedError(result["error"])


def _drain_result_queue_after_writer_failure(error):
    """Release all queue waiters after the single Excel writer has stopped."""
    while True:
        try:
            command = RESULT_QUEUE.get_nowait()
        except queue.Empty:
            break
        try:
            if command and command[0] == "VERIFY":
                _, _row, _web, _name, result, completed = command
                result["error"] = str(error)
                completed.set()
        finally:
            RESULT_QUEUE.task_done()


def excel_writer_thread():
    """Luồng duy nhất ghi Excel; mọi thao tác COM đều có retry chống OLE busy."""
    global _EXCEL_WRITER_STATUS
    pythoncom = None
    dirty = False
    last_save = time.time()
    _EXCEL_WRITER_ALIVE.set()
    _EXCEL_WRITER_STATUS = "Đang kết nối Excel"
    try:
        try:
            import pythoncom
            pythoncom.CoInitialize()
        except ImportError:
            pythoncom = None

        wb, sh = _run_excel_with_retry("kết nối workbook", get_real_sheet)
        _assert_selected_workbook(wb)

        def save_selected_workbook():
            _assert_selected_workbook(wb)
            wb.save()

        _EXCEL_WRITER_STATUS = "Đang chờ lệnh"
        _EXCEL_WRITER_READY.set()
        while True:
            command = RESULT_QUEUE.get()
            try:
                if command[0] == "STOP":
                    if dirty:
                        _run_excel_with_retry(
                            "lưu Excel khi kết thúc", save_selected_workbook
                        )
                    _EXCEL_WRITER_STATUS = "Đã dừng an toàn"
                    break

                if command[0] == "VERIFY":
                    _, row, expected_web, expected_name, result, completed = command
                    try:
                        actual_web = str(sh.range(f"{COL_WEB}{row}").value or "").strip()
                        actual_name = str(sh.range(f"{COL_NAME}{row}").value or "").strip()
                        if (
                            _queue_text(actual_web) != _queue_text(expected_web)
                            or _queue_text(actual_name) != _queue_text(expected_name)
                        ):
                            result["error"] = (
                                f"Dòng {row} đã thay đổi sau khi nạp RAM. "
                                f"Mong đợi [{expected_web} | {expected_name}], "
                                f"thực tế [{actual_web} | {actual_name}]. Không ghi dữ liệu."
                            )
                    finally:
                        completed.set()
                    continue

                if command[0] == "WRITE":
                    _, row, col, value = command

                    def write_cell():
                        _assert_selected_workbook(wb)
                        target = sh.range(f"{col}{row}")
                        target.value = value
                        if col in {
                            COL_ARTICLE_STATUS, COL_ARTICLE_ERROR, COL_BRIEF_STATUS,
                            COL_DONE, COL_RETRY_COUNT, COL_RETRY_STEP,
                            COL_RETRY_ERROR, COL_RETRY_TIME,
                        }:
                            target.api.WrapText = False

                    _run_excel_with_retry(f"ghi {col}{row}", write_cell)
                    dirty = True

                if command[0] == "SAVE" or (
                    dirty and time.time() - last_save >= WRITER_SAVE_INTERVAL_SECONDS
                ):
                    _run_excel_with_retry("lưu workbook", save_selected_workbook)
                    dirty = False
                    last_save = time.time()
            finally:
                # Chỉ tới đây khi lệnh đã thành công; retry giữ lệnh hiện tại tại chỗ.
                RESULT_QUEUE.task_done()
    except BaseException as exc:
        _EXCEL_WRITER_FAILED.set()
        _EXCEL_WRITER_STATUS = f"LỖI NGHIÊM TRỌNG: {str(exc)[:180]}"
        STOP_EVENT.set()
        RUN_EVENT.set()
        _drain_result_queue_after_writer_failure(exc)
        print(f"❌ [EXCEL WRITER] Lỗi nghiêm trọng; đã dừng Worker: {exc}")
        return
    finally:
        _EXCEL_WRITER_READY.clear()
        _EXCEL_WRITER_ALIVE.clear()
        try:
            if pythoncom:
                pythoncom.CoUninitialize()
        except Exception:
            pass


def word_worker_thread(rows):
    """Luồng duy nhất sở hữu Microsoft Word/VBA và xử lý tuần tự mọi snapshot."""
    global _WORD_CURRENT_ROW, _WORD_CURRENT_STATUS
    import pythoncom

    pythoncom.CoInitialize()
    _THREAD_CONTEXT.is_worker = True
    _THREAD_CONTEXT.worker_id = "WORD"
    _THREAD_CONTEXT.rows = rows
    try:
        while True:
            job = WORD_QUEUE.get()
            try:
                if job is None:
                    break

                if STOP_EVENT.is_set():
                    _WORD_CURRENT_STATUS = "Đã bỏ hàng chờ vì Excel Writer dừng"
                    continue

                # Tạm dừng chỉ có hiệu lực tại ranh giới an toàn giữa hai file Word.
                while not RUN_EVENT.wait(timeout=0.5):
                    _WORD_CURRENT_STATUS = "Đã tạm dừng tại điểm an toàn"

                if STOP_EVENT.is_set():
                    _WORD_CURRENT_STATUS = "Đã bỏ hàng chờ vì Excel Writer dừng"
                    continue

                row = job["row"]
                _WORD_CURRENT_ROW = row
                _WORD_CURRENT_STATUS = f"Đang xử lý dòng {row}"
                print(
                    f"\n[WORD] Bắt đầu dòng {row}: {job['name']} | "
                    f"còn chờ {WORD_QUEUE.qsize()} bài"
                )
                errors = []
                saved = False
                for attempt in (1, 2):
                    try:
                        copy_and_save_snapshot(job["article"], job["word_path"])
                        saved = True
                        break
                    except Exception as exc:
                        errors.append(f"Lần {attempt}: {exc}")
                        print(f"⚠️ [WORD] Dòng {row} lỗi lần {attempt}: {exc}")
                        time.sleep(1)

                if saved:
                    write_article_success(
                        row, job["word_path"], job["chat_url"], job["word_count"]
                    )
                    print(f"✅ [WORD] Dòng {row}: Đã lưu Word + chạy VBA thành công.")
                    _WORD_CURRENT_STATUS = f"Đã xong dòng {row}"
                else:
                    error_text = " | ".join(errors)
                    write_word_worker_error(row, error_text)
                    write_retry_note(row, 9, "WORD_WORKER", "WORD_QUEUE_ERROR", error_text)
                    print(f"❌ [WORD] Dòng {row}: Thất bại sau 2 lần, chuyển bài kế tiếp.")
                    _WORD_CURRENT_STATUS = f"Lỗi dòng {row}; đã chuyển tiếp"
            except Exception as exc:
                row = job.get("row") if isinstance(job, dict) else None
                if row is not None:
                    write_word_worker_error(row, exc)
                print(f"❌ [WORD] Lỗi ngoài dự kiến: {exc}")
            finally:
                if isinstance(job, dict):
                    with _WORD_PENDING_LOCK:
                        _WORD_PENDING_ROWS.discard(job["row"])
                    if is_word_ok(job.get("word_path")):
                        mark_done_if_complete(job["row"])
                WORD_QUEUE.task_done()
    finally:
        _WORD_CURRENT_ROW = None
        _WORD_CURRENT_STATUS = "Đã dừng"
        pythoncom.CoUninitialize()


def wait_worker_pause(worker_id, progress, row):
    skip_event = SKIP_PAUSE_EVENTS[worker_id]
    skip_event.clear()
    remaining = int(SHORT_ARTICLE_PAUSE_MINUTES * 60)
    print(
        f"⏸ [WORKER {worker_id}] Đủ {SHORT_ARTICLE_STREAK_LIMIT} lần liên tiếp: "
        f"bài thành công dưới {SHORT_ARTICLE_WORD_LIMIT} từ hoặc bài lỗi dưới "
        f"{MIN_WORDS} từ. Nghỉ {SHORT_ARTICLE_PAUSE_MINUTES} phút."
    )
    while (
        remaining > 0
        and not STOP_EVENT.is_set()
        and not SOFT_STOP_EVENT.is_set()
        and not skip_event.is_set()
    ):
        minutes, seconds = divmod(remaining, 60)
        UI_QUEUE.put((
            "PAUSE", worker_id, row,
            f"Đang nghỉ do 2 bài ngắn/lỗi <700: {minutes:02d}:{seconds:02d}",
            0,
        ))
        skip_event.wait(timeout=1)
        remaining -= 1
    progress.update(row, "Tiếp tục làm việc")


def worker_loop(worker_id, rows):
    # Khởi tạo COM theo thread cho các API Windows phụ trợ. Word/VBA chỉ do
    # word_worker_thread sở hữu, Worker Edge không trực tiếp lưu Word.
    import pythoncom
    pythoncom.CoInitialize()
    _THREAD_CONTEXT.is_worker = True
    _THREAD_CONTEXT.worker_id = worker_id
    _THREAD_CONTEXT.rows = rows
    _THREAD_CONTEXT.stop_worker_after_row = False
    progress = WorkerProgress(worker_id)
    driver = wait = None
    short_streak = 0
    completed = 0
    recycled = 0
    try:
        driver, wait = create_shared_driver()
        while not STOP_EVENT.is_set():
            wait_until_system_resumed(worker_id)
            # Dừng mềm chỉ có hiệu lực trước khi nhận task mới. Task hiện tại
            # vẫn chạy xong để các checkpoint Word/Excel được ghi đầy đủ.
            if STOP_EVENT.is_set() or SOFT_STOP_EVENT.is_set():
                break
            try:
                row = TASK_QUEUE.get_nowait()
            except queue.Empty:
                break
            try:
                current_task = read_task(row)
                verify_row_identity(row)
                _WORKER_CURRENT_NAMES[worker_id] = current_task.get("name", "")
                before_word = is_word_ok(current_task["word_path"])
                _THREAD_CONTEXT.article_too_short_failed = False
                progress.update(row, "Bắt đầu xử lý", completed)
                try:
                    driver, wait = process_row(row, driver, wait, progress=progress)
                except DriverTransportError as first_driver_error:
                    progress.update(
                        row,
                        "EdgeDriver mất phản hồi; đang mở lại đúng Worker và thử dòng này lần cuối",
                        completed,
                    )
                    print(
                        f"⚠️ [WORKER {worker_id}] Dòng {row}: EdgeDriver mất phản hồi: "
                        f"{first_driver_error}"
                    )
                    write_retry_note(
                        row,
                        1,
                        "EDGE_DRIVER",
                        "DRIVER_TRANSPORT_RESTART",
                        first_driver_error,
                    )
                    try:
                        driver.quit()
                    except Exception:
                        pass
                    driver = wait = None
                    time.sleep(1)
                    driver, wait = create_shared_driver()
                    _THREAD_CONTEXT.article_too_short_failed = False
                    print(
                        f"-> [WORKER {worker_id}] Đã mở Edge mới bằng đúng profile; "
                        f"thử lại dòng {row} đúng 1 lần."
                    )
                    try:
                        driver, wait = process_row(
                            row, driver, wait, progress=progress
                        )
                    except DriverTransportError as second_driver_error:
                        write_retry_note(
                            row,
                            2,
                            "EDGE_DRIVER",
                            "DRIVER_TRANSPORT_FINAL",
                            second_driver_error,
                        )
                        raise RuntimeError(
                            "EdgeDriver vẫn mất phản hồi sau khi đã khởi động lại "
                            f"Worker {worker_id}: {second_driver_error}"
                        ) from second_driver_error
                after = read_task(row)
                if _THREAD_CONTEXT.article_too_short_failed:
                    short_streak += 1
                    print(
                        f"⚠️ [WORKER {worker_id}] Dòng {row} thất bại vì dưới "
                        f"{MIN_WORDS} từ; bộ đếm nghỉ = {short_streak}/"
                        f"{SHORT_ARTICLE_STREAK_LIMIT}."
                    )
                elif not before_word and (
                    is_word_ok(after["word_path"])
                    or str(after["article_status"]).upper() == STATUS_WORD_QUEUED
                ):
                    try:
                        word_count = int(float(_THREAD_CONTEXT.rows[row].get(COL_WORD_COUNT)))
                    except (TypeError, ValueError):
                        word_count = None
                    if word_count is not None:
                        short_streak = short_streak + 1 if word_count < SHORT_ARTICLE_WORD_LIMIT else 0
                completed += 1
                recycled += 1
                if short_streak >= SHORT_ARTICLE_STREAK_LIMIT:
                    wait_worker_pause(worker_id, progress, row)
                    short_streak = 0
                if recycled >= RECYCLE_EVERY_N_ROWS:
                    progress.update(row, "Khởi động lại Edge để xả RAM", completed)
                    try:
                        driver.quit()
                    except Exception:
                        pass
                    driver, wait = create_shared_driver()
                    recycled = 0
                progress.update(row, "Đã xong phần Edge; hệ thống tự chờ Word", completed)
            except RowIdentityChangedError as exc:
                RUN_EVENT.clear()
                STOP_EVENT.set()
                progress.update(row, "Dừng an toàn: dòng Excel đã thay đổi", completed)
                print(f"❌ [BẢO HIỂM DÒNG] {exc}")
            except ChatGPTSendUnconfirmedError as exc:
                # Không gửi lại trong cùng lượt, nhưng cũng không được làm chết Worker.
                # COL_DONE khác OK nên load_excel_tasks_once() sẽ tự đưa dòng này vào
                # hàng đợi khi chạy lại; checkpoint Word/Brief/Img vẫn được giữ nguyên.
                write_retry_note(row, 9, "SKIPPED", "CHATGPT_SEND_UNCONFIRMED", exc)
                write_value(row, COL_DONE, f"Tạm bỏ lượt: ChatGPT chưa xác nhận gửi ({str(exc)[:100]})")
                _THREAD_CONTEXT.stop_worker_after_row = False
                driver, wait = get_active_driver(driver, wait)
                progress.update(row, "Tạm bỏ dòng; tiếp tục bài kế tiếp", completed)
                print(
                    f"⏭ [WORKER {worker_id}] Dòng {row}: không xác nhận được prompt; "
                    "tạm bỏ trong lượt này, tiếp tục hàng đợi."
                )
            except Exception as exc:
                write_retry_note(row, 9, "TOTAL", "TOTAL_ERROR", exc)
                write_value(row, COL_DONE, f"Lỗi tổng: {str(exc)[:120]}")
                driver, wait = get_active_driver(driver, wait)
                progress.update(row, f"Lỗi: {str(exc)[:100]}", completed)
            finally:
                TASK_QUEUE.task_done()
    finally:
        if driver:
            try:
                driver.quit()
            except Exception:
                pass
        UI_QUEUE.put(("STATUS", worker_id, None, "Đã dừng", completed))
        pythoncom.CoUninitialize()


def choose_worker_ids():
    """
    Hiện hộp chọn profile Worker 1-5 trước khi chạy.
    Trả về danh sách ID đã chọn, ví dụ [2, 3].
    Nếu giao diện Tkinter không mở được, chuyển sang nhập dạng 2,3 ở console.
    """
    try:
        import tkinter as tk
        from tkinter import messagebox

        selected = []
        root = tk.Tk()
        root.title(f"{VERSION} - Chọn Worker")
        root.geometry("430x360+120+100")
        root.resizable(False, False)
        root.attributes("-topmost", True)

        tk.Label(
            root,
            text="CHỌN PROFILE WORKER SẼ CHẠY",
            font=("Segoe UI", 13, "bold"),
            pady=12,
        ).pack()
        tk.Label(
            root,
            text="Tick một hoặc nhiều Worker. Mỗi Worker dùng đúng profile cùng số.",
            font=("Segoe UI", 9),
            wraplength=380,
            justify="center",
        ).pack(pady=(0, 8))

        vars_by_id = {}
        options_frame = tk.Frame(root)
        options_frame.pack(fill="x", padx=75, pady=5)
        for worker_id in range(1, 6):
            var = tk.BooleanVar(value=(worker_id <= NUM_WORKERS))
            vars_by_id[worker_id] = var
            tk.Checkbutton(
                options_frame,
                text=f"Worker {worker_id}  →  profile worker_{worker_id}",
                variable=var,
                font=("Segoe UI", 10),
                anchor="w",
            ).pack(fill="x", pady=3)

        select_all_var = tk.BooleanVar(value=False)

        def toggle_all():
            value = bool(select_all_var.get())
            for var in vars_by_id.values():
                var.set(value)

        tk.Checkbutton(
            root,
            text="Chọn tất cả 5 Worker",
            variable=select_all_var,
            command=toggle_all,
            font=("Segoe UI", 9, "bold"),
        ).pack(pady=5)

        def start_selected():
            chosen = [wid for wid, var in vars_by_id.items() if var.get()]
            if not chosen:
                messagebox.showwarning(
                    "Chưa chọn Worker",
                    "Hãy chọn ít nhất một Worker để chạy.",
                    parent=root,
                )
                return
            selected.extend(chosen)
            root.destroy()

        def cancel_run():
            root.destroy()

        buttons = tk.Frame(root)
        buttons.pack(pady=12)
        tk.Button(
            buttons,
            text="BẮT ĐẦU",
            command=start_selected,
            width=14,
            bg="#1f8f4e",
            fg="white",
            font=("Segoe UI", 10, "bold"),
        ).pack(side="left", padx=7)
        tk.Button(
            buttons,
            text="HỦY",
            command=cancel_run,
            width=10,
            font=("Segoe UI", 10),
        ).pack(side="left", padx=7)

        root.protocol("WM_DELETE_WINDOW", cancel_run)
        root.mainloop()
        return selected

    except Exception as gui_error:
        print(f"⚠️ Không mở được hộp chọn Worker: {gui_error}")
        while True:
            try:
                answer = input(
                    "Nhập Worker muốn chạy, cách nhau bằng dấu phẩy "
                    "(ví dụ 2,3 hoặc 1,3,5): "
                ).strip()
            except (EOFError, KeyboardInterrupt):
                return []
            try:
                chosen = sorted({int(item.strip()) for item in answer.split(",") if item.strip()})
            except ValueError:
                chosen = []
            if chosen and all(1 <= worker_id <= 5 for worker_id in chosen):
                return chosen
            print("Lựa chọn không hợp lệ. Chỉ dùng các số 1-5, ví dụ: 2,3")

def main():
    global NUM_WORKERS, SELECTED_WORKER_IDS
    SELECTED_WORKER_IDS = choose_worker_ids()
    if not SELECTED_WORKER_IDS:
        print("Đã hủy chạy chương trình vì chưa chọn Worker.")
        return
    NUM_WORKERS = len(SELECTED_WORKER_IDS)
    worker_text = ", ".join(map(str, SELECTED_WORKER_IDS))
    print(
        f"{VERSION}: nạp Excel một lần, khởi chạy {NUM_WORKERS} Worker Edge "
        f"(profile: {worker_text}) và 1 Worker Word + VBA chuyên dụng."
    )
    rows, task_rows = load_excel_tasks_once()
    if not task_rows:
        print("Không có dòng nào cần xử lý.")
        return

    for worker_id in SELECTED_WORKER_IDS:
        SKIP_PAUSE_EVENTS[worker_id] = threading.Event()
    for row in task_rows:
        TASK_QUEUE.put(row)

    MultiWorkerMonitor(SELECTED_WORKER_IDS)
    writer = threading.Thread(target=excel_writer_thread, name="ExcelWriter")
    writer.start()
    while not _EXCEL_WRITER_READY.wait(timeout=0.5):
        if not writer.is_alive():
            raise RuntimeError("Excel Writer không khởi động được; chưa chạy Worker nào.")
    word_worker = threading.Thread(
        target=word_worker_thread,
        args=(rows,),
        name="WordWorker",
    )
    word_worker.start()
    workers = [
        threading.Thread(
            target=worker_loop,
            args=(worker_id, rows),
            name=f"Worker-{worker_id}",
        )
        for worker_id in SELECTED_WORKER_IDS
    ]
    for worker in workers:
        worker.start()
    for worker in workers:
        worker.join()

    # Worker Edge đã giao hết bài. Chờ Worker Word vét sạch hàng rồi mới dừng writer.
    WORD_QUEUE.join()
    WORD_QUEUE.put(None)
    word_worker.join()

    RESULT_QUEUE.join()
    if writer.is_alive():
        RESULT_QUEUE.put(("STOP",))
        writer.join()
    if _EXCEL_WRITER_FAILED.is_set():
        raise ExcelWriterUnavailableError(
            "Flow đã dừng vì Excel Writer không phục hồi trong 30 giây; "
            "không tiếp tục ghi để tránh nhầm workbook."
        )
    print("V2 CODEX đã kết thúc; toàn bộ Result Queue đã được ghi xuống Excel.")


if __name__ == "__main__":
    main()
